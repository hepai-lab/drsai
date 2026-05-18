"""
docx_processor.py
~~~~~~~~~~~~~~~~~
A lightweight .docx editing utility based on python-docx.
No LLM/agent dependencies — used by the frontend docx editing API.
"""

import os
import re
import logging
from pathlib import Path
from typing import Any, Dict, List, Optional

logger = logging.getLogger(__name__)


# ==================== CJK Font Helpers ====================

_CJK_FONT_KEYWORDS = (
    '宋', '黑', '楷', '仿宋', '隶书', '幼圆', '华文', '微软雅黑',
    'SimSun', 'NSimSun', 'SimHei', 'KaiTi', 'FangSong',
    'STSong', 'STHeiti', 'STKaiti', 'STFangsong', 'STZhongsong',
    'MingLiU', 'PMingLiU', 'MS Mincho', 'MS Gothic',
    'Malgun', 'Batang', 'Gulim', 'DengXian', '等线',
)


def _contains_chinese(text: str) -> bool:
    return bool(re.search(r'[\u4e00-\u9fff]', text))


def _is_cjk_font(font_name: str) -> bool:
    return any(kw in font_name for kw in _CJK_FONT_KEYWORDS)


def _set_east_asia_font(run_element, font_name: str):
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    rPr = run_element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), font_name)


def _apply_font_name_to_run(run, font_name: str):
    if not font_name:
        return
    run.font.name = font_name
    if _is_cjk_font(font_name):
        _set_east_asia_font(run._element, font_name)


def _apply_font_name_to_style(style, font_name: str):
    if not font_name:
        return
    from docx.oxml.ns import qn

    style.font.name = font_name
    if _is_cjk_font(font_name):
        rPr = style.element.get_or_add_rPr()
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is None:
            from docx.oxml import OxmlElement
            rFonts = OxmlElement('w:rFonts')
            rPr.insert(0, rFonts)
        rFonts.set(qn('w:eastAsia'), font_name)


# ==================== DocumentProcessor ====================

class DocumentProcessor:
    """
    Lightweight .docx editor using python-docx.
    Used by the frontend docx editing API to apply structured edits to .docx files.
    """

    def __init__(self, workspace_dir: str):
        self.workspace_dir = Path(workspace_dir)
        self.workspace_dir.mkdir(parents=True, exist_ok=True)

    # ------------------------------------------------------------------ #
    #  Main public API: update_from_html                                   #
    # ------------------------------------------------------------------ #

    def update_from_html(self, file_path: str, html: str) -> Dict[str, Any]:
        """
        Parse HTML from the frontend editor and update the .docx file accordingly.

        This is the main entry point for the frontend editing API. It:
        1. Reads the existing .docx paragraphs
        2. Parses the incoming HTML
        3. Matches paragraphs by index (or by approximate text match)
        4. Updates only changed paragraphs via python-docx

        Args:
            file_path: Path to the .docx file on disk.
            html: HTML string from TipTap editor.

        Returns:
            {"success": True/False, "changes": [...], "message": "..."}
        """
        try:
            import docx
            from docx.oxml.ns import qn

            file_path = Path(file_path)
            if not file_path.exists():
                return {"success": False, "message": f"File not found: {file_path}"}

            doc = docx.Document(file_path)
            changes = []

            # Extract raw paragraphs from HTML
            html_paragraphs = _parse_html_to_paragraphs(html)

            # Match with existing doc paragraphs
            doc_para_count = len(doc.paragraphs)

            for idx, html_para in enumerate(html_paragraphs):
                if idx >= doc_para_count:
                    # HTML has more content than original — append new paragraphs
                    _add_paragraph_from_dict(doc, html_para)
                    changes.append(f"Added new paragraph at index {idx}")
                else:
                    original_text = doc.paragraphs[idx].text
                    new_text = html_para.get("text", "")

                    if original_text != new_text:
                        # Text changed — replace content while preserving original formatting
                        _replace_paragraph_text(doc.paragraphs[idx], new_text, doc)
                        changes.append(f"Updated paragraph {idx}")

            # If HTML has fewer paragraphs than original doc, delete the extras
            if len(html_paragraphs) < doc_para_count:
                excess_start = len(html_paragraphs)
                # Remove excess paragraphs from end to start (reverse order to keep indices stable)
                for i in range(doc_para_count - 1, excess_start - 1, -1):
                    p = doc.paragraphs[i]._element
                    p.getparent().remove(p)
                changes.append(f"Removed {doc_para_count - excess_start} excess paragraph(s)")

            # Save back
            doc.save(file_path)

            return {
                "success": True,
                "message": f"Document updated. {len(changes)} paragraph(s) changed.",
                "changes": changes,
            }

        except Exception as e:
            logger.error(f"update_from_html failed: {e}")
            return {"success": False, "message": str(e)}

    # ------------------------------------------------------------------ #
    #  Structured edit API (same as DocMaster)                           #
    # ------------------------------------------------------------------ #

    def edit_docx(self, file_path: str, edits: List[Dict[str, Any]], overwrite_original: bool = True) -> Dict[str, Any]:
        """
        Apply structured edits to a .docx file.

        Edit types:
          - add_paragraph, add_heading, add_table
          - replace_text, delete_text, delete_paragraph
          - modify_style, format_text, format_paragraph
          - add_page_break, set_table_style
          - add_bullet_list, add_numbered_list
          - insert_image: Insert an image file
              Required: image_path (str). Optional: position (int, default = end), width_inches (float)

        Args:
            file_path: Path to the .docx file.
            edits: List of edit operation dicts.
            overwrite_original: If True, save in-place. If False, save a new file.

        Returns:
            {"success": True/False, "message": "...", "changes": [...], ...}
        """
        try:
            import docx
            from docx.shared import Pt, RGBColor
            from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_BREAK
            from docx.oxml.ns import qn

            file_path = Path(file_path)
            if not file_path.exists():
                return {"success": False, "message": f"File not found: {file_path}"}

            doc = docx.Document(file_path)
            changes_made = []

            # ----- Helpers -----

            def get_alignment_value(alignment):
                mapping = {"left": WD_PARAGRAPH_ALIGNMENT.LEFT,
                           "center": WD_PARAGRAPH_ALIGNMENT.CENTER,
                           "right": WD_PARAGRAPH_ALIGNMENT.RIGHT,
                           "justify": WD_PARAGRAPH_ALIGNMENT.JUSTIFY}
                return mapping.get(str(alignment).lower())

            def normalize_color(color_value):
                if color_value is None:
                    return None
                s = str(color_value).strip().lstrip('#')
                return s.upper() if len(s) == 6 else None

            def set_run(run, formatting):
                if not formatting:
                    return
                font_name = formatting.get("font_name") or formatting.get("font")
                font_size = formatting.get("font_size")
                bold = formatting.get("bold")
                italic = formatting.get("italic")
                underline = formatting.get("underline")
                color = normalize_color(formatting.get("color"))

                if font_name:
                    _apply_font_name_to_run(run, font_name)
                    if _contains_chinese(run.text) and not _is_cjk_font(font_name):
                        _set_east_asia_font(run._element, "宋体")
                if font_size is not None:
                    run.font.size = Pt(font_size)
                if bold is not None:
                    run.font.bold = bold
                if italic is not None:
                    run.font.italic = italic
                if underline is not None:
                    run.font.underline = underline
                if color:
                    run.font.color.rgb = RGBColor.from_string(color)

            def set_paragraph(paragraph, formatting):
                if not formatting:
                    return
                al = get_alignment_value(formatting.get("alignment"))
                if al is not None:
                    paragraph.alignment = al
                if "spacing_before" in formatting:
                    paragraph.paragraph_format.space_before = Pt(formatting["spacing_before"])
                if "spacing_after" in formatting:
                    paragraph.paragraph_format.space_after = Pt(formatting["spacing_after"])
                if "line_spacing" in formatting:
                    paragraph.paragraph_format.line_spacing = formatting["line_spacing"]

            def insert_para(content, position, style=None):
                if isinstance(position, int) and 0 <= position < len(doc.paragraphs):
                    p = doc.paragraphs[position].insert_paragraph_before(content)
                else:
                    p = doc.add_paragraph(content)
                if style:
                    p.style = style
                return p

            def iter_table_paragraphs():
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for p in cell.paragraphs:
                                yield p

            def replace_para_text(paragraph, old_text, new_text):
                """Replace old_text with new_text in a paragraph, preserving run formatting."""
                if not old_text or old_text not in paragraph.text:
                    return 0

                if not paragraph.runs:
                    paragraph.text = paragraph.text.replace(old_text, new_text)
                    return 1

                from docx.oxml import OxmlElement

                source_run = next((r for r in paragraph.runs if old_text in r.text), None)
                if not source_run:
                    paragraph.text = paragraph.text.replace(old_text, new_text)
                    return 1

                run_text = source_run.text
                idx = run_text.find(old_text)

                run_elem = source_run._element
                p_elem = run_elem.getparent()

                before = run_text[:idx]
                after = run_text[idx + len(old_text):]

                # Create "after" run
                new_run = OxmlElement("w:r")
                rPr = run_elem.find(qn("w:rPr"))
                if rPr is not None:
                    new_run.append(rPr.copy())
                t_after = OxmlElement("w:t")
                t_after.set(qn("xml:space"), "preserve")
                t_after.text = after
                new_run.append(t_after)

                # Update source run to [before + new_text]
                t_src = run_elem.find(qn("w:t"))
                if t_src is None:
                    t_src = OxmlElement("w:t")
                    t_src.set(qn("xml:space"), "preserve")
                    run_elem.append(t_src)
                t_src.text = before + new_text

                p_elem.insert(list(p_elem).index(run_elem) + 1, new_run)
                return 1

            # ----- Process each edit -----
            for edit in edits:
                edit_type = edit.get("type", "add_paragraph")

                try:
                    if edit_type == "add_paragraph":
                        p = insert_para(edit.get("content", ""), edit.get("position"))
                        set_paragraph(p, edit)
                        for run in p.runs:
                            set_run(run, edit)
                        changes_made.append(f"Added paragraph: {edit.get('content', '')[:30]}")

                    elif edit_type == "add_heading":
                        level = int(edit.get("level", 1))
                        style_name = "Title" if level <= 0 else f"Heading {min(level, 9)}"
                        p = insert_para(edit.get("content", ""), edit.get("position"), style_name)
                        set_paragraph(p, edit)
                        for run in p.runs:
                            set_run(run, edit)
                        changes_made.append(f"Added heading (level {level}): {edit.get('content', '')}")

                    elif edit_type == "replace_text":
                        cnt = 0
                        for p in list(doc.paragraphs) + list(iter_table_paragraphs()):
                            cnt += replace_para_text(p, edit.get("old_text", ""), edit.get("new_text", ""))
                        changes_made.append(f"Replaced '{edit.get('old_text', '')}' in {cnt} place(s)")

                    elif edit_type == "delete_text":
                        target = edit.get("old_text", "")
                        if target:
                            for p in doc.paragraphs:
                                if target in p.text:
                                    p.text = p.text.replace(target, "")
                                    changes_made.append(f"Deleted '{target}'")
                                    break

                    elif edit_type == "delete_paragraph":
                        pos = edit.get("position")
                        count = edit.get("count", 1)
                        if isinstance(pos, int) and 0 <= pos < len(doc.paragraphs):
                            for i in range(count):
                                idx_to_remove = pos
                                if idx_to_remove < len(doc.paragraphs):
                                    p_elem = doc.paragraphs[idx_to_remove]._element
                                    p_elem.getparent().remove(p_elem)
                            changes_made.append(f"Deleted {count} paragraph(s) at position {pos}")

                    elif edit_type == "modify_style":
                        style_name = edit.get("style_name", "Normal")
                        try:
                            style = doc.styles[style_name]
                        except KeyError:
                            from docx.enum.style import WD_STYLE_TYPE
                            style = doc.styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
                        _apply_font_name_to_style(style, edit.get("font_name") or edit.get("font") or "")
                        if edit.get("font_size") is not None:
                            style.font.size = Pt(edit["font_size"])
                        al = get_alignment_value(edit.get("alignment"))
                        if al is not None:
                            style.paragraph_format.alignment = al
                        changes_made.append(f"Modified style '{style_name}'")

                    elif edit_type == "add_table":
                        data = edit.get("data", [])
                        rows = len(data) if data else edit.get("rows", 1)
                        cols = max((len(r) if isinstance(r, (list, tuple)) else 1) for r in data) if data else edit.get("cols", 1)
                        table = doc.add_table(rows=rows, cols=cols)
                        for ri, row in enumerate(data or []):
                            row_data = list(row) if isinstance(row, (list, tuple)) else [str(row)]
                            for ci, val in enumerate(row_data[:cols]):
                                table.cell(ri, ci).text = str(val)
                        if edit.get("table_style") or edit.get("style"):
                            table.style = edit.get("table_style") or edit["style"]
                        changes_made.append(f"Added {rows}x{cols} table")

                    elif edit_type == "add_page_break":
                        pos = edit.get("position")
                        if isinstance(pos, int) and 0 <= pos < len(doc.paragraphs):
                            p = doc.paragraphs[pos].insert_paragraph_before("")
                        else:
                            p = doc.add_paragraph()
                        p.add_run().add_break(WD_BREAK.PAGE)
                        changes_made.append("Added page break")

                    elif edit_type == "set_table_style":
                        ti = edit.get("table_index", 0)
                        ts = edit.get("table_style") or edit.get("style")
                        if 0 <= ti < len(doc.tables) and ts:
                            doc.tables[ti].style = ts
                            changes_made.append(f"Applied style '{ts}' to table {ti}")

                    elif edit_type == "add_bullet_list":
                        items = edit.get("items", [])
                        position = edit.get("position")
                        for item in items:
                            text = item if isinstance(item, str) else item.get("text", "")
                            level = item.get("level", 0) if isinstance(item, dict) else 0
                            if not text:
                                continue
                            if isinstance(position, int) and 0 <= position < len(doc.paragraphs):
                                p = doc.paragraphs[position].insert_paragraph_before(text)
                            else:
                                p = doc.add_paragraph(text)
                            style_name = f"List Bullet {level + 1}" if level > 0 else "List Bullet"
                            try:
                                p.style = doc.styles[style_name]
                            except KeyError:
                                try:
                                    p.style = doc.styles["List Bullet"]
                                except KeyError:
                                    p.style = doc.styles["Normal"]
                            if isinstance(position, int):
                                position += 1
                        changes_made.append(f"Added bullet list with {len(items)} item(s)")

                    elif edit_type == "add_numbered_list":
                        items = edit.get("items", [])
                        position = edit.get("position")
                        for item in items:
                            text = item if isinstance(item, str) else item.get("text", "")
                            level = item.get("level", 0) if isinstance(item, dict) else 0
                            if not text:
                                continue
                            if isinstance(position, int) and 0 <= position < len(doc.paragraphs):
                                p = doc.paragraphs[position].insert_paragraph_before(text)
                            else:
                                p = doc.add_paragraph(text)
                            style_name = f"List Number {level + 1}" if level > 0 else "List Number"
                            try:
                                p.style = doc.styles[style_name]
                            except KeyError:
                                try:
                                    p.style = doc.styles["List Number"]
                                except KeyError:
                                    p.style = doc.styles["Normal"]
                            if isinstance(position, int):
                                position += 1
                        changes_made.append(f"Added numbered list with {len(items)} item(s)")

                    elif edit_type == "format_paragraph":
                        pos = edit.get("position")
                        if isinstance(pos, int) and 0 <= pos < len(doc.paragraphs):
                            set_paragraph(doc.paragraphs[pos], edit)
                            changes_made.append(f"Formatted paragraph {pos}")

                    elif edit_type == "insert_image":
                        image_path = edit.get("image_path")
                        if not image_path:
                            changes_made.append("Error in insert_image: image_path is required")
                        elif not os.path.isfile(image_path):
                            changes_made.append(f"Error in insert_image: file not found: {image_path}")
                        else:
                            try:
                                from docx.shared import Inches
                                insert_pos = edit.get("position")
                                width_inches = edit.get("width_inches")
                                
                                # Determine insertion point
                                if isinstance(insert_pos, int) and 0 <= insert_pos < len(doc.paragraphs):
                                    p = doc.paragraphs[insert_pos].insert_paragraph_before("")
                                else:
                                    p = doc.add_paragraph()
                                
                                run = p.add_run()
                                if width_inches:
                                    run.add_picture(image_path, width=Inches(width_inches))
                                else:
                                    run.add_picture(image_path)
                                
                                changes_made.append(f"Inserted image: {os.path.basename(image_path)}")
                            except Exception as img_err:
                                changes_made.append(f"Error in insert_image: {img_err}")

                except Exception as e:
                    changes_made.append(f"Error in {edit_type}: {e}")

            # ----- Save -----
            if overwrite_original:
                doc.save(file_path)
            else:
                timestamp = __import__("datetime").datetime.now().strftime("%Y%m%d_%H%M%S")
                doc.save(file_path.parent / f"{file_path.stem}_edited_{timestamp}.docx")

            return {
                "success": True,
                "message": f"Edits applied. {len(changes_made)} change(s).",
                "changes_made": changes_made,
                "file_path": str(file_path),
            }

        except Exception as e:
            logger.error(f"edit_docx failed: {e}")
            return {"success": False, "message": str(e)}


# ==================== HTML Parsing Helpers ====================

def _parse_html_to_paragraphs(html: str) -> List[Dict[str, Any]]:
    """
    Very lightweight HTML parser — extracts text and basic inline formatting.

    Converts HTML paragraphs/rdings to a list of dicts like:
      {"text": "...", "bold": True, "italic": False, "alignment": "left"}

    Limitation: Only handles inline <b>, <i>, <strong>, <em> and headings.
    Tables, images, and complex formatting are stripped.
    """
    import re

    paragraphs = []

    # Strip HTML comments
    html = re.sub(r"<!--.*?-->", "", html, flags=re.DOTALL)

    # Remove scripts and styles
    html = re.sub(r"<script[^>]*>.*?</script>", "", html, flags=re.DOTALL | re.IGNORECASE)
    html = re.sub(r"<style[^>]*>.*?</style>", "", html, flags=re.DOTALL | re.IGNORECASE)

    # Split block-level elements into "paragraphs"
    block_pattern = re.compile(
        r"<(p|div|h[1-6]|li|blockquote|tr)(\s[^>]*)?>(.*?)</\1>",
        re.DOTALL | re.IGNORECASE,
    )

    for match in block_pattern.finditer(html):
        tag = match.group(1).lower()
        inner = match.group(3)

        # Extract text (strip all inner tags)
        text = re.sub(r"<[^>]+>", "", inner).strip()
        if not text:
            continue

        result: Dict[str, Any] = {"text": text}

        # Detect heading level
        if re.match(r"^h[1-6]$", tag):
            result["heading_level"] = int(tag[1])

        # Detect inline formatting
        result["bold"] = bool(re.search(r"<b|strong", inner, re.IGNORECASE))
        result["italic"] = bool(re.search(r"<i|em", inner, re.IGNORECASE))

        # Detect alignment from parent style
        parent_style = match.group(2) or ""
        if "center" in parent_style or "text-align:center" in parent_style:
            result["alignment"] = "center"
        elif "right" in parent_style or "text-align:right" in parent_style:
            result["alignment"] = "right"
        elif "justify" in parent_style or "text-align:justify" in parent_style:
            result["alignment"] = "justify"

        paragraphs.append(result)

    return paragraphs


def _add_paragraph_from_dict(doc, para_dict: Dict[str, Any]):
    """Add a new paragraph to a docx document from a parsed paragraph dict."""
    import docx
    from docx.shared import Pt, RGBColor

    level = para_dict.get("heading_level")
    if level:
        style_name = "Title" if level <= 0 else f"Heading {min(level, 9)}"
        p = doc.add_heading(para_dict["text"], level if level > 0 else 0)
    else:
        p = doc.add_paragraph(para_dict["text"])

    # Apply basic inline formatting to all runs
    if para_dict.get("bold") or para_dict.get("italic"):
        if not p.runs:
            p.add_run(para_dict["text"])
        for run in p.runs:
            if para_dict.get("bold"):
                run.font.bold = True
            if para_dict.get("italic"):
                run.font.italic = True


def _safe_get_font_prop(font, attr_name: str):
    """Safely get a font property, swallowing python-docx CT_RPr/internal errors."""
    try:
        return getattr(font, attr_name)
    except AttributeError:
        # CT_RPr or other internal python-docx errors on malformed XML
        return None


def _safe_set_font_prop(run, attr_name: str, value):
    """Safely set a font property, swallowing python-docx CT_RPr/internal errors."""
    if value is None:
        return
    try:
        setattr(run.font, attr_name, value)
    except (AttributeError, TypeError):
        pass


def _replace_paragraph_text(paragraph, new_text: str, doc, formatting=None):
    """Replace a paragraph's text while preserving the original paragraph style and font.
    
    Args:
        paragraph: The docx paragraph to modify
        new_text: The new text content
        doc: The docx document
        formatting: Optional dict with bold, italic, underline from TipTap editor
    """
    from docx.oxml.ns import qn
    
    p_elem = paragraph._element
    
    # Store the original run formatting if it exists (safely, CT_RPr-safe)
    original_font_name = None
    original_font_size = None
    original_font_bold = None
    original_font_italic = None
    
    if paragraph.runs:
        orig_run = paragraph.runs[0]
        original_font_name = _safe_get_font_prop(orig_run.font, "name")
        original_font_size = _safe_get_font_prop(orig_run.font, "size")
        original_font_bold = _safe_get_font_prop(orig_run.font, "bold")
        original_font_italic = _safe_get_font_prop(orig_run.font, "italic")

    # Remove all existing <w:r> children
    for r in p_elem.findall(qn("w:r")):
        p_elem.remove(r)

    # Remove all <w:hyperlink> children (keep the paragraph itself)
    for hl in p_elem.findall(qn("w:hyperlink")):
        p_elem.remove(hl)

    # Reset to plain text (creates new run with default formatting)
    paragraph.text = new_text
    
    # Apply formatting to the new run
    if paragraph.runs:
        new_run = paragraph.runs[0]
        
        # Priority: TipTap formatting > Original formatting > Default
        
        # Bold: TipTap takes priority
        if formatting and formatting.get("bold") is not None:
            _safe_set_font_prop(new_run, "bold", formatting["bold"])
        elif original_font_bold is not None:
            _safe_set_font_prop(new_run, "bold", original_font_bold)
        
        # Italic: TipTap takes priority
        if formatting and formatting.get("italic") is not None:
            _safe_set_font_prop(new_run, "italic", formatting["italic"])
        elif original_font_italic is not None:
            _safe_set_font_prop(new_run, "italic", original_font_italic)
        
        # Font name and size: always preserve from original
        if original_font_name:
            _safe_set_font_prop(new_run, "name", original_font_name)
        if original_font_size:
            _safe_set_font_prop(new_run, "size", original_font_size)


def _copy_run_formatting(source_para, target_para, doc):
    """
    Copy run-level formatting from source paragraph to target paragraph.
    This includes font size, font name, bold, italic, etc.
    """
    logger.info(f"[FontDebug] _copy_run_formatting called")
    logger.info(f"[FontDebug] Source para text: {source_para.text[:50] if source_para.text else 'empty'}")
    logger.info(f"[FontDebug] Target para text: {target_para.text[:50] if target_para.text else 'empty'}")
    
    if not source_para.runs:
        logger.info("[FontDebug] No runs in source paragraph")
        return
    
    # Get formatting from first run (safely, CT_RPr-safe)
    sample_run = source_para.runs[0]
    sample_name = _safe_get_font_prop(sample_run.font, "name")
    sample_size = _safe_get_font_prop(sample_run.font, "size")
    sample_bold = _safe_get_font_prop(sample_run.font, "bold")
    sample_italic = _safe_get_font_prop(sample_run.font, "italic")
    logger.info(f"[FontDebug] Sample run font: name={sample_name}, size={sample_size}, bold={sample_bold}")
    
    # Copy paragraph formatting
    pf_source = source_para.paragraph_format
    pf_target = target_para.paragraph_format
    
    try:
        pf_target.alignment = pf_source.alignment
        pf_target.space_before = pf_source.space_before
        pf_target.space_after = pf_source.space_after
        pf_target.line_spacing = pf_source.line_spacing
    except Exception as e:
        logger.warning(f"[FontDebug] Error copying paragraph format: {e}")
    
    # Copy run formatting to the new paragraph's runs
    if target_para.runs:
        target_run = target_para.runs[0]
        logger.info(f"[FontDebug] Target run before: name={_safe_get_font_prop(target_run.font, 'name')}, size={_safe_get_font_prop(target_run.font, 'size')}")
        _safe_set_font_prop(target_run, "name", sample_name)
        _safe_set_font_prop(target_run, "size", sample_size)
        if sample_bold is not None:
            _safe_set_font_prop(target_run, "bold", sample_bold)
        if sample_italic is not None:
            _safe_set_font_prop(target_run, "italic", sample_italic)
        logger.info(f"[FontDebug] Target run after: name={_safe_get_font_prop(target_run.font, 'name')}, size={_safe_get_font_prop(target_run.font, 'size')}")
    else:
        logger.info("[FontDebug] No runs in target paragraph")


# ==================== Content-Based Paragraph Matching ====================

def calculate_similarity(str1: str, str2: str) -> float:
    """
    Calculate similarity between two strings using a combination of
    character-based Jaccard similarity and word-based overlap.
    
    Returns a float between 0.0 (completely different) and 1.0 (identical).
    
    This enables fuzzy paragraph matching based on content rather than position.
    """
    if not str1 or not str2:
        return 0.0
    
    # Normalize: lowercase, strip whitespace
    s1 = str1.strip().lower()
    s2 = str2.strip().lower()
    
    # Exact match
    if s1 == s2:
        return 1.0
    
    # Character-based Jaccard similarity
    set1 = set(s1)
    set2 = set(s2)
    intersection = len(set1 & set2)
    union = len(set1 | set2)
    char_similarity = (2.0 * intersection) / union if union > 0 else 0.0
    
    # Word-based similarity (for better handling of punctuation/digit changes)
    import re
    words1 = set(re.findall(r'[\u4e00-\u9fff]+|[a-zA-Z0-9]+', s1))
    words2 = set(re.findall(r'[\u4e00-\u9fff]+|[a-zA-Z0-9]+', s2))
    
    if not words1 and not words2:
        return char_similarity
    
    if not words1 or not words2:
        return char_similarity * 0.8  # Slight penalty for word mismatch
    
    word_intersection = len(words1 & words2)
    word_union = len(words1 | words2)
    word_similarity = (2.0 * word_intersection) / word_union if word_union > 0 else 0.0
    
    # Weighted average (favor word similarity for text content)
    return 0.4 * char_similarity + 0.6 * word_similarity


def find_best_match(paragraph: str, original_list: List[str], threshold: float = 0.7) -> tuple:
    """
    Find the best matching paragraph from a list of original paragraphs.
    
    Args:
        paragraph: The paragraph text to match.
        original_list: List of original paragraph texts.
        threshold: Minimum similarity score (0.0-1.0) to consider a match.
                   Default 0.7 means 70% similar is required.
    
    Returns:
        Tuple of (best_index, best_similarity_score) or (None, 0.0) if no match.
    
    Uses greedy matching - finds the most similar paragraph above threshold.
    """
    best_index = None
    best_score = threshold  # Only accept matches above this
    
    for idx, original_text in enumerate(original_list):
        score = calculate_similarity(paragraph, original_text)
        if score > best_score:
            best_score = score
            best_index = idx
    
    return best_index, best_score


def edit_docx_by_content_match(
    file_path: str,
    edits: List[Dict[str, Any]],
    original_paragraphs: List[str],
    preserve_format: bool = True
) -> Dict[str, Any]:
    """
    Apply structured edits to a .docx file using content-based paragraph matching.
    
    This function extends edit_docx() by adding content-based matching.
    When an edit references a paragraph by content (old_text), it finds
    the matching paragraph in the document by similarity, not by position.
    
    Edit types supported:
      - add_paragraph: Add new paragraph at specified content position
      - replace_text: Replace text with content-based matching
      - delete_paragraph: Delete paragraph matched by content
      - format_text: Format text (bold/italic/etc) by content match
      - format_paragraph: Format paragraph by content match
      - insert_image: Insert an image file at a specified position
          Required fields: image_path (str, path to image file)
          Optional fields: insert_after_paragraph (int, default -1 = end), width_inches (float)
    
    Args:
        file_path: Path to the .docx file.
        edits: List of edit operation dicts.
        original_paragraphs: List of original paragraph texts for matching.
        preserve_format: If True, preserve original paragraph formatting.
    
    Returns:
        {"success": True/False, "message": "...", "changes_made": [...], "file_path": "..."}
    """
    import docx
    from docx.shared import Pt, RGBColor
    from docx.oxml.ns import qn
    
    try:
        file_path = Path(file_path)
        if not file_path.exists():
            return {"success": False, "message": f"File not found: {file_path}"}
        
        doc = docx.Document(file_path)
        changes_made = []
        
        # Build current document paragraph text list
        current_paragraphs = [p.text for p in doc.paragraphs]
        
        # Track which paragraphs have been matched/edited
        matched_original_indices = set()
        matched_current_indices = set()
        
        # Pre-process: find matches for all edits that reference content
        paragraph_matches = {}  # edit_index -> (original_idx, current_idx, similarity)
        
        for edit_idx, edit in enumerate(edits):
            edit_type = edit.get("type", "")
            
            if edit_type in ("replace_text", "format_text"):
                old_text = edit.get("old_text", "")
                if old_text:
                    # Find in original paragraphs
                    orig_idx, orig_score = find_best_match(old_text, original_paragraphs)
                    # Find in current paragraphs
                    curr_idx, curr_score = find_best_match(old_text, current_paragraphs)
                    
                    if orig_idx is not None and curr_idx is not None:
                        # Use the better match
                        paragraph_matches[edit_idx] = {
                            "original_idx": orig_idx,
                            "current_idx": curr_idx,
                            "score": max(orig_score, curr_score)
                        }
                        matched_original_indices.add(orig_idx)
                        matched_current_indices.add(curr_idx)
                    elif orig_idx is not None:
                        paragraph_matches[edit_idx] = {
                            "original_idx": orig_idx,
                            "current_idx": None,
                            "score": orig_score
                        }
                        matched_original_indices.add(orig_idx)
                    elif curr_idx is not None:
                        paragraph_matches[edit_idx] = {
                            "original_idx": None,
                            "current_idx": curr_idx,
                            "score": curr_score
                        }
                        matched_current_indices.add(curr_idx)
            
            elif edit_type in ("delete_paragraph", "format_paragraph"):
                # Try to match by content or position
                target_text = edit.get("text", edit.get("content", ""))
                position = edit.get("position")
                
                if target_text:
                    orig_idx, orig_score = find_best_match(target_text, original_paragraphs)
                    curr_idx, curr_score = find_best_match(target_text, current_paragraphs)
                    
                    if orig_idx is not None and curr_idx is not None:
                        paragraph_matches[edit_idx] = {
                            "original_idx": orig_idx,
                            "current_idx": curr_idx,
                            "score": max(orig_score, curr_score)
                        }
                        matched_original_indices.add(orig_idx)
                        matched_current_indices.add(curr_idx)
                elif isinstance(position, int):
                    # Fallback to position-based matching
                    if 0 <= position < len(doc.paragraphs):
                        paragraph_matches[edit_idx] = {
                            "original_idx": position,
                            "current_idx": position,
                            "score": 1.0
                        }
                        matched_original_indices.add(position)
                        matched_current_indices.add(position)
        
        # Process edits
        for edit_idx, edit in enumerate(edits):
            edit_type = edit.get("type", "")
            match_info = paragraph_matches.get(edit_idx, {})
            target_curr_idx = match_info.get("current_idx")
            
            try:
                if edit_type == "add_paragraph":
                    content = edit.get("content", "")
                    position = edit.get("position")
                    formatting = edit.get("formatting", {})
                    logger.info(f"[FontDebug] add_paragraph called: content='{content[:50]}...', position={position}, formatting={formatting}")
                    
                    if isinstance(position, int) and 0 <= position < len(doc.paragraphs):
                        # Insert before the paragraph at position, copy its style
                        ref_paragraph = doc.paragraphs[position]
                        ref_style = ref_paragraph.style
                        logger.info(f"[FontDebug] Reference para: text='{ref_paragraph.text[:30]}...', style={ref_style}")
                        
                        # Insert new paragraph
                        p = ref_paragraph.insert_paragraph_before()
                        p.text = content
                        logger.info(f"[FontDebug] New para text set to: '{content[:30]}...'")
                        
                        # Copy style from reference paragraph
                        try:
                            if ref_style:
                                p.style = ref_style
                                logger.info(f"[FontDebug] Copied style: {ref_style}")
                            # Also try to copy direct formatting
                            _copy_run_formatting(ref_paragraph, p, doc)
                        except Exception as e:
                            logger.warning(f"[FontDebug] Error copying style: {e}")
                    else:
                        p = doc.add_paragraph(content)
                        logger.info(f"[FontDebug] Added paragraph at end (no position match)")
                    
                    # Apply TipTap formatting if provided
                    if formatting and p.runs:
                        try:
                            run = p.runs[0]
                            if formatting.get("bold") is not None:
                                run.font.bold = formatting["bold"]
                            if formatting.get("italic") is not None:
                                run.font.italic = formatting["italic"]
                        except Exception:
                            pass
                    
                    changes_made.append(f"Added paragraph: {content[:30]}...")
                
                elif edit_type == "replace_text":
                    logger.info(f"[FontDebug] replace_text called: old='{edit.get('old_text', '')[:30]}...', new='{edit.get('new_text', '')[:30]}...'")
                    
                    # Get formatting from edit (if provided)
                    formatting = edit.get("formatting", {})
                    logger.info(f"[FontDebug] Formatting from edit: {formatting}")
                    
                    if target_curr_idx is not None and 0 <= target_curr_idx < len(doc.paragraphs):
                        new_text = edit.get("new_text", "")
                        p = doc.paragraphs[target_curr_idx]
                        logger.info(f"[FontDebug] Target paragraph before replace: text='{p.text[:30]}...', runs={len(p.runs)}")
                        
                        # Preserve paragraph style and apply formatting
                        _replace_paragraph_text(p, new_text, doc, formatting=formatting)
                        
                        logger.info(f"[FontDebug] After replace: text='{p.text[:30]}...', runs={len(p.runs)}")
                        if p.runs:
                            run = p.runs[0]
                            logger.info(f"[FontDebug] New run font: name={_safe_get_font_prop(run.font, 'name')}, bold={_safe_get_font_prop(run.font, 'bold')}, italic={_safe_get_font_prop(run.font, 'italic')}")
                        
                        changes_made.append(f"Replaced text at paragraph {target_curr_idx}")
                    else:
                        # Fallback: find and replace first occurrence
                        old_text = edit.get("old_text", "")
                        new_text = edit.get("new_text", "")
                        replaced = False
                        for p in doc.paragraphs:
                            if old_text in p.text:
                                logger.info(f"[FontDebug] Fallback replace: para text='{p.text[:30]}...'")
                                _replace_paragraph_text(p, p.text.replace(old_text, new_text), doc, formatting=formatting)
                                changes_made.append(f"Replaced text: {old_text[:20]}...")
                                replaced = True
                                break
                        if not replaced:
                            changes_made.append(f"Warning: Could not find text to replace")
                
                elif edit_type == "delete_paragraph":
                    if target_curr_idx is not None and 0 <= target_curr_idx < len(doc.paragraphs):
                        p_elem = doc.paragraphs[target_curr_idx]._element
                        p_elem.getparent().remove(p_elem)
                        changes_made.append(f"Deleted paragraph {target_curr_idx}")
                    else:
                        position = edit.get("position")
                        if isinstance(position, int) and 0 <= position < len(doc.paragraphs):
                            p_elem = doc.paragraphs[position]._element
                            p_elem.getparent().remove(p_elem)
                            changes_made.append(f"Deleted paragraph at position {position}")
                
                elif edit_type == "format_text":
                    if target_curr_idx is not None and 0 <= target_curr_idx < len(doc.paragraphs):
                        p = doc.paragraphs[target_curr_idx]
                        formatting = edit.get("formatting", edit)  # Support both nested and flat format
                        for run in p.runs:
                            if formatting.get("bold") is not None:
                                _safe_set_font_prop(run, "bold", formatting["bold"])
                            if formatting.get("italic") is not None:
                                _safe_set_font_prop(run, "italic", formatting["italic"])
                            if formatting.get("underline") is not None:
                                _safe_set_font_prop(run, "underline", formatting["underline"])
                        changes_made.append(f"Formatted text at paragraph {target_curr_idx}")
                
                elif edit_type == "format_paragraph":
                    if target_curr_idx is not None and 0 <= target_curr_idx < len(doc.paragraphs):
                        p = doc.paragraphs[target_curr_idx]
                        if edit.get("alignment"):
                            from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
                            al_map = {"left": WD_PARAGRAPH_ALIGNMENT.LEFT,
                                      "center": WD_PARAGRAPH_ALIGNMENT.CENTER,
                                      "right": WD_PARAGRAPH_ALIGNMENT.RIGHT,
                                      "justify": WD_PARAGRAPH_ALIGNMENT.JUSTIFY}
                            alignment = al_map.get(edit["alignment"].lower())
                            if alignment:
                                p.alignment = alignment
                        changes_made.append(f"Formatted paragraph {target_curr_idx}")
                
                elif edit_type == "insert_image":
                    image_path = edit.get("image_path")
                    if not image_path:
                        changes_made.append("Error in insert_image: image_path is required")
                    elif not os.path.isfile(image_path):
                        changes_made.append(f"Error in insert_image: file not found: {image_path}")
                    else:
                        try:
                            from docx.shared import Inches
                            # Support both 'position' (agent convention) and 'insert_after_paragraph' (explicit)
                            raw_pos = edit.get("position", edit.get("insert_after_paragraph", -1))
                            try:
                                insert_after = int(raw_pos)
                            except (TypeError, ValueError):
                                insert_after = -1
                            width_inches = edit.get("width_inches")  # optional, for proportional scaling
                            total_paras = len(doc.paragraphs)
                            
                            if insert_after == -1 or insert_after >= total_paras:
                                # Insert at end of document
                                target_para = doc.add_paragraph()
                                log_pos = f"end (of {total_paras} paragraphs)"
                            else:
                                # Insert after the specified paragraph index
                                target_para = doc.paragraphs[insert_after]
                                # Insert a new empty paragraph after this one via XML
                                from docx.oxml import OxmlElement
                                new_p = OxmlElement("w:p")
                                target_para._element.addnext(new_p)
                                # Re-fetch (paragraphs order may shift after XML insert)
                                target_para = doc.paragraphs[insert_after + 1]
                                log_pos = f"after paragraph {insert_after} (of {total_paras})"
                            
                            # Add image to the target paragraph
                            run = target_para.add_run()
                            if width_inches:
                                run.add_picture(image_path, width=Inches(width_inches))
                            else:
                                run.add_picture(image_path)
                            
                            changes_made.append(f"Inserted image at {log_pos}: {os.path.basename(image_path)}")
                        except Exception as img_err:
                            changes_made.append(f"Error in insert_image: {str(img_err)}")
                
            except Exception as e:
                changes_made.append(f"Error in {edit_type}: {str(e)}")
        
        # Save the document
        doc.save(file_path)
        
        return {
            "success": True,
            "message": f"Edits applied using content-based matching. {len(changes_made)} change(s).",
            "changes_made": changes_made,
            "file_path": str(file_path),
        }
    
    except Exception as e:
        logger.error(f"edit_docx_by_content_match failed: {e}")
        return {"success": False, "message": str(e)}
