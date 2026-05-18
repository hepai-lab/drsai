"""Tests for numbered-list expansion of section_body_empty fills.

When the agent supplies a `1. xxx\\n2. yyy\\n3. zzz` value for an empty
body slot, DocMaster expands it into multiple paragraphs with Word's
auto-numbering applied — so the rendered document carries real list
items (clean hanging indent, editable/reorderable) rather than plain
"1. 2. 3." prefixes baked into the text.
"""
from __future__ import annotations

import sys
import tempfile
from pathlib import Path

_HERE = Path(__file__).resolve().parent
sys.path.insert(0, str(_HERE.parent))

from docx import Document  # type: ignore[import-not-found]
from docx.oxml.ns import qn  # type: ignore[import-not-found]
import docx_template_skill as dts


def _build_empty_section_doc(tmpdir: Path, heading: str = "四、双方协作内容") -> Path:
    d = Document()
    d.add_heading(heading, level=1)
    d.add_paragraph("")  # empty body — the slot target
    d.add_heading("五、其他", level=1)
    path = tmpdir / "tpl.docx"
    d.save(str(path))
    return path


def _num_id_of(paragraph) -> str | None:
    pPr = paragraph._p.pPr
    if pPr is None:
        return None
    numPr = pPr.find(qn("w:numPr"))
    if numPr is None:
        return None
    ni = numPr.find(qn("w:numId"))
    return ni.get(qn("w:val")) if ni is not None else None


def test_numbered_value_expands_into_multiple_numbered_paragraphs():
    tmp = Path(tempfile.mkdtemp())
    tpl = _build_empty_section_doc(tmp)
    skill = dts.DocxTemplateSkill(str(tmp))
    ins = skill.inspect_template(str(tpl))
    sbe = [s for s in ins["slots"] if s["kind"] == "section_body_empty"]
    assert sbe, ins["slots"]
    value = (
        "1. 甲方负责提供技术方案和设计要求。\n"
        "2. 乙方负责详细设计、材料采购、加工制造。\n"
        "3. 乙方应在关键工序节点通知甲方派员到厂见证。\n"
        "4. 甲方应在收到通知后5个工作日内安排人员到厂。"
    )
    out = tmp / "out.docx"
    skill.fill_template(str(tpl), str(out), slot_values={sbe[0]["id"]: value})

    d = Document(str(out))
    # Find the four newly-created list-item paragraphs.
    body_paras = [
        p for p in d.paragraphs
        if p.text.strip() and p.text.strip()[0] not in "四五"
        and "其他" not in p.text
    ]
    assert len(body_paras) == 4, f"expected 4 paragraphs, got {len(body_paras)}: {[p.text for p in body_paras]}"

    # Each must (a) have w:numPr applied and (b) NOT carry the literal
    # "N." text — the number comes from Word's numbering definition.
    for i, p in enumerate(body_paras, start=1):
        nid = _num_id_of(p)
        assert nid is not None, f"para {i} missing numPr: {p.text!r}"
        assert not p.text.lstrip().startswith(f"{i}."), (
            f"para {i} still carries literal '{i}.' prefix: {p.text!r}"
        )

    # All list items share the same numId so Word numbers them continuously.
    nids = {_num_id_of(p) for p in body_paras}
    assert len(nids) == 1, f"expected single numId, got {nids}"


def test_plain_value_falls_back_to_single_paragraph():
    """Non-list values keep the existing single-paragraph behavior — we
    don't want to fragment ordinary prose."""
    tmp = Path(tempfile.mkdtemp())
    tpl = _build_empty_section_doc(tmp)
    skill = dts.DocxTemplateSkill(str(tmp))
    ins = skill.inspect_template(str(tpl))
    sbe = [s for s in ins["slots"] if s["kind"] == "section_body_empty"][0]
    out = tmp / "out.docx"
    skill.fill_template(str(tpl), str(out), slot_values={sbe["id"]: "本节适用于本合同所有条款。"})
    d = Document(str(out))
    body = [p for p in d.paragraphs if "适用于" in p.text]
    assert len(body) == 1
    assert _num_id_of(body[0]) is None, "plain prose should not get numbering"


def test_numbered_list_uses_existing_decimal_definition():
    """When the template already has a decimal '1. 2. 3.' abstractNum, we
    reuse its numId rather than appending a new definition."""
    tmp = Path(tempfile.mkdtemp())
    tpl = _build_empty_section_doc(tmp)
    # Pre-existing num-count
    pre = Document(str(tpl))
    pre_numbering_part = pre.part.numbering_part
    pre_num_count = (
        len(pre_numbering_part.element.findall(qn("w:num")))
        if pre_numbering_part is not None else 0
    )

    skill = dts.DocxTemplateSkill(str(tmp))
    ins = skill.inspect_template(str(tpl))
    sbe = [s for s in ins["slots"] if s["kind"] == "section_body_empty"][0]
    out = tmp / "out.docx"
    skill.fill_template(str(tpl), str(out), slot_values={
        sbe["id"]: "1. a\n2. b\n3. c"
    })
    post = Document(str(out))
    post_numbering_part = post.part.numbering_part
    post_num_count = (
        len(post_numbering_part.element.findall(qn("w:num")))
        if post_numbering_part is not None else 0
    )
    # At most one new num was added (we'd ideally reuse, but a fresh
    # python-docx document has no decimal '%1.' definition, so a new one
    # is created exactly once).
    assert post_num_count - pre_num_count <= 1, (
        f"too many numbering definitions added: pre={pre_num_count} post={post_num_count}"
    )


def test_each_list_gets_its_own_num_id_so_it_restarts():
    """Each section_body_empty list must get its OWN <w:num> reference —
    otherwise Word continues the counter across sections (the agent
    writes `8. a / 9. b` for the second section because it thinks it's
    continuing from the first). With per-list numIds, each rendered list
    starts at 1 regardless of what numbers the agent used."""
    tmp = Path(tempfile.mkdtemp())
    d = Document()
    for heading in ("四、协作", "五、验收", "六、维护"):
        d.add_heading(heading, level=1)
        d.add_paragraph("")
    tpl = tmp / "tpl.docx"
    d.save(str(tpl))

    skill = dts.DocxTemplateSkill(str(tmp))
    ins = skill.inspect_template(str(tpl))
    sbe = [s for s in ins["slots"] if s["kind"] == "section_body_empty"]
    assert len(sbe) == 3
    # Notice section 2 starts at "3." and section 3 at "5." — agents that
    # continued the counter across sections should still produce three
    # independent lists in the rendered output.
    sv = {
        sbe[0]["id"]: "1. a\n2. b",
        sbe[1]["id"]: "3. c\n4. d",
        sbe[2]["id"]: "5. e\n6. f",
    }
    out = tmp / "out.docx"
    skill.fill_template(str(tpl), str(out), slot_values=sv)

    post = Document(str(out))
    list_paras = [p for p in post.paragraphs if _num_id_of(p) is not None]
    assert len(list_paras) == 6, [p.text for p in list_paras]
    nids = [_num_id_of(p) for p in list_paras]
    # Same numId within a list, different numIds between lists.
    assert nids[0] == nids[1]
    assert nids[2] == nids[3]
    assert nids[4] == nids[5]
    assert len(set(nids)) == 3, (
        f"expected 3 distinct numIds (one per list), got {nids}"
    )


def test_each_new_num_carries_start_override_one():
    """Every freshly-created <w:num> must include a <w:lvlOverride
    w:ilvl='0'><w:startOverride w:val='1'/></w:lvlOverride>. Without
    this, some Word versions continue the counter across sibling <w:num>
    elements that share an abstractNum — so section 五's list ends up
    numbered 9, 10, 11, 12 as a continuation of section 二's 1-4 rather
    than restarting at 1."""
    tmp = Path(tempfile.mkdtemp())
    d = Document()
    for heading in ("二、A", "四、B", "五、C"):
        d.add_heading(heading, level=1)
        d.add_paragraph("")
    tpl = tmp / "tpl.docx"
    d.save(str(tpl))
    skill = dts.DocxTemplateSkill(str(tmp))
    ins = skill.inspect_template(str(tpl))
    sbe = [s for s in ins["slots"] if s["kind"] == "section_body_empty"]
    out = tmp / "out.docx"
    skill.fill_template(str(tpl), str(out), slot_values={
        sbe[0]["id"]: "1. a\n2. b",
        sbe[1]["id"]: "5. e\n6. f",   # agent continued count
        sbe[2]["id"]: "9. i\n10. j",
    })
    post = Document(str(out))
    el = post.part.numbering_part.element
    fresh_nums = []
    for num in el.findall(qn("w:num")):
        ov = num.find(qn("w:lvlOverride"))
        if ov is None:
            continue
        so = ov.find(qn("w:startOverride"))
        if so is None:
            continue
        fresh_nums.append((num.get(qn("w:numId")), so.get(qn("w:val"))))
    # The three lists we just inserted should each have startOverride=1.
    assert len(fresh_nums) >= 3, f"expected ≥3 nums with startOverride, got {fresh_nums}"
    for nid, sov in fresh_nums:
        assert sov == "1", f"numId={nid} has startOverride={sov!r}; expected '1'"


def test_mixed_intro_prose_and_list():
    """Section value with an intro line followed by a numbered list — the
    intro becomes a plain paragraph and the list items become
    auto-numbered paragraphs underneath."""
    tmp = Path(tempfile.mkdtemp())
    tpl = _build_empty_section_doc(tmp, heading="一、项目简介及产品清单")
    skill = dts.DocxTemplateSkill(str(tmp))
    ins = skill.inspect_template(str(tpl))
    sbe = [s for s in ins["slots"] if s["kind"] == "section_body_empty"][0]
    out = tmp / "out.docx"
    skill.fill_template(str(tpl), str(out), slot_values={
        sbe["id"]: "本项目采购非标设备1套，产品清单如下：\n"
                   "1. 主真空腔体\n"
                   "2. 真空观察窗\n"
                   "3. 真空法兰接口"
    })
    d = Document(str(out))
    # Intro paragraph: no numbering.
    intro = next(p for p in d.paragraphs if "产品清单如下" in p.text)
    assert _num_id_of(intro) is None
    # List items: numbered, stripped of "N." prefix.
    items = [p for p in d.paragraphs
             if p.text.strip() in ("主真空腔体", "真空观察窗", "真空法兰接口")]
    assert len(items) == 3, [p.text for p in d.paragraphs]
    for p in items:
        assert _num_id_of(p) is not None, p.text


def test_mixed_two_lists_with_headers_between():
    """`甲方：\\n1.…\\n2.…\\n\\n乙方：\\n1.…\\n2.…` — two separate lists,
    each starts at 1 in the rendered doc, headers stay as prose."""
    tmp = Path(tempfile.mkdtemp())
    tpl = _build_empty_section_doc(tmp, heading="四、双方协作内容")
    skill = dts.DocxTemplateSkill(str(tmp))
    ins = skill.inspect_template(str(tpl))
    sbe = [s for s in ins["slots"] if s["kind"] == "section_body_empty"][0]
    out = tmp / "out.docx"
    skill.fill_template(str(tpl), str(out), slot_values={
        sbe["id"]:
            "甲方：\n1. 提供技术方案\n2. 提供场地\n\n"
            "乙方：\n1. 完成设计制造\n2. 提供文件",
    })
    d = Document(str(out))
    paras = [p for p in d.paragraphs if p.text.strip()]

    # Headers (甲方：/乙方：) are prose (no numbering).
    headers = [p for p in paras if p.text.strip() in ("甲方：", "乙方：")]
    assert len(headers) == 2
    for h in headers:
        assert _num_id_of(h) is None, h.text

    # Each list has its own numId so both render starting at 1.
    list_items = [p for p in paras
                  if p.text.strip() in (
                      "提供技术方案", "提供场地",
                      "完成设计制造", "提供文件")]
    assert len(list_items) == 4
    nids = [_num_id_of(p) for p in list_items]
    assert all(n is not None for n in nids)
    # First two share one numId; last two share another.
    assert nids[0] == nids[1]
    assert nids[2] == nids[3]
    assert nids[0] != nids[2], (
        f"expected the two lists to have distinct numIds for restart, got {nids}"
    )


def test_list_with_non_one_start_still_recognized():
    """Agents sometimes forget to restart at 1 (`8. … 9. …`). The list
    must still be recognized and rendered with Word numbering — the
    rendered output restarts at 1 because we hand out a fresh <w:num>."""
    tmp = Path(tempfile.mkdtemp())
    tpl = _build_empty_section_doc(tmp, heading="五、验收")
    skill = dts.DocxTemplateSkill(str(tmp))
    ins = skill.inspect_template(str(tpl))
    sbe = [s for s in ins["slots"] if s["kind"] == "section_body_empty"][0]
    out = tmp / "out.docx"
    skill.fill_template(str(tpl), str(out), slot_values={
        sbe["id"]: "8.验收标准\n9.验收方式\n10.验收期限",
    })
    d = Document(str(out))
    items = [p for p in d.paragraphs
             if p.text.strip() in ("验收标准", "验收方式", "验收期限")]
    assert len(items) == 3, [p.text for p in d.paragraphs]
    for p in items:
        assert _num_id_of(p) is not None
        assert not p.text.strip()[0].isdigit(), (
            f"numeric prefix leaked into text: {p.text!r}"
        )


def test_segment_value_helper():
    s = dts._segment_value
    # Pure prose
    assert s("hello") == [("prose", "hello")]
    # Pure list
    assert s("1. a\n2. b") == [("list", ["a", "b"])]
    # Intro + list
    assert s("intro\n1. a\n2. b") == [("prose", "intro"), ("list", ["a", "b"])]
    # Two lists separated by header
    assert s("甲：\n1. a\n2. b\n\n乙：\n1. c\n2. d") == [
        ("prose", "甲："),
        ("list", ["a", "b"]),
        ("prose", "乙："),
        ("list", ["c", "d"]),
    ]
    # Single-item "list" → folded back into prose
    assert s("1. only one") == [("prose", "1. only one")]


def test_split_numbered_lines_helper():
    """Unit tests for the parser."""
    f = dts._split_numbered_lines
    assert f("1. a\n2. b\n3. c") == ["a", "b", "c"]
    assert f("1、a\n2、b") == ["a", "b"]
    assert f("(1) a\n(2) b") == ["a", "b"]
    # Single line — not a list.
    assert f("1. a") is None
    # Any starting number is OK as long as they increment by 1.
    assert f("2. a\n3. b") == ["a", "b"]
    assert f("8. a\n9. b\n10. c") == ["a", "b", "c"]
    # Non-incrementing.
    assert f("1. a\n3. b") is None
    # Plain prose with no numbering.
    assert f("a\nb") is None
    # Empty.
    assert f("") is None
    assert f(None) is None  # type: ignore[arg-type]


def _run_all():
    failed = []
    fns = [(name, fn) for name, fn in globals().items()
           if name.startswith("test_") and callable(fn)]
    for name, fn in fns:
        try:
            fn()
            print(f"  OK   {name}")
        except AssertionError as exc:
            print(f"  FAIL {name}: {exc}")
            failed.append(name)
        except Exception as exc:
            import traceback
            print(f"  ERR  {name}: {type(exc).__name__}: {exc}")
            traceback.print_exc()
            failed.append(name)
    print()
    if failed:
        print(f"{len(failed)} of {len(fns)} test(s) failed: {failed}")
        return 1
    print(f"All {len(fns)} tests passed.")
    return 0


if __name__ == "__main__":
    raise SystemExit(_run_all())
