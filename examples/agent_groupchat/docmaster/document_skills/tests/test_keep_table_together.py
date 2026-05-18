"""Tests for the keep-table-together pass.

CN ministry templates often place an instruction like
`（全文排版确保下表不跨页。此句话非正文，应删除）` immediately before a
signature/seal table. The instruction itself is a meta-note (deleted on
fill) but encodes a layout intent: the next table must render on a single
page. DocMaster honors this by hardening the table with `w:cantSplit` and
`w:keepNext` before deleting the instruction.
"""
from __future__ import annotations

import sys
import tempfile
from pathlib import Path

_HERE = Path(__file__).resolve().parent
sys.path.insert(0, str(_HERE.parent))

from docx import Document  # type: ignore[import-not-found]
import docx_template_skill as dts

NS = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"


def _build_doc_with_instruction_and_table(tmpdir: Path, instruction: str) -> Path:
    d = Document()
    d.add_paragraph("正文上方一些内容。")
    d.add_paragraph(instruction)
    tbl = d.add_table(rows=3, cols=2)
    for ri in range(3):
        tbl.rows[ri].cells[0].text = f"R{ri}C0"
        tbl.rows[ri].cells[1].text = f"R{ri}C1"
    path = tmpdir / "tpl.docx"
    d.save(str(path))
    return path


def _count_cant_split_rows(tbl) -> int:
    return sum(
        1
        for tr in tbl._tbl.findall(f"{NS}tr")
        if (tp := tr.find(f"{NS}trPr")) is not None
        and tp.find(f"{NS}cantSplit") is not None
    )


def _count_keep_next_paragraphs(tbl) -> int:
    return sum(
        1
        for p in tbl._tbl.iter(f"{NS}p")
        if (pp := p.find(f"{NS}pPr")) is not None
        and pp.find(f"{NS}keepNext") is not None
    )


def test_keep_table_together_applied_when_instruction_removed():
    """`下表不跨页` instruction + removal → next table gets cantSplit on
    every row and keepNext on every paragraph except the last."""
    tmp = Path(tempfile.mkdtemp())
    tpl = _build_doc_with_instruction_and_table(
        tmp, "（全文排版确保下表不跨页。此句话非正文，应删除）"
    )
    skill = dts.DocxTemplateSkill(str(tmp))
    ins = skill.inspect_template(str(tpl))
    rm_ids = [r["id"] for r in ins["removals"]]
    assert rm_ids, f"expected a removal candidate; got {ins['removals']}"

    out = tmp / "out.docx"
    skill.fill_template(str(tpl), str(out), slot_values={}, removal_ids=rm_ids)

    d = Document(str(out))
    assert len(d.tables) == 1
    tbl = d.tables[0]
    n_rows = len(tbl.rows)
    assert _count_cant_split_rows(tbl) == n_rows, (
        f"expected cantSplit on all {n_rows} rows"
    )
    paras = list(tbl._tbl.iter(f"{NS}p"))
    # All paragraphs except the last carry keepNext.
    assert _count_keep_next_paragraphs(tbl) == max(0, len(paras) - 1), (
        f"expected keepNext on {len(paras)-1} paragraphs, got "
        f"{_count_keep_next_paragraphs(tbl)}"
    )
    # And the instruction paragraph is gone.
    body = "\n".join(p.text for p in d.paragraphs)
    assert "下表不跨页" not in body, body


def test_keep_table_together_skipped_for_unrelated_removal():
    """A non-keep-together removal (just `应删除` with no table reference)
    must NOT modify any table on the document."""
    tmp = Path(tempfile.mkdtemp())
    tpl = _build_doc_with_instruction_and_table(
        tmp, "（此句话非正文，应删除）"
    )
    skill = dts.DocxTemplateSkill(str(tmp))
    ins = skill.inspect_template(str(tpl))
    rm_ids = [r["id"] for r in ins["removals"]]
    assert rm_ids

    out = tmp / "out.docx"
    skill.fill_template(str(tpl), str(out), slot_values={}, removal_ids=rm_ids)
    d = Document(str(out))
    tbl = d.tables[0]
    assert _count_cant_split_rows(tbl) == 0, "no cantSplit expected"
    assert _count_keep_next_paragraphs(tbl) == 0, "no keepNext expected"


def test_keep_table_together_english_phrasing():
    """English `keep table together` / `keep following table on one page`
    is also honored."""
    tmp = Path(tempfile.mkdtemp())
    tpl = _build_doc_with_instruction_and_table(
        tmp, "(Keep the following table on one page — delete this note before publishing.)"
    )
    skill = dts.DocxTemplateSkill(str(tmp))
    ins = skill.inspect_template(str(tpl))
    rm_ids = [r["id"] for r in ins["removals"]]
    assert rm_ids

    out = tmp / "out.docx"
    skill.fill_template(str(tpl), str(out), slot_values={}, removal_ids=rm_ids)
    d = Document(str(out))
    tbl = d.tables[0]
    assert _count_cant_split_rows(tbl) == len(tbl.rows)


def test_apply_keep_table_together_helper_is_idempotent():
    """Applying the pass twice doesn't duplicate cantSplit / keepNext."""
    tmp = Path(tempfile.mkdtemp())
    d = Document()
    d.add_paragraph("Header")
    d.add_table(rows=2, cols=2)
    path = tmp / "doc.docx"
    d.save(str(path))
    d = Document(str(path))
    tbl = d.tables[0]
    assert dts._apply_keep_table_together(tbl._tbl)
    assert dts._apply_keep_table_together(tbl._tbl)
    # Verify exactly one cantSplit per row.
    for tr in tbl._tbl.findall(f"{NS}tr"):
        trPr = tr.find(f"{NS}trPr")
        assert trPr is not None
        assert len(trPr.findall(f"{NS}cantSplit")) == 1


def _run_all():
    failed = []
    fns = [(name, fn) for name, fn in globals().items()
           if name.startswith("test_") and callable(fn)]
    for name, fn in fns:
        try:
            fn()
            print(f"  OK   {name}")
        except AssertionError as exc:
            print(f"  FAIL {name}: {exc}")
            failed.append(name)
        except Exception as exc:
            import traceback
            print(f"  ERR  {name}: {type(exc).__name__}: {exc}")
            traceback.print_exc()
            failed.append(name)
    print()
    if failed:
        print(f"{len(failed)} of {len(fns)} test(s) failed: {failed}")
        return 1
    print(f"All {len(fns)} tests passed.")
    return 0


if __name__ == "__main__":
    raise SystemExit(_run_all())
