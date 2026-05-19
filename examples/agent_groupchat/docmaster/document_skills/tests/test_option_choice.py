"""Tests for the 二选一 / 三选一 option_choice slot kind.

Plain-assert style so it runs with `python test_option_choice.py` without
pytest. Pytest will still discover the `test_*` functions if available.
"""
from __future__ import annotations

import sys
import tempfile
from pathlib import Path

_HERE = Path(__file__).resolve().parent
sys.path.insert(0, str(_HERE.parent))

from docx import Document  # type: ignore[import-not-found]
import docx_template_skill as dts


def _make_two_options_template(tmpdir: Path) -> Path:
    """Build an in-memory template mimicking the 技术成果归属 section from a
    real Chinese 非标 contract: a heading, a 'choose-one' instruction, two
    `（●第N种：…）` headers each with multiple body paragraphs, and a
    trailing section heading so body collection terminates cleanly."""
    d = Document()
    d.add_heading("八、技术成果归属", level=1)
    d.add_paragraph(
        "（以下两种选择适合的一种，可以根据需要再进行改写，不选的一种请删除）"
    )
    d.add_paragraph("（●第一种：方案A 标题）")
    d.add_paragraph("8.1 合同履行中产生的知识产权归甲方所有。")
    d.add_paragraph("8.2 乙方保证不在其他场合复用本合同的技术成果。")
    d.add_paragraph("（●第二种：方案B 标题）")
    d.add_paragraph("8.1 合同履行过程中产生的知识产权由双方共有。")
    d.add_paragraph("8.2 本合同签订前双方各自拥有的知识产权归各自所有。")
    d.add_paragraph("8.3 甲方提供的所有技术信息均视为甲方的专有信息。")
    d.add_heading("九、违约责任", level=1)
    d.add_paragraph("9.1 任何一方违反本合同应承担违约责任。")
    path = tmpdir / "tpl.docx"
    d.save(str(path))
    return path


def _inspect_oc(skill, tpl_path):
    ins = skill.inspect_template(str(tpl_path))
    return ins, [s for s in ins["slots"] if s["kind"] == "option_choice"]


# ── Detection ────────────────────────────────────────────────────────────


def test_emits_one_option_choice_slot():
    tmp = Path(tempfile.mkdtemp())
    tpl = _make_two_options_template(tmp)
    skill = dts.DocxTemplateSkill(str(tmp))
    ins, oc_slots = _inspect_oc(skill, tpl)
    assert len(oc_slots) == 1
    s = oc_slots[0]
    assert len(s["options"]) == 2
    assert s["options"][0]["index"] == 1
    assert s["options"][1]["index"] == 2
    # Public view exposes per-option preview text drawn from the body.
    assert "归甲方所有" in s["options"][0]["preview"]
    assert "由双方共有" in s["options"][1]["preview"]
    # Public view includes fill_policy hint.
    assert "第一种" in s["fill_policy"]


def test_prompt_not_double_emitted_as_removal():
    tmp = Path(tempfile.mkdtemp())
    tpl = _make_two_options_template(tmp)
    skill = dts.DocxTemplateSkill(str(tmp))
    ins, _ = _inspect_oc(skill, tpl)
    assert ins["removals"] == [], (
        "The option-choice prompt paragraph should not be emitted as a "
        "separate removal — the slot fill owns its deletion."
    )


def test_option_body_paragraphs_not_emitted_as_slots():
    tmp = Path(tempfile.mkdtemp())
    tpl = _make_two_options_template(tmp)
    skill = dts.DocxTemplateSkill(str(tmp))
    ins, oc = _inspect_oc(skill, tpl)
    # Only the option_choice slot should be emitted (no underscores, no
    # label_blank, no section_body_empty inside the option group).
    other_slots = [s for s in ins["slots"] if s["kind"] != "option_choice"]
    assert other_slots == [], f"unexpected non-OC slots: {other_slots}"


def test_no_false_positive_without_options():
    """Prompt-like text but NO `第N种` headers → no group, falls back to the
    normal removal heuristic."""
    tmp = Path(tempfile.mkdtemp())
    d = Document()
    d.add_paragraph(
        "（以下两种选择适合的一种，不选的一种请删除）"
    )
    d.add_paragraph("普通正文一段。")
    d.add_paragraph("普通正文二段。")
    tpl = tmp / "tpl.docx"
    d.save(str(tpl))
    skill = dts.DocxTemplateSkill(str(tmp))
    ins = skill.inspect_template(str(tpl))
    oc = [s for s in ins["slots"] if s["kind"] == "option_choice"]
    assert oc == []
    # The prompt should still be detected as a regular removal.
    assert any("请删除" in r["text"] or "请删除" in r["reason"]
               for r in ins["removals"])


# ── Fill ─────────────────────────────────────────────────────────────────


def test_fill_with_option_one_keeps_option_one():
    tmp = Path(tempfile.mkdtemp())
    tpl = _make_two_options_template(tmp)
    skill = dts.DocxTemplateSkill(str(tmp))
    _, oc = _inspect_oc(skill, tpl)
    out = tmp / "out.docx"
    result = skill.fill_template(str(tpl), str(out), slot_values={oc[0]["id"]: 1})
    assert result.get("success")
    body = "\n".join(p.text for p in Document(str(out)).paragraphs)
    assert "以下两种选择" not in body
    assert "第二种" not in body
    assert "由双方共有" not in body
    assert "归甲方所有" in body
    assert "乙方保证不在其他场合" in body
    # Surrounding sections preserved.
    assert "八、技术成果归属" in body
    assert "九、违约责任" in body


def test_fill_with_chinese_label_keeps_option_two():
    tmp = Path(tempfile.mkdtemp())
    tpl = _make_two_options_template(tmp)
    skill = dts.DocxTemplateSkill(str(tmp))
    _, oc = _inspect_oc(skill, tpl)
    out = tmp / "out.docx"
    skill.fill_template(str(tpl), str(out), slot_values={oc[0]["id"]: "第二种"})
    body = "\n".join(p.text for p in Document(str(out)).paragraphs)
    assert "由双方共有" in body
    assert "8.3 甲方提供的所有技术信息" in body
    assert "归甲方所有。" not in body
    assert "第一种" not in body
    assert "以下两种选择" not in body


def test_fill_with_bad_value_leaves_document_intact():
    tmp = Path(tempfile.mkdtemp())
    tpl = _make_two_options_template(tmp)
    skill = dts.DocxTemplateSkill(str(tmp))
    _, oc = _inspect_oc(skill, tpl)
    out = tmp / "out.docx"
    result = skill.fill_template(str(tpl), str(out), slot_values={oc[0]["id"]: "garbage"})
    body = "\n".join(p.text for p in Document(str(out)).paragraphs)
    # Every original paragraph remains.
    assert "以下两种选择" in body
    assert "归甲方所有" in body
    assert "由双方共有" in body
    assert "第一种" in body
    assert "第二种" in body
    # Warning surfaced.
    warns = result.get("warnings") or []
    assert any("did not match" in w for w in warns), f"warnings={warns}"


def _add_numbered_paragraph(doc, text: str, num_id: int = 1, ilvl: int = 1):
    """Append a paragraph whose only "section marker" comes from w:numPr
    (auto-numbered list) — mimicking Chinese contracts that style 不可抗力 /
    违约责任 as numbered list items rather than as Word Heading paragraphs."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    p = doc.add_paragraph(text)
    pPr = p._p.get_or_add_pPr()
    numPr = OxmlElement("w:numPr")
    ilvl_el = OxmlElement("w:ilvl")
    ilvl_el.set(qn("w:val"), str(ilvl))
    numId_el = OxmlElement("w:numId")
    numId_el.set(qn("w:val"), str(num_id))
    numPr.append(ilvl_el)
    numPr.append(numId_el)
    pPr.append(numPr)
    return p


def test_fill_does_not_absorb_auto_numbered_sections_after_option_two():
    """Regression: option 2's body must stop at an auto-numbered section
    header (w:numPr) even when the run text doesn't carry a literal
    "九、" / "十、" prefix. Otherwise the entire body of the next sections
    gets swallowed and dropped along with the rejected option."""
    tmp = Path(tempfile.mkdtemp())
    d = Document()
    d.add_paragraph("技术成果的归属和分享")
    d.add_paragraph("（以下两种选择适合的一种，可以根据需要再进行改写，不选的一种请删除）")
    d.add_paragraph("（●第一种：方案A 标题）")
    d.add_paragraph("8.1 合同履行中产生的知识产权归甲方所有。")
    d.add_paragraph("8.2 乙方保证不在其他场合复用本合同的技术成果。")
    d.add_paragraph("（●第二种：方案B 标题）")
    d.add_paragraph("8.1 合同履行过程中产生的知识产权由双方共有。")
    d.add_paragraph("8.2 本合同签订前双方各自拥有的知识产权归各自所有。")
    # Auto-numbered section headers (Word renders "九、" / "十、" from the
    # numbering definition; the run text below is just the title).
    _add_numbered_paragraph(d, "不可抗力")
    d.add_paragraph("一方因不可抗力不能履行合同的，按照影响程度免除责任。")
    _add_numbered_paragraph(d, "违约责任")
    d.add_paragraph("10.1 乙方违约责任")
    d.add_paragraph("乙方迟延交货按合同总额0.5%/周支付违约金。")
    tpl = tmp / "tpl.docx"
    d.save(str(tpl))

    skill = dts.DocxTemplateSkill(str(tmp))
    _, oc = _inspect_oc(skill, tpl)
    assert len(oc) == 1, f"expected one option_choice slot, got {len(oc)}"

    out = tmp / "out.docx"
    skill.fill_template(str(tpl), str(out), slot_values={oc[0]["id"]: 1})
    body = "\n".join(p.text for p in Document(str(out)).paragraphs)

    # Chosen option 1 body kept; option 2 body dropped.
    assert "归甲方所有" in body
    assert "由双方共有" not in body
    # Sections AFTER the option_choice block must survive — these are the
    # paragraphs the old code wrongly absorbed into option 2's body.
    assert "不可抗力" in body
    assert "一方因不可抗力" in body
    assert "违约责任" in body
    assert "10.1 乙方违约责任" in body
    assert "乙方迟延交货" in body


def test_normalize_accepts_varied_inputs():
    n = dts._normalize_option_choice
    assert n(1, 2) == 1
    assert n("2", 2) == 2
    assert n("二", 2) == 2
    assert n("第二种", 2) == 2
    assert n("second", 2) == 2
    assert n("（第一种）", 2) == 1
    assert n(3, 2) is None
    assert n("third", 2) is None
    assert n("", 2) is None
    assert n(None, 2) is None
    assert n(True, 2) is None  # bool is not a valid index


# ── Runner ───────────────────────────────────────────────────────────────


def _run_all():
    failed = []
    fns = [(name, fn) for name, fn in globals().items()
           if name.startswith("test_") and callable(fn)]
    for name, fn in fns:
        try:
            fn()
            print(f"  OK   {name}")
        except AssertionError as exc:
            print(f"  FAIL {name}: {exc}")
            failed.append(name)
        except Exception as exc:
            import traceback
            print(f"  ERR  {name}: {type(exc).__name__}: {exc}")
            traceback.print_exc()
            failed.append(name)
    print()
    if failed:
        print(f"{len(failed)} of {len(fns)} test(s) failed: {failed}")
        return 1
    print(f"All {len(fns)} tests passed.")
    return 0


if __name__ == "__main__":
    raise SystemExit(_run_all())
