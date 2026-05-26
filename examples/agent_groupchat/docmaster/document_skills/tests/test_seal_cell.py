"""Tests for seal/date placeholder cell detection and fill.

Chinese contract signature blocks use a vertically-merged tall cell with
template text like "合 同 章 \\n\\n 年  月  日". DocMaster must:
  - emit ONE fillable placeholder_cell slot per merged region (not one per
    spanned row),
  - point it at the merge anchor (the cell with the actual visible
    placeholder text), not a phantom empty cell below the merge,
  - suppress phantom empty_cell slots that sit below the merged region,
  - clear the placeholder content on fill so the stamp/date sits in the
    seal area rather than appearing at the bottom of the table.
"""
from __future__ import annotations

import sys
import tempfile
from pathlib import Path

_HERE = Path(__file__).resolve().parent
sys.path.insert(0, str(_HERE.parent))

from docx import Document  # type: ignore[import-not-found]
from docx.oxml import OxmlElement  # type: ignore[import-not-found]
from docx.oxml.ns import qn  # type: ignore[import-not-found]
import docx_template_skill as dts


def _set_vmerge(cell, val: str):
    """Set <w:vMerge w:val='restart'/> (or 'continue') on a cell."""
    tc = cell._tc
    tcPr = tc.find(qn("w:tcPr"))
    if tcPr is None:
        tcPr = OxmlElement("w:tcPr")
        tc.insert(0, tcPr)
    # Remove any existing vMerge
    for old in tcPr.findall(qn("w:vMerge")):
        tcPr.remove(old)
    vMerge = OxmlElement("w:vMerge")
    if val == "restart":
        vMerge.set(qn("w:val"), "restart")
    tcPr.append(vMerge)


def _build_seal_table_doc(tmpdir: Path) -> Path:
    """Build a minimal table mimicking the 甲方 signature block:
    - 6 rows × 2 cols
    - Col 0: field labels (单位名称, 住所, ...)
    - Col 1: rows 0-3 vertically merged into a tall "合同章 / 年 月 日" cell;
             row 4 = a real value cell; row 5 = a phantom empty cell below
             the merge."""
    d = Document()
    tbl = d.add_table(rows=6, cols=2)
    labels = ["单位名称", "住  所", "法定代表人", "甲方代表", "E-mail", "（phantom）"]
    for ri, lbl in enumerate(labels):
        tbl.rows[ri].cells[0].text = lbl
    # Put seal text into row 0 col 1 and merge rows 0-3
    seal_cell = tbl.rows[0].cells[1]
    seal_cell.text = "合 同 章\n\n\n\n年  月  日"
    _set_vmerge(seal_cell, "restart")
    for ri in range(1, 4):
        _set_vmerge(tbl.rows[ri].cells[1], "continue")
    # Row 4: a real cell with content
    tbl.rows[4].cells[1].text = "user@example.com"
    # Row 5: phantom empty cell below the merge (no vMerge)
    tbl.rows[5].cells[1].text = ""
    path = tmpdir / "tpl.docx"
    d.save(str(path))
    return path


def test_one_placeholder_slot_per_merged_seal_cell():
    """Exactly one placeholder_cell slot for the merged seal region — not
    one per spanned row."""
    tmp = Path(tempfile.mkdtemp())
    tpl = _build_seal_table_doc(tmp)
    skill = dts.DocxTemplateSkill(str(tmp))
    ins = skill.inspect_template(str(tpl))
    pc = [s for s in ins["slots"] if s["kind"] == "placeholder_cell"]
    assert len(pc) == 1, f"expected 1 placeholder_cell, got {len(pc)}: {pc}"
    assert "合" in pc[0]["label"] and "章" in pc[0]["label"]


def test_phantom_empty_cell_below_merged_seal_is_suppressed():
    """The empty cell in row 5 (below the merged seal region) must not be
    surfaced as a separate empty_cell slot. It's a layout phantom; filling
    it scatters the stamp text under the signature block."""
    tmp = Path(tempfile.mkdtemp())
    tpl = _build_seal_table_doc(tmp)
    skill = dts.DocxTemplateSkill(str(tmp))
    ins = skill.inspect_template(str(tpl))
    ec = [s for s in ins["slots"] if s["kind"] == "empty_cell"]
    # The phantom in row 5 col 1 must NOT be in the slot list.
    for s in ec:
        ctx = s.get("context", "")
        assert "phantom" not in ctx, f"phantom cell wrongly emitted: {s}"


def test_seal_fill_replaces_anchor_content():
    """Filling the placeholder_cell slot must overwrite the seal anchor's
    text (so the stamp/date appears IN the merged area, not appended after
    the original 合同章/年月日 placeholder)."""
    tmp = Path(tempfile.mkdtemp())
    tpl = _build_seal_table_doc(tmp)
    skill = dts.DocxTemplateSkill(str(tmp))
    ins = skill.inspect_template(str(tpl))
    pc = [s for s in ins["slots"] if s["kind"] == "placeholder_cell"][0]
    out = tmp / "out.docx"
    skill.fill_template(str(tpl), str(out), slot_values={
        pc["id"]: "（甲方盖章）\n2026年5月18日"
    })
    d = Document(str(out))
    seal_cell = d.tables[0].rows[0].cells[1]
    text = "\n".join(p.text for p in seal_cell.paragraphs)
    # Original placeholder gone; new content present.
    assert "合 同 章" not in text, f"placeholder leaked through: {text!r}"
    assert "（甲方盖章）" in text, f"seal value missing: {text!r}"
    assert "2026年5月18日" in text, f"date missing: {text!r}"


def test_real_seal_cell_with_only_date_placeholder():
    """A cell with just `年  月  日` (date placeholder, no seal marker) is
    also picked up — common in document headers and footers."""
    tmp = Path(tempfile.mkdtemp())
    d = Document()
    tbl = d.add_table(rows=1, cols=2)
    tbl.rows[0].cells[0].text = "签订日期"
    tbl.rows[0].cells[1].text = "      年    月    日"
    tpl = tmp / "tpl.docx"
    d.save(str(tpl))
    skill = dts.DocxTemplateSkill(str(tmp))
    ins = skill.inspect_template(str(tpl))
    pc = [s for s in ins["slots"] if s["kind"] == "placeholder_cell"]
    assert pc, f"expected date placeholder cell to be detected; got {ins['slots']}"


def test_many_highlighted_cells_all_detected():
    """Regression: a table with many distinct highlighted cells must emit a
    slot for EACH of them. Previously the per-cell dedup used `id(tc)`,
    which CPython recycles across gc'd lxml proxy objects — so most cells
    after the first were silently dropped. With the lxml-element dedup,
    every non-merged cell is scanned exactly once."""
    tmp = Path(tempfile.mkdtemp())
    d = Document()
    tbl = d.add_table(rows=8, cols=2)
    labels = [
        "单位名称", "住所", "法定代表人", "乙方代表",
        "联系人及电话", "E-mail或传真", "开户银行", "帐号",
    ]
    from docx.enum.text import WD_COLOR_INDEX
    for ri, lbl in enumerate(labels):
        tbl.rows[ri].cells[0].text = lbl
        p = tbl.rows[ri].cells[1].paragraphs[0]
        r = p.add_run("必填")
        r.font.highlight_color = WD_COLOR_INDEX.YELLOW
    tpl = tmp / "tpl.docx"
    d.save(str(tpl))

    skill = dts.DocxTemplateSkill(str(tmp))
    ins = skill.inspect_template(str(tpl))
    # The author wrote '必填' in every value cell — a meaningless filler. The
    # detector emits one highlighted slot per cell, and the meaningless-label
    # enricher rewrites each label to the row label ("单位名称", "住所", ...)
    # so the agent can tell the slots apart. The original "必填" is preserved
    # in _meta.original_label for debugging.
    highlighted_in_table = [
        s for s in ins["slots"]
        if s["kind"] == "highlighted" and s.get("original_label") == "必填"
    ]
    assert len(highlighted_in_table) == 8, (
        f"expected one slot per row (originally '必填'), got "
        f"{len(highlighted_in_table)}: {highlighted_in_table}"
    )
    enriched_labels = sorted(s["label"] for s in highlighted_in_table)
    assert enriched_labels == sorted(labels), (
        f"expected each slot to carry its row label, got {enriched_labels}"
    )


def test_non_placeholder_cell_unaffected():
    """A regular table with real content (no seal/date markers) is not
    affected by the new detection."""
    tmp = Path(tempfile.mkdtemp())
    d = Document()
    tbl = d.add_table(rows=2, cols=2)
    tbl.rows[0].cells[0].text = "项目"
    tbl.rows[0].cells[1].text = "数量"
    tbl.rows[1].cells[0].text = "螺丝"
    tbl.rows[1].cells[1].text = ""  # this is a normal empty cell — keep it
    tpl = tmp / "tpl.docx"
    d.save(str(tpl))
    skill = dts.DocxTemplateSkill(str(tmp))
    ins = skill.inspect_template(str(tpl))
    pc = [s for s in ins["slots"] if s["kind"] == "placeholder_cell"]
    ec = [s for s in ins["slots"] if s["kind"] == "empty_cell"]
    assert not pc, f"no placeholder_cell expected; got {pc}"
    assert any(s.get("label") == "数量" for s in ec), f"expected normal empty_cell; got {ec}"


def _run_all():
    failed = []
    fns = [(name, fn) for name, fn in globals().items()
           if name.startswith("test_") and callable(fn)]
    for name, fn in fns:
        try:
            fn()
            print(f"  OK   {name}")
        except AssertionError as exc:
            print(f"  FAIL {name}: {exc}")
            failed.append(name)
        except Exception as exc:
            import traceback
            print(f"  ERR  {name}: {type(exc).__name__}: {exc}")
            traceback.print_exc()
            failed.append(name)
    print()
    if failed:
        print(f"{len(failed)} of {len(fns)} test(s) failed: {failed}")
        return 1
    print(f"All {len(fns)} tests passed.")
    return 0


if __name__ == "__main__":
    raise SystemExit(_run_all())
