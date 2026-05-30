"""Tests for detection enhancements added for CN ministry templates:

- Inline whitespace blanks bounded by label colons / currency symbols.
- Auto-numbered (w:numPr) section headings as section_body_empty triggers.
- Explicit removal markers like "应删除" / "此句话非正文".
- Suppression: a paragraph flagged as a full-paragraph removal should not
  also surface as a fillable highlighted slot.

Plain-assert style — runs with `python test_inline_blanks_and_removals.py`.
"""
from __future__ import annotations

import sys
import tempfile
from pathlib import Path

_HERE = Path(__file__).resolve().parent
sys.path.insert(0, str(_HERE.parent))

from docx import Document  # type: ignore[import-not-found]
from docx.oxml import OxmlElement  # type: ignore[import-not-found]
from docx.oxml.ns import qn  # type: ignore[import-not-found]
from docx.enum.text import WD_COLOR_INDEX  # type: ignore[import-not-found]
import docx_template_skill as dts


def _add_numbered_paragraph(doc, text: str, num_id: int = 1, ilvl: int = 1):
    p = doc.add_paragraph(text)
    pPr = p._p.get_or_add_pPr()
    numPr = OxmlElement("w:numPr")
    ilvl_el = OxmlElement("w:ilvl")
    ilvl_el.set(qn("w:val"), str(ilvl))
    numId_el = OxmlElement("w:numId")
    numId_el.set(qn("w:val"), str(num_id))
    numPr.append(ilvl_el)
    numPr.append(numId_el)
    pPr.append(numPr)
    return p


# ── Inline whitespace blanks ─────────────────────────────────────────────


def test_inline_blank_between_colon_and_punctuation():
    """`交货地点：     ；rest` — whitespace gap bounded by '：' and '；'
    must be detected as a fillable slot."""
    tmp = Path(tempfile.mkdtemp())
    d = Document()
    d.add_paragraph("交货地点：         ；乙方负责运输。")
    tpl = tmp / "tpl.docx"
    d.save(str(tpl))
    skill = dts.DocxTemplateSkill(str(tmp))
    ins = skill.inspect_template(str(tpl))
    inline = [s for s in ins["slots"]
              if s["kind"] == "underscores" and "交货地点" in (s.get("label") or "")]
    assert inline, f"expected an inline blank slot for 交货地点; got {ins['slots']}"
    out = tmp / "out.docx"
    skill.fill_template(str(tpl), str(out),
                        slot_values={inline[0]["id"]: "北京市石景山区"})
    body = "\n".join(p.text for p in Document(str(out)).paragraphs)
    assert "交货地点：北京市石景山区；乙方负责运输" in body, body


def test_inline_blank_money_paragraph_two_gaps():
    """The 总经费 line has TWO whitespace gaps (`：    元` and `￥    元`)
    plus a highlighted drafting note at the tail. All three must be picked
    up — previously only the highlighted tail was emitted, and the fill
    leaked into the wrong position."""
    tmp = Path(tempfile.mkdtemp())
    d = Document()
    p = d.add_paragraph()
    p.add_run("（一）本项目总经费：")
    # First whitespace gap with underline (Word's "fill-in line" trick).
    r_gap1 = p.add_run("            ")
    r_gap1.font.underline = True
    p.add_run("元（元），即￥")
    r_gap2 = p.add_run("        ")
    r_gap2.font.underline = True
    p.add_run("元（小写）。报价明细表见本合同附件。")
    r_hl = p.add_run("（非标制造产品，通常应该提供报价明细）")
    r_hl.font.highlight_color = WD_COLOR_INDEX.YELLOW
    tpl = tmp / "tpl.docx"
    d.save(str(tpl))
    skill = dts.DocxTemplateSkill(str(tmp))
    ins = skill.inspect_template(str(tpl))
    # Find the two underscore-kind blanks pointing at this paragraph.
    blanks = [s for s in ins["slots"] if s["kind"] == "underscores"]
    assert len(blanks) >= 2, f"expected ≥2 inline blanks, got {len(blanks)}: {blanks}"
    # The highlighted span must still be emitted as a separate slot.
    hls = [s for s in ins["slots"] if s["kind"] == "highlighted"]
    assert hls, "expected highlighted drafting note to still surface as a slot"


def test_uw_span_absorbs_slash_draft_mark():
    """`label：   /   unit；` — the '/' between two underlined gaps is a
    draft mark, NOT alternatives. The fill must replace the whole region
    (leading gap + slash + trailing gap) cleanly, leaving '平方米' intact."""
    tmp = Path(tempfile.mkdtemp())
    d = Document()
    p = d.add_paragraph()
    p.add_run("建筑面积：")
    r1 = p.add_run("       ")
    r1.font.underline = True
    p.add_run("/")
    r2 = p.add_run("        ")
    r2.font.underline = True
    p.add_run("平方米；")
    tpl = tmp / "tpl.docx"
    d.save(str(tpl))
    skill = dts.DocxTemplateSkill(str(tmp))
    ins = skill.inspect_template(str(tpl))
    blanks = [s for s in ins["slots"]
              if s["kind"] == "underscores" and "建筑面积" in (s.get("label") or "")]
    assert blanks, f"expected a slot for 建筑面积; got {ins['slots']}"
    out = tmp / "out.docx"
    skill.fill_template(str(tpl), str(out), slot_values={blanks[0]["id"]: "8500"})
    body = "\n".join(p.text for p in Document(str(out)).paragraphs)
    assert "/" not in body, f"slash draft mark should be replaced: {body!r}"
    assert "8500" in body and "平方米" in body, body


def test_uw_span_absorbs_sample_text():
    """`label：  <sample>  ` — when an underlined gap follows sample text
    (no hard boundary in between), the slot must cover both so the fill
    REPLACES the sample instead of appending alongside it."""
    tmp = Path(tempfile.mkdtemp())
    d = Document()
    p = d.add_paragraph()
    p.add_run("承包方式：  包工包料")
    r = p.add_run("                                     ")
    r.font.underline = True
    tpl = tmp / "tpl.docx"
    d.save(str(tpl))
    skill = dts.DocxTemplateSkill(str(tmp))
    ins = skill.inspect_template(str(tpl))
    blanks = [s for s in ins["slots"]
              if s["kind"] == "underscores" and "承包方式" in (s.get("label") or "")]
    assert blanks, f"expected a slot for 承包方式; got {ins['slots']}"
    out = tmp / "out.docx"
    skill.fill_template(str(tpl), str(out),
                        slot_values={blanks[0]["id"]: "总包"})
    body = "\n".join(p.text for p in Document(str(out)).paragraphs)
    # Sample '包工包料' must be gone; new value '总包' present exactly once.
    assert "包工包料" not in body, f"sample text must be replaced: {body!r}"
    assert body.count("总包") == 1, body


def test_uw_spans_merge_across_sample_list():
    """`label：  <gap1> <sample,with,顿号> <gap2>` — two underlined gaps
    separated by sample text containing only 顿号 ('、') must merge into one
    slot, since 顿号 is intra-list punctuation, not a slot boundary."""
    tmp = Path(tempfile.mkdtemp())
    d = Document()
    p = d.add_paragraph()
    p.add_run("承包范围：")
    r1 = p.add_run("    ")
    r1.font.underline = True
    p.add_run("包括防水、门窗更换、粉刷等")
    r2 = p.add_run("                ")
    r2.font.underline = True
    tpl = tmp / "tpl.docx"
    d.save(str(tpl))
    skill = dts.DocxTemplateSkill(str(tmp))
    ins = skill.inspect_template(str(tpl))
    blanks = [s for s in ins["slots"]
              if s["kind"] == "underscores" and "承包范围" in (s.get("label") or "")]
    # Should be exactly ONE merged slot, not two.
    assert len(blanks) == 1, f"expected 1 merged slot, got {len(blanks)}: {blanks}"
    out = tmp / "out.docx"
    skill.fill_template(str(tpl), str(out),
                        slot_values={blanks[0]["id"]: "外墙防水、门窗更换"})
    body = "\n".join(p.text for p in Document(str(out)).paragraphs)
    assert "包括防水" not in body and "粉刷等" not in body, body
    assert "外墙防水、门窗更换" in body, body


def test_underscore_pair_with_slash_draft_mark_absorbs_separator():
    """`____/____` between a label and a unit is a single fillable region
    where `/` is a draft mark, not a separator preserving formatting. Past
    behavior emitted a 'composite' slot covering both underscore runs but
    left the literal `/` between them — output was `12000/平方米` instead of
    `12000 平方米`."""
    tmp = Path(tempfile.mkdtemp())
    d = Document()
    p = d.add_paragraph()
    p.add_run("建筑面积：____/____平方米")
    tpl = tmp / "tpl.docx"
    d.save(str(tpl))
    skill = dts.DocxTemplateSkill(str(tmp))
    ins = skill.inspect_template(str(tpl))
    slot = next(s for s in ins["slots"] if s["kind"] == "underscores")
    out = tmp / "out.docx"
    skill.fill_template(
        str(tpl), str(out), slot_values={slot["id"]: "12000"}
    )
    body = "\n".join(p.text for p in Document(str(out)).paragraphs)
    assert "/" not in body, f"draft slash should be absorbed: {body!r}"
    assert "12000" in body and "平方米" in body, body


def test_underscore_pair_with_backslash_draft_mark_absorbs_separator():
    """Same as the slash case, but with backslash. Both are equivalently
    used as draft marks in CN templates."""
    tmp = Path(tempfile.mkdtemp())
    d = Document()
    p = d.add_paragraph()
    p.add_run("结构类型：____\\____ ；檐高：____ 米")
    tpl = tmp / "tpl.docx"
    d.save(str(tpl))
    skill = dts.DocxTemplateSkill(str(tmp))
    ins = skill.inspect_template(str(tpl))
    structure_slot = next(
        s for s in ins["slots"]
        if s["kind"] == "underscores" and "结构类型" in (s.get("label") or "")
    )
    out = tmp / "out.docx"
    skill.fill_template(
        str(tpl), str(out), slot_values={structure_slot["id"]: "框架剪力墙"}
    )
    body = "\n".join(p.text for p in Document(str(out)).paragraphs)
    assert "\\" not in body, f"draft backslash should be absorbed: {body!r}"
    assert "框架剪力墙" in body, body


def test_underscore_digit_cell_composite_preserved():
    """`¥ ____ ____ ____` with PURE whitespace between cells stays a
    digit-cell composite — first cell takes the value, others blank. This
    must keep working after the slash-absorption fix."""
    tmp = Path(tempfile.mkdtemp())
    d = Document()
    p = d.add_paragraph()
    p.add_run("总价：¥ ____ ____ ____ 元")
    tpl = tmp / "tpl.docx"
    d.save(str(tpl))
    skill = dts.DocxTemplateSkill(str(tmp))
    ins = skill.inspect_template(str(tpl))
    slot = next(s for s in ins["slots"] if s["kind"] == "underscores")
    out = tmp / "out.docx"
    skill.fill_template(
        str(tpl), str(out), slot_values={slot["id"]: "8500"}
    )
    body = "\n".join(p.text for p in Document(str(out)).paragraphs)
    # Value placed once; '元' and '¥' still in place.
    assert body.count("8500") == 1, body
    assert "¥" in body and "元" in body, body


def test_inline_blank_with_slash_draft_mark_detected():
    """`label：    /    ；` — plain (no-underline) whitespace gap broken by
    a single slash draft mark. The inline-blank scanner missed this because
    its regex required contiguous whitespace, so the slot was never emitted
    and the draft mark survived the fill."""
    tmp = Path(tempfile.mkdtemp())
    d = Document()
    d.add_paragraph("建筑面积：    /    平方米；")
    tpl = tmp / "tpl.docx"
    d.save(str(tpl))
    skill = dts.DocxTemplateSkill(str(tmp))
    ins = skill.inspect_template(str(tpl))
    blanks = [
        s for s in ins["slots"]
        if s["kind"] == "underscores" and "建筑面积" in (s.get("label") or "")
    ]
    assert blanks, f"expected an inline-blank slot for 建筑面积; got {ins['slots']}"
    out = tmp / "out.docx"
    skill.fill_template(
        str(tpl), str(out), slot_values={blanks[0]["id"]: "12000"}
    )
    body = "\n".join(p.text for p in Document(str(out)).paragraphs)
    assert "/" not in body, f"draft slash should be absorbed: {body!r}"
    assert "12000" in body and "平方米" in body, body


def test_inline_blank_with_backslash_draft_mark_detected():
    """Same case but with backslash."""
    tmp = Path(tempfile.mkdtemp())
    d = Document()
    d.add_paragraph("承包范围：    \\    ；")
    tpl = tmp / "tpl.docx"
    d.save(str(tpl))
    skill = dts.DocxTemplateSkill(str(tmp))
    ins = skill.inspect_template(str(tpl))
    blanks = [
        s for s in ins["slots"]
        if s["kind"] == "underscores" and "承包范围" in (s.get("label") or "")
    ]
    assert blanks, f"expected an inline-blank slot for 承包范围; got {ins['slots']}"
    out = tmp / "out.docx"
    skill.fill_template(
        str(tpl), str(out), slot_values={blanks[0]["id"]: "外墙翻新"}
    )
    body = "\n".join(p.text for p in Document(str(out)).paragraphs)
    assert "\\" not in body, f"draft backslash should be absorbed: {body!r}"
    assert "外墙翻新" in body, body


def test_inline_blank_normal_slash_in_label_preserved():
    """Counter-test: a slash that's part of a real label like '檐高/跨度'
    must NOT be treated as a draft mark — it's structural punctuation
    inside the label, not a fillable region. The slot's RIGHT side
    (after the colon) is what gets filled."""
    tmp = Path(tempfile.mkdtemp())
    d = Document()
    d.add_paragraph("檐高/跨度：    平方米；")
    tpl = tmp / "tpl.docx"
    d.save(str(tpl))
    skill = dts.DocxTemplateSkill(str(tmp))
    ins = skill.inspect_template(str(tpl))
    blanks = [s for s in ins["slots"] if s["kind"] == "underscores"]
    assert blanks, f"expected a slot; got {ins['slots']}"
    out = tmp / "out.docx"
    skill.fill_template(
        str(tpl), str(out), slot_values={blanks[0]["id"]: "48"}
    )
    body = "\n".join(p.text for p in Document(str(out)).paragraphs)
    # The '檐高/跨度' label slash must remain intact.
    assert "檐高/跨度" in body, body
    assert "48" in body, body


def test_force_fresh_skips_continuation_autodetect():
    """When `force_fresh=True`, fill_template must overwrite an existing
    output_path with a fresh fill from the original template, even if the
    output file already exists. Past sessions had the agent loop because
    the implicit continuation auto-detect kept swapping the source to a
    botched partial fill — `force_fresh` is the documented escape hatch."""
    tmp = Path(tempfile.mkdtemp())
    d = Document()
    p = d.add_paragraph()
    p.add_run("甲方：")
    r = p.add_run("                ")
    r.font.underline = True
    p2 = d.add_paragraph()
    p2.add_run("乙方：")
    r2 = p2.add_run("                ")
    r2.font.underline = True
    tpl = tmp / "tpl.docx"
    d.save(str(tpl))
    skill = dts.DocxTemplateSkill(str(tmp))
    ins = skill.inspect_template(str(tpl))
    slots_by_label = {s.get("label"): s["id"] for s in ins["slots"]}
    out = tmp / "out.docx"
    # First fill: WRONG values (simulate a botched run).
    skill.fill_template(
        str(tpl), str(out),
        slot_values={
            slots_by_label["甲方"]: "WRONG_PARTY_A",
            slots_by_label["乙方"]: "WRONG_PARTY_B",
        },
    )
    body = "\n".join(p.text for p in Document(str(out)).paragraphs)
    assert "WRONG_PARTY_A" in body, body
    # Second fill WITHOUT force_fresh: should trigger continuation mode and
    # surface the notice. Slots were already consumed by the partial doc, so
    # canonical ids no longer resolve — the agent's intent (re-fill from the
    # original) is silently lost.
    res_continuation = skill.fill_template(
        str(tpl), str(out),
        slot_values={
            slots_by_label["甲方"]: "Acme",
            slots_by_label["乙方"]: "Beta",
        },
    )
    assert res_continuation.get("chunked_continuation") is True, res_continuation
    assert res_continuation.get("continuation_notice"), res_continuation
    # Third fill WITH force_fresh: overwrite from the original template.
    res_fresh = skill.fill_template(
        str(tpl), str(out),
        slot_values={
            slots_by_label["甲方"]: "Acme",
            slots_by_label["乙方"]: "Beta",
        },
        force_fresh=True,
    )
    assert res_fresh.get("chunked_continuation") is False, res_fresh
    assert "continuation_notice" not in res_fresh, res_fresh
    body = "\n".join(p.text for p in Document(str(out)).paragraphs)
    assert "WRONG_PARTY_A" not in body and "WRONG_PARTY_B" not in body, body
    assert "Acme" in body and "Beta" in body, body


def test_legacy_slot_id_rejected_at_fill_boundary():
    """When inspect_template returned descriptive ids like
    `slot_000_<label>`, a caller passing the bare legacy form `slot_0` must
    be rejected before any document mutation — silently routing on the
    numeric prefix corrupted real fills in past incidents."""
    tmp = Path(tempfile.mkdtemp())
    d = Document()
    p = d.add_paragraph()
    p.add_run("甲方：")
    r = p.add_run("                ")
    r.font.underline = True
    tpl = tmp / "tpl.docx"
    d.save(str(tpl))
    skill = dts.DocxTemplateSkill(str(tmp))
    ins = skill.inspect_template(str(tpl))
    slot = next(s for s in ins["slots"] if s["kind"] == "underscores")
    canonical = slot["id"]
    assert canonical != "slot_0", (
        f"expected descriptive canonical id, got {canonical!r}"
    )
    out = tmp / "out.docx"
    # Caller uses the legacy bare form — must be refused.
    result = skill.fill_template(
        str(tpl), str(out), slot_values={"slot_0": "Acme Corp"}
    )
    sf = result.get("slot_fill") or result
    assert sf.get("rejected") is True, sf
    assert "slot_0" in (sf.get("rejected_legacy_slot_ids") or []), sf
    assert sf.get("legacy_canonical_map", {}).get("slot_0") == canonical, sf
    # No mutation should reach the output document — fill_template signals
    # failure when slot-fill is the only requested operation.
    assert result.get("success") is False, result


def test_inline_blank_not_triggered_in_normal_prose():
    """Plain prose with a colon and only a single space afterward must NOT
    be detected as a blank — the regex requires ≥3 whitespace chars."""
    tmp = Path(tempfile.mkdtemp())
    d = Document()
    d.add_paragraph("注意：本合同自双方签字盖章后生效。")
    tpl = tmp / "tpl.docx"
    d.save(str(tpl))
    skill = dts.DocxTemplateSkill(str(tmp))
    ins = skill.inspect_template(str(tpl))
    assert not any(s["kind"] == "underscores" for s in ins["slots"]), ins["slots"]


# ── Date scaffold (`__年__月__日` with whitespace gaps) ────────────────


def test_date_scaffold_year_month_day():
    """`签订日期：    年    月    日` — the whole scaffold (gaps + 年/月/日)
    becomes one slot whose fill replaces the entire span cleanly. Without
    this, only the leading gap is detected and the value gets sandwiched
    against leftover `年   月   日` template text."""
    tmp = Path(tempfile.mkdtemp())
    d = Document()
    d.add_paragraph("签订日期：       年     月     日")
    tpl = tmp / "tpl.docx"
    d.save(str(tpl))
    skill = dts.DocxTemplateSkill(str(tmp))
    ins = skill.inspect_template(str(tpl))
    date_slots = [s for s in ins["slots"]
                  if s["kind"] == "underscores" and "签订日期" in (s.get("label") or s.get("context") or "")]
    assert len(date_slots) == 1, f"expected single date slot; got {date_slots}"
    out = tmp / "out.docx"
    skill.fill_template(str(tpl), str(out), slot_values={date_slots[0]["id"]: "2026年5月18日"})
    body = "\n".join(p.text for p in Document(str(out)).paragraphs)
    # Clean replacement — no leftover template scaffolding.
    assert "签订日期：2026年5月18日" in body, body
    # No stray `年   月` text outside the value.
    assert "日年" not in body, f"leftover scaffold after fill: {body!r}"
    assert body.count("年") == 1
    assert body.count("月") == 1
    # `签订日期` itself contains a `日`, so the total count of 日 in
    # `签订日期：2026年5月18日` is two — one in the label, one in the value.
    assert body.count("日") == 2


def test_date_scaffold_year_month_only():
    """Two-segment scaffold `    年    月` (no 日) also becomes one slot."""
    tmp = Path(tempfile.mkdtemp())
    d = Document()
    d.add_paragraph("合同期限：       年     月")
    tpl = tmp / "tpl.docx"
    d.save(str(tpl))
    skill = dts.DocxTemplateSkill(str(tmp))
    ins = skill.inspect_template(str(tpl))
    date_slots = [s for s in ins["slots"]
                  if s["kind"] == "underscores" and "合同期限" in (s.get("label") or s.get("context") or "")]
    assert len(date_slots) == 1, f"expected single slot; got {date_slots}"
    out = tmp / "out.docx"
    skill.fill_template(str(tpl), str(out), slot_values={date_slots[0]["id"]: "2026年5月"})
    body = "\n".join(p.text for p in Document(str(out)).paragraphs)
    assert "合同期限：2026年5月" in body, body
    assert "月年" not in body, body


def test_date_scaffold_no_false_positive_for_normal_prose():
    """`2024年5月` written as normal prose (no whitespace gap before 年)
    must NOT be detected as a fillable scaffold."""
    tmp = Path(tempfile.mkdtemp())
    d = Document()
    d.add_paragraph("本合同自2024年5月起生效。")
    tpl = tmp / "tpl.docx"
    d.save(str(tpl))
    skill = dts.DocxTemplateSkill(str(tmp))
    ins = skill.inspect_template(str(tpl))
    assert not any(s["kind"] == "underscores" for s in ins["slots"]), ins["slots"]


# ── section_body_empty for auto-numbered headings ────────────────────────


def test_section_body_empty_under_auto_numbered_heading():
    """An empty paragraph immediately after a numbered (w:numPr) heading
    must surface as a section_body_empty slot. Previously these were
    missed because `_is_heading` only checked Heading-style names."""
    tmp = Path(tempfile.mkdtemp())
    d = Document()
    _add_numbered_paragraph(d, "技术指标及要求")
    d.add_paragraph("")  # empty body — should become a fillable slot
    _add_numbered_paragraph(d, "双方协作内容")
    d.add_paragraph("")
    tpl = tmp / "tpl.docx"
    d.save(str(tpl))
    skill = dts.DocxTemplateSkill(str(tmp))
    ins = skill.inspect_template(str(tpl))
    labels = [s.get("label") for s in ins["slots"] if s["kind"] == "section_body_empty"]
    assert "技术指标及要求" in labels, labels
    assert "双方协作内容" in labels, labels


def test_auto_numbered_body_text_not_a_heading():
    """A LONG numbered paragraph (body sub-item, not a section title) must
    NOT be treated as a heading — otherwise the empty paragraph after it
    would wrongly emit a section_body_empty slot."""
    tmp = Path(tempfile.mkdtemp())
    d = Document()
    _add_numbered_paragraph(
        d,
        "本合同技术开发风险责任由乙方承担，乙方需按时交付符合标准的产品。",
        ilvl=4,
    )
    d.add_paragraph("")
    tpl = tmp / "tpl.docx"
    d.save(str(tpl))
    skill = dts.DocxTemplateSkill(str(tmp))
    ins = skill.inspect_template(str(tpl))
    assert not any(s["kind"] == "section_body_empty" for s in ins["slots"]), ins["slots"]


# ── Removal markers ──────────────────────────────────────────────────────


def test_removal_marker_yingshanchu():
    """`应删除` / `此句话非正文` in a fully-parenthesized paragraph must
    be detected as a whole-paragraph removal, not a fillable highlighted
    slot."""
    tmp = Path(tempfile.mkdtemp())
    d = Document()
    p = d.add_paragraph()
    r = p.add_run("（全文排版确保下表不跨页。此句话非正文，应删除）")
    r.font.highlight_color = WD_COLOR_INDEX.YELLOW
    tpl = tmp / "tpl.docx"
    d.save(str(tpl))
    skill = dts.DocxTemplateSkill(str(tmp))
    ins = skill.inspect_template(str(tpl))
    # Removal emitted as full-paragraph.
    rms = [r for r in ins["removals"] if r["kind"] == "instruction_paragraph"]
    assert rms, f"expected an instruction_paragraph removal; got {ins['removals']}"
    # No duplicate highlighted slot for the same paragraph.
    hls = [s for s in ins["slots"]
           if s["kind"] == "highlighted" and "应删除" in (s.get("context") or "")]
    assert not hls, f"highlighted slot should be suppressed; got {hls}"


def test_run_level_removal_unsuppressed():
    """A run-level removal (instruction_run) inside otherwise-normal prose
    must NOT suppress fillable slots elsewhere in that paragraph."""
    tmp = Path(tempfile.mkdtemp())
    d = Document()
    p = d.add_paragraph()
    r_hl = p.add_run("HT-PLACEHOLDER")
    r_hl.font.highlight_color = WD_COLOR_INDEX.YELLOW
    p.add_run(" 合同编号请填写完整 ")
    # Add a non-italic run-level removal phrase that won't dominate the paragraph.
    p.add_run("应删除")
    tpl = tmp / "tpl.docx"
    d.save(str(tpl))
    skill = dts.DocxTemplateSkill(str(tmp))
    ins = skill.inspect_template(str(tpl))
    # The highlighted slot should still surface.
    hls = [s for s in ins["slots"] if s["kind"] == "highlighted"]
    assert hls, f"expected highlighted slot to remain; got {ins['slots']}"


# ── Container section heading suppression ───────────────────────────────


def test_container_section_heading_no_label_blank():
    """A top-level numbered heading like '四、交货方式及时间、地点：' whose
    body lives in the following paragraphs must NOT be emitted as a
    label_blank — otherwise a value gets appended after the heading's colon
    in addition to (or instead of) filling the structured sub-fields below."""
    tmp = Path(tempfile.mkdtemp())
    d = Document()
    d.add_paragraph("四、交货方式及时间、地点：")
    d.add_paragraph("交货方式：由乙方负责送货或办理货物托运，运输及保险费由乙方负担；")
    d.add_paragraph("交货时间：合同签订之日起      周内；")
    d.add_paragraph("交货地点：甲方单位（北京市石景山区玉泉路19号乙院）。")
    tpl = tmp / "tpl.docx"
    d.save(str(tpl))
    skill = dts.DocxTemplateSkill(str(tmp))
    ins = skill.inspect_template(str(tpl))
    bad = [s for s in ins["slots"]
           if s["kind"] == "label_blank"
           and "交货方式及时间" in (s.get("label") or "")]
    assert not bad, f"container heading should not produce label_blank; got {bad}"
    # The sub-field heading "交货方式：" is also a Label: paragraph, but it
    # legitimately has body content right after the colon ("由乙方负责..."),
    # so it should be left alone (no label_blank emitted on a filled line).
    filled_sub = [s for s in ins["slots"]
                  if s["kind"] == "label_blank"
                  and "交货方式" == (s.get("label") or "").strip()]
    assert not filled_sub, f"already-filled sub-field should not get a slot; got {filled_sub}"


def test_container_section_heading_auto_numbered():
    """Same as above but the leading '四、' comes from Word's auto-numbering
    (w:numPr), so the run text is just '交货方式及时间、地点：' with no
    visible section marker. Real CN ministry templates render the section
    number from a numbering definition — the suppression must still fire."""
    tmp = Path(tempfile.mkdtemp())
    d = Document()
    _add_numbered_paragraph(d, "交货方式及时间、地点：", ilvl=1)
    d.add_paragraph("交货方式：由乙方负责送货或办理货物托运，运输及保险费由乙方负担；")
    d.add_paragraph("交货时间：合同签订之日起      周内；")
    d.add_paragraph("交货地点：甲方单位（北京市石景山区玉泉路19号乙院）。")
    tpl = tmp / "tpl.docx"
    d.save(str(tpl))
    skill = dts.DocxTemplateSkill(str(tmp))
    ins = skill.inspect_template(str(tpl))
    bad = [s for s in ins["slots"]
           if s["kind"] == "label_blank"
           and "交货方式及时间" in (s.get("label") or "")]
    assert not bad, f"auto-numbered container heading must not produce label_blank; got {bad}"


def test_container_section_heading_with_prose_body():
    """Auto-numbered section heading like '产品的保修期及售后服务：' followed
    by a plain prose body (no '：' label on the next paragraph). This must
    still suppress the heading's label_blank — the value should go in the
    prose body, not be appended after the heading colon."""
    tmp = Path(tempfile.mkdtemp())
    d = Document()
    _add_numbered_paragraph(d, "产品的保修期及售后服务：", ilvl=0)
    d.add_paragraph("本合同产品从甲方验收合格之日起，保修期或质量保证期    年。"
                    "保修期内非因甲方人为原因产品如发生故障乙方应负责免费提供维修。")
    tpl = tmp / "tpl.docx"
    d.save(str(tpl))
    skill = dts.DocxTemplateSkill(str(tmp))
    ins = skill.inspect_template(str(tpl))
    bad = [s for s in ins["slots"]
           if s["kind"] == "label_blank" and "保修期" in (s.get("label") or "")]
    assert not bad, f"container heading with prose body must not get label_blank; got {bad}"


def test_container_section_heading_with_deeper_numbered_body():
    """A heading at ilvl=0 followed by body items at ilvl=1 (same numId but
    deeper level) — the heading is a parent, the items are its children.
    Real example: '九、违约责任：' (ilvl=0) over '甲方无正当理由...' (ilvl=1)."""
    tmp = Path(tempfile.mkdtemp())
    d = Document()
    _add_numbered_paragraph(d, "违约责任：", ilvl=0)
    _add_numbered_paragraph(d,
        "甲方无正当理由不按照本合同约定的期限和金额支付价款，应向乙方支付迟延付款违约金。",
        ilvl=1)
    _add_numbered_paragraph(d,
        "除因不可抗力，乙方若不按合同规定的时间向甲方交货，甲方有权解除合同。",
        ilvl=1)
    tpl = tmp / "tpl.docx"
    d.save(str(tpl))
    skill = dts.DocxTemplateSkill(str(tmp))
    ins = skill.inspect_template(str(tpl))
    bad = [s for s in ins["slots"]
           if s["kind"] == "label_blank" and "违约责任" in (s.get("label") or "")]
    assert not bad, f"heading over deeper-ilvl body must not get label_blank; got {bad}"


def test_numbered_siblings_label_blank_retained():
    """Two auto-numbered sibling paragraphs `项目名称：` / `项目编号：xxx`
    must NOT trigger container suppression — they're peers, not parent/
    child. The first one is still a fillable label_blank."""
    tmp = Path(tempfile.mkdtemp())
    d = Document()
    _add_numbered_paragraph(d, "项目名称：", ilvl=1)
    _add_numbered_paragraph(d, "项目编号：XYZ-001", ilvl=1)
    tpl = tmp / "tpl.docx"
    d.save(str(tpl))
    skill = dts.DocxTemplateSkill(str(tmp))
    ins = skill.inspect_template(str(tpl))
    lbs = [s for s in ins["slots"]
           if s["kind"] == "label_blank" and "项目名称" in (s.get("label") or "")]
    assert lbs, f"numbered sibling label_blank must be retained; got {ins['slots']}"


def test_label_only_heading_retained_when_followed_by_section_break():
    """A `Label:` paragraph followed by another top-level section break
    (no body content under it) must still emit a label_blank — the value
    legitimately goes after the colon."""
    tmp = Path(tempfile.mkdtemp())
    d = Document()
    d.add_paragraph("四、项目名称：")
    d.add_paragraph("五、技术要求：详见附件。")
    tpl = tmp / "tpl.docx"
    d.save(str(tpl))
    skill = dts.DocxTemplateSkill(str(tmp))
    ins = skill.inspect_template(str(tpl))
    lbs = [s for s in ins["slots"]
           if s["kind"] == "label_blank" and "项目名称" in (s.get("label") or "")]
    assert lbs, f"expected label_blank for 项目名称; got {ins['slots']}"


def test_underscore_slot_absorbs_trailing_slash_plus_plain_gap():
    """`____/         平方米` — single underscore run followed by a slash and
    plenty of whitespace before the unit. The slot must extend rightward over
    the slash + whitespace so the fill replaces them all (otherwise output is
    `12000/         平方米` with the draft mark surviving)."""
    tmp = Path(tempfile.mkdtemp())
    d = Document()
    p = d.add_paragraph()
    p.add_run("建筑面积：________/         平方米；")
    tpl = tmp / "tpl.docx"
    d.save(str(tpl))
    skill = dts.DocxTemplateSkill(str(tmp))
    ins = skill.inspect_template(str(tpl))
    slot = next(
        s for s in ins["slots"]
        if s["kind"] == "underscores" and "建筑面积" in (s.get("label") or "")
    )
    out = tmp / "out.docx"
    skill.fill_template(
        str(tpl), str(out), slot_values={slot["id"]: "12000"}
    )
    body = "\n".join(p.text for p in Document(str(out)).paragraphs)
    assert "/" not in body, f"draft slash should be absorbed: {body!r}"
    assert "12000" in body and "平方米" in body, body


def test_underscore_slot_absorbs_trailing_slash_at_end_of_line():
    """`京发改〔2025〕第0128号/` — value already in place, but author left a
    trailing `/` draft mark with nothing after it. The fillable region (the
    underscore slot) should extend to consume that trailing slash."""
    tmp = Path(tempfile.mkdtemp())
    d = Document()
    p = d.add_paragraph()
    p.add_run("批准文号：________/")
    tpl = tmp / "tpl.docx"
    d.save(str(tpl))
    skill = dts.DocxTemplateSkill(str(tmp))
    ins = skill.inspect_template(str(tpl))
    slot = next(
        s for s in ins["slots"]
        if s["kind"] == "underscores" and "批准文号" in (s.get("label") or "")
    )
    out = tmp / "out.docx"
    skill.fill_template(
        str(tpl), str(out),
        slot_values={slot["id"]: "京发改〔2025〕第0128号"},
    )
    body = "\n".join(p.text for p in Document(str(out)).paragraphs)
    assert not body.rstrip().endswith("/"), (
        f"trailing draft slash should be absorbed: {body!r}"
    )
    assert "京发改〔2025〕第0128号" in body, body


def test_underscore_decorator_group_absorbs_trailing_slash_plus_gap():
    """Decorator-gap group `____/____` followed by `/         平方米` — both
    the inter-pair slash AND the trailing slash+whitespace should be absorbed
    into one slot (the prior fix handled the inter-pair slash; this test
    nails the trailing extension)."""
    tmp = Path(tempfile.mkdtemp())
    d = Document()
    p = d.add_paragraph()
    p.add_run("结构类型：____\\____         ；檐高：____ 米")
    tpl = tmp / "tpl.docx"
    d.save(str(tpl))
    skill = dts.DocxTemplateSkill(str(tmp))
    ins = skill.inspect_template(str(tpl))
    structure_slot = next(
        s for s in ins["slots"]
        if s["kind"] == "underscores" and "结构类型" in (s.get("label") or "")
    )
    out = tmp / "out.docx"
    skill.fill_template(
        str(tpl), str(out),
        slot_values={structure_slot["id"]: "框架剪力墙"},
    )
    body = "\n".join(p.text for p in Document(str(out)).paragraphs)
    # The 结构类型 fill must absorb the inter-pair backslash AND the trailing
    # whitespace gap before `；` (no draft mark survives anywhere in the
    # 结构类型 region).
    assert "\\" not in body, f"draft backslash should be absorbed: {body!r}"
    # No `/` in the structure-type half either (檐高's underscores haven't been
    # filled, but they don't have a slash before them).
    structure_segment = body.split("檐高")[0]
    assert "/" not in structure_segment, (
        f"draft slashes should be absorbed in 结构类型 region: {body!r}"
    )
    assert "框架剪力墙" in body, body


def test_strip_stranded_draft_marks_handles_baked_in_value_with_unit():
    """Template with value already baked in: `建筑面积：12000/        平方米`.
    No underscores remain, so the scanner can't attach a slot — but the
    stranded `/        ` between `12000` and `平方米` must still be cleaned
    up during fill. Without the pre-save cleanup pass, the output keeps the
    draft mark and the user sees `12000/        平方米` in the result."""
    tmp = Path(tempfile.mkdtemp())
    d = Document()
    d.add_paragraph("建筑面积：12000/        平方米；层数：6/              ")
    # Add a fillable slot somewhere else so fill_template doesn't bail out
    # on "nothing to fill" (the cleanup runs on the final document regardless).
    d.add_paragraph("项目名称：________")
    tpl = tmp / "tpl.docx"
    d.save(str(tpl))
    skill = dts.DocxTemplateSkill(str(tmp))
    ins = skill.inspect_template(str(tpl))
    slot = next(s for s in ins["slots"]
                if s["kind"] == "underscores" and "项目名称" in (s.get("label") or ""))
    out = tmp / "out.docx"
    skill.fill_template(
        str(tpl), str(out),
        slot_values={slot["id"]: "示例工程"},
        force_fresh=True,
    )
    body = "\n".join(p.text for p in Document(str(out)).paragraphs)
    assert "12000/" not in body, f"stranded slash should be cleaned: {body!r}"
    assert "12000" in body and "平方米" in body, body
    assert "6/" not in body, f"trailing draft slash should be cleaned: {body!r}"


def test_strip_stranded_draft_marks_preserves_real_separator():
    """A real `XX/YY` separator (no whitespace between the slash and the
    value on each side) must NOT be cleaned — only `value / ≥3 spaces /
    boundary` patterns are stranded draft marks."""
    tmp = Path(tempfile.mkdtemp())
    d = Document()
    d.add_paragraph("檐高/跨度：48米")
    d.add_paragraph("日期：2026/05/20")
    d.add_paragraph("规格：1/3 标准")
    d.add_paragraph("占位：________")
    tpl = tmp / "tpl.docx"
    d.save(str(tpl))
    skill = dts.DocxTemplateSkill(str(tmp))
    ins = skill.inspect_template(str(tpl))
    slot = next(s for s in ins["slots"]
                if s["kind"] == "underscores" and "占位" in (s.get("label") or ""))
    out = tmp / "out.docx"
    skill.fill_template(
        str(tpl), str(out),
        slot_values={slot["id"]: "x"},
        force_fresh=True,
    )
    body = "\n".join(p.text for p in Document(str(out)).paragraphs)
    assert "檐高/跨度" in body, body
    assert "2026/05/20" in body, body
    assert "1/3" in body, body


def test_strip_stranded_draft_marks_at_end_of_paragraph():
    """Trailing `<value>/<whitespace>` at end of paragraph (no boundary
    token after) must also be cleaned. Real CN templates: `京发改第0128号/`
    followed only by trailing whitespace."""
    tmp = Path(tempfile.mkdtemp())
    d = Document()
    d.add_paragraph("批准文号：京发改〔2025〕第0128号/             ")
    d.add_paragraph("占位：________")
    tpl = tmp / "tpl.docx"
    d.save(str(tpl))
    skill = dts.DocxTemplateSkill(str(tmp))
    ins = skill.inspect_template(str(tpl))
    slot = next(s for s in ins["slots"]
                if s["kind"] == "underscores" and "占位" in (s.get("label") or ""))
    out = tmp / "out.docx"
    skill.fill_template(
        str(tpl), str(out),
        slot_values={slot["id"]: "x"},
        force_fresh=True,
    )
    body = "\n".join(p.text for p in Document(str(out)).paragraphs)
    # No `/` should follow the 批准文号 value.
    assert "0128号/" not in body, f"end-of-line draft slash should be cleaned: {body!r}"
    assert "0128号" in body, body


def test_uw_span_stops_at_open_paren_option_marker():
    """`归 <underlined gap> （甲、乙、双）方所有。` — the UW walker must stop
    at `（`, NOT swallow `（甲、乙、双）方所有` until `。`. The option marker
    plus trailing prose ('方所有') must survive the fill."""
    tmp = Path(tempfile.mkdtemp())
    d = Document()
    p = d.add_paragraph()
    p.add_run("第十八条  乙方利用研究开发经费所购置与研究开发工作有关的设备、器材、资料等财产，归")
    r_gap = p.add_run("           ")
    r_gap.font.underline = True
    p.add_run("（甲、乙、双）方所有。")
    tpl = tmp / "tpl.docx"
    d.save(str(tpl))
    skill = dts.DocxTemplateSkill(str(tmp))
    ins = skill.inspect_template(str(tpl))
    blanks = [s for s in ins["slots"] if s["kind"] == "underscores"]
    assert len(blanks) == 1, f"expected 1 slot, got {len(blanks)}: {blanks}"
    out = tmp / "out.docx"
    skill.fill_template(str(tpl), str(out),
                        slot_values={blanks[0]["id"]: "乙"})
    body = "\n".join(p.text for p in Document(str(out)).paragraphs)
    assert "（甲、乙、双）方所有。" in body, (
        f"option marker + trailing prose must survive: {body!r}"
    )
    assert "乙" in body, body


def test_uw_span_stops_at_open_paren_inside_option_body():
    """`1．<gap>（甲、乙、双）方享有...。` — inline option marker after a UW
    gap inside an option-numbered paragraph. Same pattern as above."""
    tmp = Path(tempfile.mkdtemp())
    d = Document()
    p = d.add_paragraph()
    p.add_run("1．")
    r_gap = p.add_run("          ")
    r_gap.font.underline = True
    p.add_run("（甲、乙、双）方享有申请专利的权利。")
    tpl = tmp / "tpl.docx"
    d.save(str(tpl))
    skill = dts.DocxTemplateSkill(str(tmp))
    ins = skill.inspect_template(str(tpl))
    blanks = [s for s in ins["slots"] if s["kind"] == "underscores"]
    assert len(blanks) == 1, f"expected 1 slot, got {len(blanks)}: {blanks}"
    out = tmp / "out.docx"
    skill.fill_template(str(tpl), str(out),
                        slot_values={blanks[0]["id"]: "双"})
    body = "\n".join(p.text for p in Document(str(out)).paragraphs)
    assert "（甲、乙、双）方享有申请专利的权利。" in body, (
        f"option marker + trailing prose must survive: {body!r}"
    )
    assert "双" in body, body


def test_uw_left_anchor_does_not_cross_sentence_period():
    """`...解决。协商、调解不成的，确定按以下第<gap>种方式处理：` — the left
    anchor must NOT cross the `。` from the previous sentence. Without this
    guard, the slot absorbs the prose between `。` and the gap, which then
    gets overwritten by the fill value."""
    tmp = Path(tempfile.mkdtemp())
    d = Document()
    p = d.add_paragraph()
    p.add_run("第二十四条：双方因履行本合同而发生的争议，应协商、调解解决。"
              "协商、调解不成的，确定按以下第")
    r_gap = p.add_run("      ")
    r_gap.font.underline = True
    p.add_run("种方式处理：")
    tpl = tmp / "tpl.docx"
    d.save(str(tpl))
    skill = dts.DocxTemplateSkill(str(tmp))
    ins = skill.inspect_template(str(tpl))
    blanks = [s for s in ins["slots"] if s["kind"] == "underscores"]
    assert len(blanks) == 1, f"expected 1 slot, got {len(blanks)}: {blanks}"
    out = tmp / "out.docx"
    skill.fill_template(str(tpl), str(out),
                        slot_values={blanks[0]["id"]: "1"})
    body = "\n".join(p.text for p in Document(str(out)).paragraphs)
    assert "协商、调解不成的，确定按以下第" in body, (
        f"prose between previous 。 and the gap must survive: {body!r}"
    )
    assert "种方式处理：" in body, body
    assert "1" in body, body


def test_uw_right_walk_does_not_swallow_prose_to_period():
    """`由乙方以<gap>的方式使用。` — Chinese particles (`的`, `方`, `式`) and
    flowing prose contain no boundary chars, so the OLD walk would slurp
    everything up to `。`. The new two-stage walk stops at the trailing-
    whitespace edge instead, leaving `的方式使用` intact."""
    tmp = Path(tempfile.mkdtemp())
    d = Document()
    p = d.add_paragraph()
    p.add_run("第六条  本合同的研究开发经费由乙方以")
    r_gap = p.add_run("              ")
    r_gap.font.underline = True
    p.add_run("的方式使用。")
    tpl = tmp / "tpl.docx"
    d.save(str(tpl))
    skill = dts.DocxTemplateSkill(str(tmp))
    ins = skill.inspect_template(str(tpl))
    blanks = [s for s in ins["slots"] if s["kind"] == "underscores"]
    assert len(blanks) == 1, f"expected 1 slot, got {len(blanks)}: {blanks}"
    out = tmp / "out.docx"
    skill.fill_template(str(tpl), str(out),
                        slot_values={blanks[0]["id"]: "专款专用、独立核算"})
    body = "\n".join(p.text for p in Document(str(out)).paragraphs)
    assert "的方式使用。" in body, (
        f"prose after the gap must survive the fill: {body!r}"
    )
    assert "专款专用、独立核算" in body, body


def test_uw_two_gaps_in_same_paragraph_stay_isolated():
    """Two UW gaps separated by a full sentence (`<gap1>的方式使用。甲方有
    权以<gap2>的方式...`) must NOT merge. Each gap is its own slot, and the
    prose between them must survive."""
    tmp = Path(tempfile.mkdtemp())
    d = Document()
    p = d.add_paragraph()
    p.add_run("第六条  本合同的研究开发经费由乙方以")
    r1 = p.add_run("              ")
    r1.font.underline = True
    p.add_run("的方式使用。甲方有权以")
    r2 = p.add_run("        ")
    r2.font.underline = True
    p.add_run("的方式检查乙方进行研究开发工作和使用研究开发经费的情况，"
              "但不得妨碍乙方的正常工作。")
    tpl = tmp / "tpl.docx"
    d.save(str(tpl))
    skill = dts.DocxTemplateSkill(str(tmp))
    ins = skill.inspect_template(str(tpl))
    blanks = [s for s in ins["slots"] if s["kind"] == "underscores"]
    assert len(blanks) == 2, f"expected 2 separate slots, got {len(blanks)}: {blanks}"
    out = tmp / "out.docx"
    skill.fill_template(
        str(tpl), str(out),
        slot_values={
            blanks[0]["id"]: "专款专用、独立核算",
            blanks[1]["id"]: "定期审查财务报告和项目进度报告",
        },
    )
    body = "\n".join(p.text for p in Document(str(out)).paragraphs)
    assert "的方式使用。" in body, body
    assert "的方式检查乙方进行研究开发工作和使用研究开发经费的情况" in body, body
    assert "但不得妨碍乙方的正常工作。" in body, body
    assert "专款专用、独立核算" in body, body
    assert "定期审查财务报告和项目进度报告" in body, body


def test_uw_span_stops_at_fen_unit():
    """`本合同一式 <underlined gap> 份，具有同等法律效力。` — the gap before
    '份' is the fillable slot. The walker MUST stop at '份' (a unit char) and
    MUST NOT continue past the comma into the rest of the clause; otherwise
    the trailing prose '份，具有同等法律效力' gets sucked into the slot and
    overwritten by the fill value."""
    tmp = Path(tempfile.mkdtemp())
    d = Document()
    p = d.add_paragraph()
    p.add_run("第二十八条  本合同一式")
    r_gap = p.add_run("            ")
    r_gap.font.underline = True
    p.add_run("份，具有同等法律效力。")
    tpl = tmp / "tpl.docx"
    d.save(str(tpl))
    skill = dts.DocxTemplateSkill(str(tmp))
    ins = skill.inspect_template(str(tpl))
    blanks = [s for s in ins["slots"] if s["kind"] == "underscores"]
    assert len(blanks) == 1, f"expected exactly 1 slot, got {len(blanks)}: {blanks}"
    out = tmp / "out.docx"
    skill.fill_template(str(tpl), str(out),
                        slot_values={blanks[0]["id"]: "四"})
    body = "\n".join(p.text for p in Document(str(out)).paragraphs)
    assert "份，具有同等法律效力。" in body, (
        f"trailing clause must survive the fill: {body!r}"
    )
    assert "四" in body, body


# ── Runner ───────────────────────────────────────────────────────────────


def _run_all():
    failed = []
    fns = [(name, fn) for name, fn in globals().items()
           if name.startswith("test_") and callable(fn)]
    for name, fn in fns:
        try:
            fn()
            print(f"  OK   {name}")
        except AssertionError as exc:
            print(f"  FAIL {name}: {exc}")
            failed.append(name)
        except Exception as exc:
            import traceback
            print(f"  ERR  {name}: {type(exc).__name__}: {exc}")
            traceback.print_exc()
            failed.append(name)
    print()
    if failed:
        print(f"{len(failed)} of {len(fns)} test(s) failed: {failed}")
        return 1
    print(f"All {len(fns)} tests passed.")
    return 0


if __name__ == "__main__":
    raise SystemExit(_run_all())
