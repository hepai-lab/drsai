"""
DOCX Template Fill Skill for DocMaster.

Lets DocMaster generate filled DOCX files from a user-uploaded template.
Supports three input styles:

- Jinja (via docxtpl): {{ name }}, {% for x in xs %}…{% endfor %},
  {%tr for row in rows %}…{%tr endtr %} for table rows.
- Bracket: [NAME], [DATE] — literal substitution via python-docx.
- Heuristic slots: for templates that don't use any placeholder syntax —
  underscore lines (`______`), "Label:" with empty tail, and empty table
  cells under a header. The skill returns slot candidates from
  inspect_template; the caller maps slot ids to values and passes
  `slot_values` to fill_template.

Formatting note: bracket and slot fills only modify the runs that span the
matched text; other runs in the same paragraph keep their original
formatting intact.

Public API:
    skill = DocxTemplateSkill(workspace_dir)
    skill.inspect_template(template_path)
    skill.fill_template(template_path, output_path, context, slot_values, mode="auto")
"""

from __future__ import annotations

import logging
import re
import zipfile
from collections import defaultdict
from pathlib import Path
from typing import Any, Callable, Dict, Iterable, List, Optional, Set, Tuple

try:
    from .chinese_amount import (
        chinese_to_decimal,
        decimal_to_chinese,
        format_arabic_amount,
        parse_arabic_amount,
    )
except ImportError:  # support running this file directly
    from chinese_amount import (  # type: ignore[no-redef]
        chinese_to_decimal,
        decimal_to_chinese,
        format_arabic_amount,
        parse_arabic_amount,
    )

logger = logging.getLogger(__name__)


_BRACKET_RE = re.compile(r"\[([A-Z][A-Z0-9_\-]*)\]")
_JINJA_VAR_RE = re.compile(r"\{\{\s*([a-zA-Z_][a-zA-Z0-9_]*)")
_JINJA_LOOP_TOKENS = ("{% for", "{%- for", "{%p for", "{%tr for")
_JINJA_IF_TOKENS = ("{% if", "{%- if", "{%p if", "{%tr if")

# Underscores eligible as fill markers. Includes ASCII '_' and full-width '＿',
# in runs of 3+. A second pattern catches spaced sequences like "_ _ _ _" or
# "＿ ＿ ＿" (2+ separated underscores) which are common in Chinese forms.
_UNDERSCORE_RUN_RE = re.compile(r"[_＿]{3,}|(?:[_＿](?:[ 　]+[_＿]){1,}){1}")
# Label preceding an underscore run, e.g. "Patient Name: ______".
# Captures the label text (without the trailing colon).
_LABEL_BEFORE_RE = re.compile(r"([^\s:：][^:：\n]{0,79})\s*[:：]\s*$")
# Whole paragraph ending in "Label:" with nothing after.
_LABEL_ONLY_RE = re.compile(r"^\s*(.{1,80}?)\s*[:：]\s*$")

# Signature / seal indicators — these slots must be LEFT BLANK rather than
# filled. Looked up against the label found before the underscore and against
# any short trailing text after it (e.g. "________  (signature)").
_SIGNATURE_RE = re.compile(
    r"(?:签\s*字|签\s*名|签\s*章|盖\s*章|落\s*款|手\s*印|"
    r"signature|signed\s*by|sign\s*here|signatory|autograph|initials?|seal)",
    re.IGNORECASE,
)

# <placeholder> or 《占位符》 — angle-bracketed slot tokens
_ANGLE_BRACKET_RE = re.compile(r"<([^<>\n]{1,80})>|《([^《》\n]{1,80})》")

# Seal / stamp / date-placeholder cell — Chinese contract signature blocks
# often contain a vertically-merged cell with template text like
# "合 同 章 \n\n 年  月  日". This is a fillable region (the user signs/stamps
# and writes the date here), not real body content. Detected and emitted as
# a `placeholder_cell` slot pointing at the vMerge anchor.
_SEAL_PLACEHOLDER_RE = re.compile(
    r"合\s*同\s*章|公\s*章|印\s*章|盖\s*章|签\s*字\s*盖\s*章|seal\s+here|stamp\s+here",
    re.IGNORECASE,
)
_DATE_PLACEHOLDER_RE = re.compile(r"年\s{1,}月\s{1,}日|yyyy[/-]mm[/-]dd", re.IGNORECASE)

# Date scaffold inside a paragraph: a sequence of whitespace gaps with bare
# CJK date markers (`年 / 月 / 日`) acting as in-place scaffolding. Common
# in CN forms — the author writes `签订日期：    年    月    日` and expects
# the user to fill in `YYYY 年 MM 月 DD 日` by replacing the gaps. Without
# detecting the whole scaffold we end up filling only the first gap and
# producing `2026年5月14日年    月    日` — the year/month/day markers from
# the template stay behind.
#
# We require a leading whitespace gap (≥2) before `年` so prose like
# `2024年5月` (no gap) doesn't match.
_DATE_SCAFFOLD_RE = re.compile(
    r"[ \t　]{2,}年"
    r"(?:[ \t　]+月)?"
    r"(?:[ \t　]+日)?"
)

# Inline whitespace blank: a whitespace run (3+ chars) bounded by a label
# colon, currency symbol, or amount-label marker on the left, and sentence
# punctuation, a unit marker, or another amount-label on the right. Used
# to detect fields like:
#   "交货地点：         ；乙方负责..."   → blank between '：' and '；'
#   "本项目总经费：    元（元）"          → blank between '：' and '元'
#   "即￥        元（小写）"              → blank between '￥' and '元'
#   "总计金额（含税）：（大写）   （小写）   "
#       → two blanks: one after '（大写）' bounded by '（小写）',
#         one after '（小写）' bounded by '$' (end of paragraph).
# The gap may carry underline formatting (caught earlier by
# _find_underlined_whitespace_spans); this regex is the no-underline fallback.
# Extension: a `/` or `\` between two whitespace runs is a draft mark
# placed by the template author ("建筑面积：    /    平方米；") — the gap
# on either side of the slash is fillable, and the value should REPLACE
# the slash. Detected here so the gap+slash+gap region becomes one slot.
_INLINE_BLANK_RE = re.compile(
    r"(?P<lm>[:：¥$￥€£]|（大写）|\(大写\)|（小写）|\(小写\))"
    r"(?P<gap>[ \t　]{2,}[/\\][ \t　]{2,}|[ \t　]{3,})"
    r"(?=[;；,，.。、元年月日天周个种号份％%米]"
    r"|工作日|平方米|小时|分钟|（小写）|\(小写\)|（大写）|\(大写\)|$)"
)
# Match an amount-label marker on its own (for clean-label slot tagging).
_DAXIE_MARKER_RE = re.compile(r"^[（(]大[写寫][)）]$")
_XIAOXIE_MARKER_RE = re.compile(r"^[（(]小[写寫][)）]$")

# Phrase bank for placeholder-like prose (bilingual). Matched as whole paragraph
# OR as a parenthesised tail. Case-insensitive.
_PLACEHOLDER_PHRASES = [
    r"your\s+(?:text|name|content|answer|response|bio|title|address|email|phone)\s+(?:here|goes\s+here)?",
    r"insert\s+(?:your\s+)?[\w\s]{1,40}\s+here",
    r"provide\s+(?:your\s+|a\s+|the\s+)?[\w\s]{1,40}\s+here",
    r"replace\s+(?:this|with)\s+[\w\s]{1,40}",
    r"enter\s+(?:your\s+)?[\w\s]{1,40}\s+here",
    r"(?:click|tap)\s+(?:here\s+)?to\s+(?:enter|add)\s+[\w\s]{1,40}",
    r"to\s+be\s+(?:completed|filled|determined|added)",
    r"\bt\.?b\.?d\.?\b",
    r"\bt\.?o\.?d\.?o\b",
    r"lorem\s+ipsum[\w\s,\.]*",
    r"\[\s*(?:insert|enter|provide|placeholder|your)\b[^\]]{1,80}\]",
    r"\(\s*(?:insert|enter|provide|describe|fill\s+in|replace)\b[^)]{1,80}\)",
    # Chinese variants
    r"请\s*(?:填写|输入|填入|补充|说明|填空|完善)[一-鿿\w，。、\s]{0,40}",
    r"待\s*(?:填写|补充|完成|确定)",
    r"此处\s*(?:填写|输入|填入|添加|填空)[一-鿿\w，。、\s]{0,30}",
    r"（\s*(?:请填写|此处|填写|输入|示例)[一-鿿\w，。、\s]{0,40}\s*）",
]
_RE_PHRASE_BANK = re.compile(
    "|".join(f"(?:{p})" for p in _PLACEHOLDER_PHRASES),
    re.IGNORECASE,
)

# Phrase bank for instructional / meta text that should be removed before publishing.
_REMOVAL_PHRASES = [
    r"delete\s+(?:this|the\s+following|before|prior\s+to)\b[^.\n]{0,80}",
    r"remove\s+(?:this|the\s+following|before|prior\s+to|in\s+final)\b[^.\n]{0,80}",
    r"note\s+to\s+(?:author|reviewer|editor|self|reader)\b[^.\n]{0,200}",
    r"\b(?:template\s+)?instructions?\s*[:：]",
    r"this\s+section\s+is\s+(?:for|only|optional|a\s+placeholder|an\s+example)\b[^.\n]{0,200}",
    r"for\s+(?:internal|template)\s+use\s+only\b[^.\n]{0,80}",
    r"example\s+only[\s\-—:：]",
    r"sample\s+(?:text|content|paragraph)\b[^.\n]{0,80}",
    r"e\.g\.[\s\-—:：]",
    r"^\s*example\s*[:：]",
    # Chinese variants
    r"请\s*删除\s*(?:本段|此段|此句|本节|这段|本部分)?",
    r"模板\s*(?:说明|示例|提示|备注)",
    r"仅供参考",
    r"以下\s*(?:内容|部分)?\s*为?\s*示例",
    r"使用\s*前\s*请\s*删除",
    r"作者\s*备注\s*[:：]",
    # "应删除" / "应当删除" — directive form (vs. polite 请删除).
    r"应\s*(?:当\s*)?(?:删除|删去|去除|去掉)",
    # "此句话非正文" / "本段非正文" — explicit "this is not body text" markers.
    r"(?:此|这|该|本)\s*(?:句话|句|段|节|部分|条)\s*(?:非|不是|不属于)\s*正文",
    # "非正文，应删除" — combined marker (covered by the two above, but the
    # joined form is the most common phrasing in CN ministry templates).
    r"非\s*正文[^。\n]{0,20}应\s*(?:删除|删去)",
    # "正式文件（中/里/内）...删除" — drafting note explicitly scoped to
    # "the formal/published version". Common phrasing:
    #   "正式文件中这句话删除"
    #   "正式版本删除此段"
    #   "正式稿删去"
    r"正式\s*(?:文件|文档|稿|版本|版|提交\s*稿)[^。\n]{0,40}(?:删除|删去|去掉)",
    # Bare "此/这/该/本 + (句|句话|段|节|条|行|部分) + 删除"
    # — directive form without preceding "应", "请", "需". Covers
    # "此句话删除" / "这段删除" / "本行删去" / "该节去掉".
    r"(?:此|这|该|本)\s*(?:句|句话|段|节|条|行|部分)\s*(?:话)?\s*(?:删除|删去|去除|去掉)",
    # "空着" / "留空" — drafting note telling the reader to leave the field
    # blank (e.g. "签订日期：    年  月  日（空着，后盖章的一方手写即可）").
    # These are instructions to the signer, not template prose — strip them.
    r"空\s*着",
    r"留\s*空",
    # "由X方手写/填写/签字/盖章" — explicit instruction that some party will
    # complete the field by hand at signing. Same intent as 空着 but spelt out.
    # The 'X方' alternation covers 甲/乙/丙/丁/双 + 方, '当事人', '签订方',
    # 'X方' written out by role ('盖章方', '签字方', '收货方', etc.), or up
    # to 6 CJK chars + 方 (catches '招标方', '中标方', '承包方', '发包方').
    r"由\s*(?:[甲乙丙丁双]\s*方|当事人|签订?方|[一-鿿]{1,6}方)[^。\n]{0,30}"
    r"(?:手\s*写|填\s*写|签\s*字|盖\s*章)",
    # "(后)盖章(的)(一)?方 ... 手写/填写" — "the side that seals later
    # will handwrite this." Covers "后盖章的一方手写即可",
    # "盖章方填写", "盖章一方手写".
    r"(?:后\s*)?盖\s*章\s*(?:的\s*)?(?:一?方)?[^。\n]{0,15}"
    r"(?:手\s*写|填\s*写)",
    # In-paren drafting notes — author addressing the future reader of the
    # template ("here we suggest", "to avoid mistakes", "you can leave it
    # blank"). These are the most common CN-template highlight-yellow notes
    # and are NEVER part of the final document, but the inner text doesn't
    # contain "删除" so the bank used to miss them.
    #
    # Anchoring on a clear instructional verb keeps false positives down —
    # parenthetical proper nouns ("(中国科学院高能物理研究所)"), citations
    # ("(参见附件1)" / "(下同)"), dates and code identifiers don't contain any
    # of these stems, so they stay as template prose.
    r"(?:这\s*里|此\s*处)\s*(?:建议|提示|应当?|应该|说明|注意)",
    # "建议 + verb" — "建议不要写", "建议填入", "建议另起", "建议保留".
    # Requires a following action verb so a parenthetical that only contains
    # "建议" as a noun ("(我的建议)") doesn't trip the bank.
    r"建\s*议\s*(?:不\s*要|不\s*写|不\s*填|写|填\s*入|填\s*写|保\s*留|"
    r"删\s*除|删\s*去|另\s*起|改\s*为|采\s*用|参\s*考)",
    # "避免 + 错/不一致/混淆" — explanatory drafting note ("...to avoid
    # mistakes / inconsistency with the invoice").
    r"避\s*免\s*[^。\n]{0,20}"
    r"(?:写\s*错|出\s*错|错\s*误|混\s*淆|歧\s*义|不\s*一\s*致|与[^。\n]{1,20}不\s*一\s*致)",
    # "可(不|以不) + 填/写" — "(this field can be left blank)".
    r"可\s*(?:不|以\s*不)\s*(?:填|写|填\s*写)",
    # "仅供参考" already in the bank above; add "供参考(用)" for plural variants
    # ("以下条款仅供参考用").
    r"供\s*参\s*考(?:\s*用)?",
]
_RE_REMOVAL_BANK = re.compile(
    "|".join(f"(?:{p})" for p in _REMOVAL_PHRASES),
    re.IGNORECASE,
)

# "Keep the following table on one page" instructions. When a removed
# instruction paragraph matches this regex, the next sibling <w:tbl> after
# the paragraph is hardened: cantSplit on every row + keepNext on every
# paragraph (except the last) so Word renders the whole table on one page.
_KEEP_TABLE_TOGETHER_RE = re.compile(
    # 不/不要/不应/不能/不可/不得/不许/勿 — catch the full negation family
    # since CN drafting notes vary ("下表不跨页", "下表不应跨页", "下表勿跨页").
    r"下\s*表\s*(?:不(?:\s*(?:要|应|能|可|得|许))?\s*跨\s*页"
    r"|不(?:\s*(?:要|应|能|可|得|许))?\s*分\s*页"
    r"|保持\s*在?\s*同\s*一\s*页"
    r"|勿\s*跨\s*页)"
    r"|表\s*格\s*(?:不(?:\s*(?:要|应|能|可|得|许))?\s*跨\s*页"
    r"|不(?:\s*(?:要|应|能|可|得|许))?\s*分\s*页"
    r"|保持\s*在?\s*同\s*一\s*页)"
    r"|本\s*表\s*(?:不(?:\s*(?:要|应|能|可|得|许))?\s*跨\s*页"
    r"|不(?:\s*(?:要|应|能|可|得|许))?\s*分\s*页)"
    r"|确\s*保\s*下\s*表\s*不(?:\s*(?:要|应|能|可|得|许))?\s*跨\s*页"
    r"|keep\s+(?:the\s+)?(?:following\s+|next\s+)?table\s+(?:together|on\s+one\s+page)"
    r"|table\s+(?:must\s+)?(?:stay|fit)\s+on\s+(?:one|a\s+single)\s+page",
    re.IGNORECASE,
)

# Word's built-in heading style names; covers both English and Chinese variants.
_HEADING_STYLE_RE = re.compile(r"^(?:Heading\s*\d+|Title|Subtitle|标题|副标题)", re.IGNORECASE)

# Headings that *introduce* a following list/table — body text should NOT go
# directly after them; the structured content (rows / list items) is the right
# fill target. e.g. "一、甲方委托乙方提供以下维修服务：" → the table below.
_LIST_INTRO_RE = re.compile(
    r"(?:以下|如下|下表|下列|下面|"
    r"following|below|as\s+follows)",
    re.IGNORECASE,
)

# ── Scaffold patterns inside a highlighted span ──────────────────────────────
# When a highlighted span looks like "<variable><scaffold>" (e.g. "15个工作日",
# "¥850", "50%", "2025年5月14日"), we split it so the user's bare reply (just
# the number) still produces a correctly-scaffolded output ("20个工作日").
#
# Each entry: (compiled regex with named groups 'pre','var','suf', kind_label)
_SCAFFOLD_PATTERNS = [
    # Currency prefix + number  (¥850, $1,200, €99.50)
    (re.compile(r"^(?P<pre>[¥￥$€£])\s*(?P<var>[\d,]+(?:\.\d+)?)(?P<suf>)$"), "currency"),
    # Number + Chinese unit  (15个工作日, 30天, 5次, 100元)
    (re.compile(r"^(?P<pre>)(?P<var>\d+(?:\.\d+)?)\s*(?P<suf>[一-鿿]{1,8})$"), "number_zh_unit"),
    # Number + ASCII unit / percent  (50%, 10kg, 3 days, 12 hrs)
    (re.compile(
        r"^(?P<pre>)(?P<var>\d+(?:\.\d+)?)\s*"
        r"(?P<suf>%|kg|km|m|cm|mm|g|lbs?|days?|weeks?|months?|years?|hours?|hrs?|mins?|sec(?:ond)?s?)$",
        re.IGNORECASE), "number_unit"),
    # Date — YYYY年MM月DD日 — only the year/month/day numbers are variable;
    # we surface the whole thing as one scaffold (full string is the variable,
    # no auto-attach). Marked so the agent knows.
    (re.compile(r"^(?P<pre>)(?P<var>\d{4}年\d{1,2}月\d{1,2}日)(?P<suf>)$"), "date_zh"),
    (re.compile(r"^(?P<pre>)(?P<var>\d{4}-\d{1,2}-\d{1,2})(?P<suf>)$"), "date_iso"),
    (re.compile(r"^(?P<pre>)(?P<var>\d{4}/\d{1,2}/\d{1,2})(?P<suf>)$"), "date_slash"),
]

# A "bare variable" reply from the user (just digits, optionally with separators,
# decimals, or a single currency-style sign). If user-provided value matches
# this AND the original had non-empty scaffolding, we auto-attach the scaffold.
_BARE_NUMERIC_RE = re.compile(r"^[+-]?[\d,]+(?:\.\d+)?$")

# Labels marking the two halves of a paired money slot. Detected in
# slot.label / slot.context to find 大写/小写 pairs for deterministic
# reconciliation (see _reconcile_amount_pairs).
_DAXIE_LABEL_RE = re.compile(r"大\s*[写寫]|in\s*words|capital\s*amount|capitalized", re.IGNORECASE)
_XIAOXIE_LABEL_RE = re.compile(r"小\s*[写寫]|in\s*figures|in\s*numerals|金\s*额|[¥￥]")

# Two-options-pattern detection. Chinese contract templates often have:
#
#   （以下两种选择适合的一种，可以根据需要再进行改写，不选的一种请删除）
#   （●第一种：方案A …）
#     option 1 body paragraphs …
#   （●第二种：方案B …）
#     option 2 body paragraphs …
#
# The fill pass keeps the chosen option's body and drops everything else.
_OPTION_PROMPT_RE = re.compile(
    r"(?:以下|下列)[\s（()）]*[两二三四五六七八九十]\s*种.*?(?:选\s*择|选)[^。]*?请\s*删\s*除"
    r"|[两二三四五六七八九十]\s*选\s*一"
)
_OPTION_HEADER_RE = re.compile(
    r"^\s*[（(]?\s*[●○•·▪◆◇■□*\-—]*\s*第\s*([一二三四五六七八九十])\s*[种種项項条條]"
)
_SECTION_BREAK_RE = re.compile(
    # CJK list marker followed by 、 ． or . — "一、" / "二．" — but NOT
    # sub-section like "一.二" (digit after the period would be a sub-key).
    r"^\s*(?:[一二三四五六七八九十]+\s*[、．](?!\d)"
    r"|[一二三四五六七八九十]+\s*\.\s+"
    # Top-level digit list marker: "1、" "1．" "1. " — but never "8.1".
    r"|\d+\s*[、．](?!\d)"
    r"|\d+\s*\.\s+(?=\D)"
    # "第N章/节/条"
    r"|第\s*[一二三四五六七八九十]+\s*[章节節条條]"
    r"|附\s*[件录録]"
    r"|[甲乙丙丁]\s*方\s*[：:])"
)

# Chinese-number map used by _normalize_option_choice (small/local; do not
# import from chinese_amount to keep this self-contained).
_CN_NUM_TO_INT = {
    "一": 1, "壹": 1, "二": 2, "贰": 2, "两": 2, "三": 3, "叁": 3,
    "四": 4, "肆": 4, "五": 5, "伍": 5, "六": 6, "陆": 6, "陸": 6,
    "七": 7, "柒": 7, "八": 8, "捌": 8, "九": 9, "玖": 9, "十": 10, "拾": 10,
}
_ENGLISH_ORDINAL_TO_INT = {
    "first": 1, "1st": 1, "second": 2, "2nd": 2, "third": 3, "3rd": 3,
    "fourth": 4, "4th": 4, "fifth": 5, "5th": 5, "sixth": 6, "6th": 6,
    "seventh": 7, "7th": 7, "eighth": 8, "8th": 8, "ninth": 9, "9th": 9,
    "tenth": 10, "10th": 10,
}


def _normalize_option_choice(value: Any, n_options: int) -> Optional[int]:
    """Coerce a user-provided option choice to a 1-based index ≤ n_options.

    Accepts: int 1..n; numeric strings; Chinese ordinals ('一', '第一种',
    '第二种', '贰' …); English ordinals ('first', '2nd'); plain English
    cardinals ('1', '2'). Returns None on garbage / out-of-range.
    """
    if value is None:
        return None
    if isinstance(value, bool):  # bool is an int subclass — reject explicitly
        return None
    if isinstance(value, int):
        return value if 1 <= value <= n_options else None
    if not isinstance(value, str):
        return None
    t = value.strip()
    if not t:
        return None
    # Drop common decorators
    t = t.strip("（()）()[]【】《》〈〉「」 　")
    # "第N种" / "第N项" / etc.
    m = re.match(r"^第\s*([一二三四五六七八九十\d]+)\s*[种種项項条條]?$", t)
    if m:
        token = m.group(1)
        if token.isdigit():
            n = int(token)
            return n if 1 <= n <= n_options else None
        if token in _CN_NUM_TO_INT:
            n = _CN_NUM_TO_INT[token]
            return n if 1 <= n <= n_options else None
        return None
    # Bare digit
    if t.isdigit():
        n = int(t)
        return n if 1 <= n <= n_options else None
    # Bare Chinese ordinal
    if t in _CN_NUM_TO_INT:
        n = _CN_NUM_TO_INT[t]
        return n if 1 <= n <= n_options else None
    # English ordinal/cardinal
    low = t.lower()
    if low in _ENGLISH_ORDINAL_TO_INT:
        n = _ENGLISH_ORDINAL_TO_INT[low]
        return n if 1 <= n <= n_options else None
    return None


def _split_scaffold(text: str) -> Optional[Dict[str, str]]:
    """If `text` matches one of the scaffold patterns, return
    {'kind': ..., 'pre': ..., 'var': ..., 'suf': ...}. Else None.
    """
    if not text:
        return None
    t = text.strip()
    if not t:
        return None
    for pat, kind in _SCAFFOLD_PATTERNS:
        m = pat.match(t)
        if m:
            return {
                "kind": kind,
                "pre": m.group("pre") or "",
                "var": m.group("var") or "",
                "suf": m.group("suf") or "",
            }
    return None


def _reconstruct_with_scaffold(scaffold: Dict[str, str], user_value: str) -> str:
    """Given a scaffold split (pre/var/suf) and a user-provided value, decide
    whether to reattach the scaffolding.

    Rule: if the original variable was numeric/currency-like AND the user value
    is a bare number (no letters, no Chinese, no existing prefix/suffix), wrap
    it: pre + user_value + suf. Otherwise pass user_value through unchanged
    (the user gave us a full replacement on purpose).
    """
    pre, suf = scaffold.get("pre", ""), scaffold.get("suf", "")
    if not (pre or suf):
        return user_value
    val = user_value.strip()
    if not val:
        return user_value
    # If the user's reply already contains the scaffold parts, don't double-wrap
    if pre and val.startswith(pre):
        return user_value
    if suf and val.endswith(suf):
        return user_value
    # Only auto-wrap when the user gave a pure-number reply (the most common
    # source of the "lost units" bug). For dates (no pre/suf) this no-ops.
    if _BARE_NUMERIC_RE.match(val):
        return f"{pre}{val}{suf}"
    return user_value


class DocxTemplateSkill:
    """Fill a DOCX template to generate a new DOCX."""

    def __init__(self, workspace_dir: Optional[str] = None) -> None:
        if workspace_dir is None:
            workspace_dir = str(Path(__file__).resolve().parent.parent / "workspace")
        self.workspace_dir = Path(workspace_dir)
        self.workspace_dir.mkdir(parents=True, exist_ok=True)

    # ------------------------------------------------------------------ #
    # Inspection                                                          #
    # ------------------------------------------------------------------ #

    def inspect_template(self, template_path: str) -> Dict[str, Any]:
        path = Path(template_path)
        if not path.exists():
            return {
                "success": False,
                "error": "File not found",
                "message": f"Template not found: {template_path}",
            }
        if path.suffix.lower() != ".docx":
            return {
                "success": False,
                "error": "Unsupported file type",
                "message": f"Template must be a .docx file (got {path.suffix})",
            }

        try:
            xml_blobs = _read_docx_text_xml(path)
        except Exception as exc:
            return {
                "success": False,
                "error": str(exc),
                "message": f"Could not read template XML: {exc}",
            }

        merged_xml = "\n".join(xml_blobs)
        jinja_vars = _scan_jinja_vars(merged_xml)
        has_loops = any(tok in merged_xml for tok in _JINJA_LOOP_TOKENS)
        has_conditionals = any(tok in merged_xml for tok in _JINJA_IF_TOKENS)

        try:
            from docx import Document
        except ImportError as exc:
            return {
                "success": False,
                "error": "python-docx not installed",
                "message": f"python-docx required for template inspection: {exc}",
            }

        try:
            doc = Document(str(path))
        except Exception as exc:
            return {
                "success": False,
                "error": str(exc),
                "message": f"Could not open template: {exc}",
            }

        paragraph_texts = ["".join((r.text or "") for r in p.runs)
                           for p in _iter_all_paragraphs(doc)]
        bracket_tokens = _scan_bracket_tokens(paragraph_texts)
        slots_full = _scan_slots(doc)
        slots_public = [_slot_public_view(s) for s in slots_full]
        removals_full = _scan_removals(doc)
        # Suppress removals whose paragraph is already owned by an
        # option_choice slot — the slot fill handles their deletion.
        consumed = _collect_option_choice_paragraph_ids(slots_full)
        if consumed:
            removals_full = [
                r for r in removals_full
                if _paragraph_element_id(
                    (r.get("_meta") or {}).get("paragraph")
                ) not in consumed
            ]
        # Suppress fillable slots that target a paragraph already marked for
        # whole-paragraph removal. Otherwise a meta-note like "（…此句话非正
        # 文，应删除）" surfaces as BOTH a fillable highlighted slot AND a
        # removal — the agent typically fills one and leaves the meta-note
        # behind. The removal owns the paragraph; drop the slot.
        removal_para_ids = {
            _paragraph_element_id((r.get("_meta") or {}).get("paragraph"))
            for r in removals_full
            if r.get("kind") == "instruction_paragraph"
        }
        removal_para_ids.discard(None)
        if removal_para_ids:
            slots_full = [
                s for s in slots_full
                if _paragraph_element_id(
                    (s.get("_meta") or {}).get("paragraph")
                ) not in removal_para_ids
            ]
            slots_public = [_slot_public_view(s) for s in slots_full]
        removals_public = [_slot_public_view(r) for r in removals_full]

        mode_detected = _detect_mode(jinja_vars, bracket_tokens, has_loops, has_conditionals)

        warnings: List[str] = []
        if mode_detected == "both":
            warnings.append(
                "Template contains both Jinja ({{ }}) and bracket ([TOKEN]) placeholders; "
                "auto mode will render Jinja first, then bracket-pass."
            )
        if mode_detected == "none" and not slots_public:
            warnings.append(
                "No placeholders or fillable slots detected. Ask the user to point out "
                "where values should go, or to mark fields with [TOKEN] / {{ var }} / underscores."
            )
        elif mode_detected == "none" and slots_public:
            warnings.append(
                f"No explicit placeholders found, but {len(slots_public)} heuristic slot(s) "
                "were detected. Confirm each with the user before filling."
            )
        if removals_public:
            warnings.append(
                f"{len(removals_public)} instructional/meta-text passage(s) detected as "
                "removal candidates. Read each to the user and confirm before deleting; "
                "pass the confirmed ids via fill_docx_template_tool's removal_ids."
            )

        return {
            "success": True,
            "template_path": str(path),
            "mode_detected": mode_detected,
            "jinja_variables": sorted(jinja_vars),
            "bracket_tokens": sorted(bracket_tokens),
            "has_loops": has_loops,
            "has_conditionals": has_conditionals,
            "slots": slots_public,
            "removals": removals_public,
            "warnings": warnings,
            "message": (
                f"mode={mode_detected}, "
                f"{len(jinja_vars)} jinja var(s), "
                f"{len(bracket_tokens)} bracket token(s), "
                f"{len(slots_public)} heuristic slot(s), "
                f"{len(removals_public)} removal candidate(s)"
            ),
        }

    # ------------------------------------------------------------------ #
    # Fill                                                                #
    # ------------------------------------------------------------------ #

    def fill_template(
        self,
        template_path: str,
        output_path: str,
        context: Optional[Dict[str, Any]] = None,
        mode: str = "auto",
        slot_values: Optional[Dict[str, Any]] = None,
        removal_ids: Optional[List[str]] = None,
        force_fresh: bool = False,
    ) -> Dict[str, Any]:
        path = Path(template_path)
        if not path.exists():
            return {
                "success": False,
                "error": "File not found",
                "message": f"Template not found: {template_path}",
            }
        context = context or {}
        slot_values = slot_values or {}
        removal_ids = list(removal_ids) if removal_ids else []

        # Chunked-fill continuation: if the same output_path already exists
        # from a previous fill_template call, switch the source to that
        # partial-fill so prior batches are preserved. Without this, an
        # agent that splits a large fill across multiple tool calls — each
        # using the original template_path — would wipe out earlier
        # batches on every call.
        #
        # IMPORTANT: when the source is swapped, slot ids shift because
        # the per-scan counter only sees REMAINING blank slots. So we ALSO
        # accept the original (canonical) id as the lookup key by matching
        # on the descriptive label tail. The caller can pass either the
        # canonical id from the FIRST inspect_template call or a fresh id
        # from a re-inspection of the partial doc — both resolve correctly.
        #
        # `force_fresh=True` is the agent's escape hatch: skip the auto-
        # detect entirely and just overwrite output_path with a fresh fill
        # from the original template. Past sessions had the agent loop on
        # this auto-detect — fill went wrong, agent retried with the same
        # output_path, the tool quietly swapped to "continuation", slot ids
        # shifted, and the loop never converged.
        chunked_continuation = False
        canonical_id_map: Optional[Dict[str, str]] = None
        out_path_check = Path(output_path)
        if (
            not force_fresh
            and out_path_check.exists()
            and out_path_check.resolve() != path.resolve()
            and slot_values
            and not context
            and not removal_ids
        ):
            try:
                with zipfile.ZipFile(str(out_path_check)) as _z:
                    is_docx = "word/document.xml" in _z.namelist()
            except (zipfile.BadZipFile, OSError):
                is_docx = False
            if is_docx:
                canonical_id_map = _build_canonical_id_map(
                    str(path), str(out_path_check)
                )
                if canonical_id_map is not None:
                    path = out_path_check
                    chunked_continuation = True
        if not isinstance(context, dict):
            return {
                "success": False,
                "error": "Invalid context",
                "message": "context must be a dict",
            }
        if not isinstance(slot_values, dict):
            return {
                "success": False,
                "error": "Invalid slot_values",
                "message": "slot_values must be a dict of {slot_id: value}",
            }

        out_path = Path(output_path)
        out_path.parent.mkdir(parents=True, exist_ok=True)

        if mode not in ("auto", "jinja", "bracket"):
            return {
                "success": False,
                "error": "Invalid mode",
                "message": f"mode must be one of 'auto', 'jinja', 'bracket' (got {mode!r})",
            }

        inspect = self.inspect_template(str(path))
        if not inspect.get("success"):
            return inspect
        detected = inspect["mode_detected"]
        warnings: List[str] = list(inspect.get("warnings", []))
        if chunked_continuation:
            warnings.append(
                f"⚠ CONTINUATION MODE: output_path {out_path_check} already "
                f"exists, so this fill is treated as a continuation of the "
                f"prior partial fill (source swapped to the partial doc). "
                f"If you did NOT intend this — i.e. you wanted to start over "
                f"from the original template — re-call fill_template with "
                f"`force_fresh=True` (do NOT use run_bash to delete the "
                f"file) or pick a different output_path."
            )
            if canonical_id_map and slot_values:
                # Translate canonical ids (from the FIRST inspect of the
                # original template) to the partial doc's current ids. The
                # counter shifts when prior fills consume blank markers, so
                # an id like 'slot_004_签订时间' from the canonical scan may
                # be e.g. 'slot_002_签订时间' after three earlier fills.
                translated: Dict[str, Any] = {}
                translations: List[str] = []
                for k, v in slot_values.items():
                    new_k = canonical_id_map.get(k, k)
                    translated[new_k] = v
                    if new_k != k:
                        translations.append(f"{k} → {new_k}")
                slot_values = translated
                if translations:
                    warnings.append(
                        "Translated canonical slot ids to current ids in "
                        "partial doc: " + ", ".join(translations[:10])
                        + ("..." if len(translations) > 10 else "")
                    )
        # Selected removals: keep the original inspect ids so we can replay them
        # on a fresh scan of the output document after fill passes complete.
        selected_removal_keys: List[Tuple[str, str, str]] = []
        if removal_ids:
            by_id = {r["id"]: r for r in inspect.get("removals", [])}
            unknown = [rid for rid in removal_ids if rid not in by_id]
            if unknown:
                warnings.append(
                    f"Unknown removal id(s): {unknown}; they will be ignored."
                )
            for rid in removal_ids:
                r = by_id.get(rid)
                if r:
                    selected_removal_keys.append(
                        (r["kind"], r["reason"], r["text"])
                    )

        nothing_to_fill = (
            detected == "none"
            and not slot_values
            and not selected_removal_keys
            and mode == "auto"
        )
        if nothing_to_fill:
            return {
                "success": False,
                "template_path": str(path),
                "output_path": str(out_path),
                "mode_used": None,
                "warnings": warnings,
                "message": (
                    "No placeholders detected and no slot_values/removal_ids provided; "
                    "nothing to fill. Either add {{ name }} / [NAME] / underscore placeholders, "
                    "or call inspect_template and pass slot_values / removal_ids."
                ),
            }

        effective_mode = mode
        if effective_mode == "auto":
            if detected in ("jinja", "both"):
                effective_mode = "jinja"
            elif detected == "bracket":
                effective_mode = "bracket"
            else:
                effective_mode = "bracket"  # slot-only path uses bracket loader

        slot_report: Dict[str, Any] = {}

        if effective_mode == "jinja":
            jinja_result = self._fill_jinja(path, out_path, context)
            if not jinja_result.get("success"):
                return {**jinja_result, "warnings": warnings}
            if mode == "auto" and detected == "both":
                bracket_pass = self._fill_brackets(out_path, out_path, context, slot_values=None)
                if bracket_pass.get("success"):
                    bracket_filled = set(bracket_pass.get("filled_keys", []))
                    jinja_result["filled_keys"] = sorted(
                        set(jinja_result.get("filled_keys", [])) | bracket_filled
                    )
                    jinja_result["missing_keys"] = sorted(
                        set(jinja_result.get("missing_keys", []))
                        | set(bracket_pass.get("missing_keys", []))
                    )
                    jinja_result["unused_keys"] = sorted(
                        set(jinja_result.get("unused_keys", [])) - bracket_filled
                    )
                    jinja_result["mode_used"] = "both"
                else:
                    warnings.append(
                        f"Jinja render succeeded but bracket-pass failed: "
                        f"{bracket_pass.get('error', 'unknown error')}"
                    )
            if slot_values:
                slot_pass = self._fill_slots_in_place(out_path, slot_values)
                slot_report = slot_pass
                if not slot_pass.get("success"):
                    warnings.append(
                        f"Slot fill failed: {slot_pass.get('error', 'unknown error')}"
                    )
                warnings.extend(slot_pass.get("reconciliation_notes") or [])
            if selected_removal_keys:
                removal_pass = _apply_removals_in_place(
                    out_path, selected_removal_keys, removal_ids
                )
                jinja_result["removals_applied"] = removal_pass.get("applied", [])
                jinja_result["removals_skipped"] = removal_pass.get("skipped", [])
            jinja_result["warnings"] = warnings
            if slot_report:
                jinja_result["slot_fill"] = slot_report
            jinja_result["chunked_continuation"] = chunked_continuation
            if chunked_continuation:
                jinja_result["continuation_notice"] = (
                    "Continuation mode triggered: output_path already existed "
                    "and was used as the source. If unintended, retry with "
                    "force_fresh=True."
                )
            return jinja_result

        # bracket / slot-only path: a single python-docx load handles both
        bracket_result = self._fill_brackets(path, out_path, context, slot_values=slot_values)
        if bracket_result.get("success") and selected_removal_keys:
            removal_pass = _apply_removals_in_place(
                out_path, selected_removal_keys, removal_ids
            )
            bracket_result["removals_applied"] = removal_pass.get("applied", [])
            bracket_result["removals_skipped"] = removal_pass.get("skipped", [])
        warnings.extend(
            (bracket_result.get("slot_fill") or {}).get("reconciliation_notes") or []
        )
        bracket_result["warnings"] = warnings
        bracket_result["chunked_continuation"] = chunked_continuation
        if chunked_continuation:
            bracket_result["continuation_notice"] = (
                "Continuation mode triggered: output_path already existed "
                "and was used as the source. If unintended, retry with "
                "force_fresh=True."
            )
        return bracket_result

    # ------------------------------------------------------------------ #
    # Jinja path (docxtpl)                                                #
    # ------------------------------------------------------------------ #

    def _fill_jinja(self, template_path: Path, output_path: Path, context: Dict[str, Any]) -> Dict[str, Any]:
        try:
            from docxtpl import DocxTemplate
        except ImportError as exc:
            return {
                "success": False,
                "error": "docxtpl not installed",
                "message": (
                    "Jinja-mode template filling requires docxtpl. "
                    "Install with: pip install docxtpl"
                ),
                "import_error": str(exc),
            }

        try:
            tpl = DocxTemplate(str(template_path))
        except Exception as exc:
            return {
                "success": False,
                "error": str(exc),
                "message": f"Could not open template: {exc}",
            }

        try:
            declared: Set[str] = set(tpl.get_undeclared_template_variables())
        except Exception:
            declared = set()

        try:
            tpl.render(context)
            tpl.save(str(output_path))
        except Exception as exc:
            return {
                "success": False,
                "error": str(exc),
                "message": f"Jinja render failed: {exc}",
                "template_path": str(template_path),
                "output_path": str(output_path),
                "mode_used": "jinja",
            }

        provided = set(context.keys())
        filled = sorted(declared & provided)
        missing = sorted(declared - provided)
        unused = sorted(provided - declared)  # nested keys not introspected
        return {
            "success": True,
            "template_path": str(template_path),
            "output_path": str(output_path),
            "mode_used": "jinja",
            "filled_keys": filled,
            "missing_keys": missing,
            "unused_keys": unused,
            "message": f"Filled {len(filled)}/{max(len(declared), 1)} declared jinja variable(s).",
        }

    # ------------------------------------------------------------------ #
    # Bracket path (python-docx) + optional slot fill in same load        #
    # ------------------------------------------------------------------ #

    def _fill_brackets(
        self,
        template_path: Path,
        output_path: Path,
        context: Dict[str, Any],
        slot_values: Optional[Dict[str, Any]] = None,
    ) -> Dict[str, Any]:
        try:
            from docx import Document
        except ImportError as exc:
            return {
                "success": False,
                "error": "python-docx not installed",
                "message": f"python-docx required for bracket-mode templates: {exc}",
            }

        try:
            doc = Document(str(template_path))
        except Exception as exc:
            return {
                "success": False,
                "error": str(exc),
                "message": f"Could not open template: {exc}",
            }

        mapping: Dict[str, str] = {str(k).upper(): "" if v is None else str(v) for k, v in context.items()}
        filled_keys: Set[str] = set()
        discovered: Set[str] = set()
        total_subs = 0

        # Bracket pass first (scans full paragraphs including tables and headers/footers)
        for paragraph in _iter_all_paragraphs(doc):
            full = "".join((r.text or "") for r in paragraph.runs)
            if "[" in full:
                for m in _BRACKET_RE.finditer(full):
                    discovered.add(m.group(1))
                total_subs += _replace_brackets_in_paragraph(paragraph, mapping, filled_keys)

        # Slot pass on the same doc, if any
        slot_outcome: Dict[str, Any] = {}
        if slot_values:
            slots = _scan_slots(doc)
            slot_outcome = _apply_slot_values(slots, slot_values)
            if slot_outcome.get("rejected"):
                return {
                    "success": False,
                    "error": "legacy_slot_ids",
                    "message": (
                        "Refused to fill: caller used bare legacy slot ids "
                        "(e.g. 'slot_4') instead of the descriptive ids "
                        "returned by inspect_template. See "
                        "slot_fill.legacy_canonical_map for the required keys."
                    ),
                    "slot_fill": slot_outcome,
                }

        _strip_stranded_draft_marks(doc)

        try:
            doc.save(str(output_path))
        except Exception as exc:
            return {
                "success": False,
                "error": str(exc),
                "message": f"Could not save filled document: {exc}",
            }

        provided_keys = {str(k).upper() for k in context.keys()}
        missing = sorted(discovered - provided_keys)
        unused = sorted(provided_keys - discovered)
        mode_used = "bracket"
        if slot_values and discovered:
            mode_used = "bracket+slots"
        elif slot_values:
            mode_used = "slots"

        result: Dict[str, Any] = {
            "success": True,
            "template_path": str(template_path),
            "output_path": str(output_path),
            "mode_used": mode_used,
            "filled_keys": sorted(filled_keys),
            "missing_keys": missing,
            "unused_keys": unused,
            "substitutions": total_subs,
            "message": (
                f"Performed {total_subs} bracket substitution(s); "
                f"{len(missing)} bracket placeholder(s) left unfilled."
            ),
        }
        if slot_outcome:
            result["slot_fill"] = slot_outcome
            result["message"] += (
                f" Filled {len(slot_outcome.get('filled_slot_ids', []))}/"
                f"{len(slot_values)} slot(s)."
            )
        return result

    def _fill_slots_in_place(self, docx_path: Path, slot_values: Dict[str, Any]) -> Dict[str, Any]:
        """Apply slot_values to a docx already on disk (used after Jinja pass)."""
        try:
            from docx import Document
        except ImportError as exc:
            return {"success": False, "error": "python-docx not installed", "message": str(exc)}
        try:
            doc = Document(str(docx_path))
        except Exception as exc:
            return {"success": False, "error": str(exc), "message": f"Could not open {docx_path}: {exc}"}
        slots = _scan_slots(doc)
        outcome = _apply_slot_values(slots, slot_values)
        if outcome.get("rejected"):
            outcome["success"] = False
            outcome["error"] = "legacy_slot_ids"
            outcome["message"] = (
                "Refused to fill: caller used bare legacy slot ids "
                "instead of the descriptive ids returned by inspect_template. "
                "See legacy_canonical_map for the required keys."
            )
            return outcome
        _strip_stranded_draft_marks(doc)
        try:
            doc.save(str(docx_path))
        except Exception as exc:
            return {"success": False, "error": str(exc), "message": f"Could not save: {exc}"}
        outcome["success"] = True
        return outcome


# ---------------------------------------------------------------------- #
# Helpers — XML / mode detection                                          #
# ---------------------------------------------------------------------- #

def _read_docx_text_xml(docx_path: Path) -> List[str]:
    """Return text-bearing XML parts (document, headers, footers) as strings."""
    blobs: List[str] = []
    with zipfile.ZipFile(str(docx_path)) as zf:
        for name in zf.namelist():
            if not name.startswith("word/"):
                continue
            if name == "word/document.xml" or name.startswith(("word/header", "word/footer")):
                if name.endswith(".xml"):
                    blobs.append(zf.read(name).decode("utf-8", errors="replace"))
    return blobs


def _scan_jinja_vars(xml_text: str) -> Set[str]:
    """Top-level Jinja variable names appearing in {{ ... }}."""
    return {m.group(1) for m in _JINJA_VAR_RE.finditer(xml_text)}


def _scan_bracket_tokens(paragraph_texts: Iterable[str]) -> Set[str]:
    tokens: Set[str] = set()
    for t in paragraph_texts:
        if "[" not in t:
            continue
        for m in _BRACKET_RE.finditer(t):
            tokens.add(m.group(1))
    return tokens


def _detect_mode(
    jinja_vars: Set[str],
    bracket_tokens: Set[str],
    has_loops: bool,
    has_conditionals: bool,
) -> str:
    has_jinja = bool(jinja_vars) or has_loops or has_conditionals
    has_bracket = bool(bracket_tokens)
    if has_jinja and has_bracket:
        return "both"
    if has_jinja:
        return "jinja"
    if has_bracket:
        return "bracket"
    return "none"


# ---------------------------------------------------------------------- #
# Helpers — document walking                                              #
# ---------------------------------------------------------------------- #

def _iter_all_paragraphs(doc):
    """Yield every paragraph reachable in a python-docx Document."""
    for p in doc.paragraphs:
        yield p
    for tbl in doc.tables:
        yield from _walk_table_paragraphs(tbl)
    for section in doc.sections:
        for hf in (
            section.header,
            getattr(section, "first_page_header", None),
            getattr(section, "even_page_header", None),
            section.footer,
            getattr(section, "first_page_footer", None),
            getattr(section, "even_page_footer", None),
        ):
            if hf is None:
                continue
            for p in hf.paragraphs:
                yield p
            for tbl in hf.tables:
                yield from _walk_table_paragraphs(tbl)


def _walk_table_paragraphs(tbl):
    for row in tbl.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                yield p
            for nested in cell.tables:
                yield from _walk_table_paragraphs(nested)


# ---------------------------------------------------------------------- #
# Helpers — format-preserving run edits                                   #
# ---------------------------------------------------------------------- #

def _replace_paragraph_span(paragraph, start: int, end: int, replacement: str) -> bool:
    """Replace the character span [start, end) in a paragraph's run text.

    Only runs that overlap the span are modified — runs entirely outside the
    span keep their text (and therefore their per-run formatting) untouched.
    The first run in the span absorbs the replacement text and inherits its
    own formatting; runs strictly inside the span get blanked.
    """
    runs = paragraph.runs
    if not runs:
        return False
    texts = [r.text or "" for r in runs]
    offsets = [0]
    for t in texts:
        offsets.append(offsets[-1] + len(t))
    total = offsets[-1]
    if start < 0 or end > total or start >= end:
        return False

    def pos_to_run(p: int) -> Tuple[int, int]:
        for i in range(len(texts)):
            if offsets[i] <= p < offsets[i + 1]:
                return i, p - offsets[i]
        return len(texts) - 1, len(texts[-1])

    r_start, o_start = pos_to_run(start)
    if end == total:
        r_end = len(texts) - 1
        o_end = len(texts[r_end])
    else:
        r_end, o_end = pos_to_run(end)

    if r_start == r_end:
        texts[r_start] = texts[r_start][:o_start] + replacement + texts[r_end][o_end:]
    else:
        texts[r_start] = texts[r_start][:o_start] + replacement
        for i in range(r_start + 1, r_end):
            texts[i] = ""
        texts[r_end] = texts[r_end][o_end:]

    for i, r in enumerate(runs):
        if r.text != texts[i]:
            r.text = texts[i]
    return True


def _replace_brackets_in_paragraph(paragraph, mapping: Dict[str, str], filled_keys: Set[str]) -> int:
    """Substitute [TOKEN] occurrences in a paragraph, preserving run formatting.

    Applies matches in reverse order so earlier-match offsets stay valid even
    when replacement lengths differ from match lengths.
    """
    if not paragraph.runs:
        return 0
    full = "".join((r.text or "") for r in paragraph.runs)
    if "[" not in full:
        return 0
    matches = list(_BRACKET_RE.finditer(full))
    fillable = [(m.start(), m.end(), m.group(1)) for m in matches if m.group(1) in mapping]
    if not fillable:
        return 0
    count = 0
    for start, end, token in reversed(fillable):
        if _replace_paragraph_span(paragraph, start, end, mapping[token]):
            count += 1
            filled_keys.add(token)
    return count


# ---------------------------------------------------------------------- #
# Helpers — slot detection                                                #
# ---------------------------------------------------------------------- #

def _slot_public_view(slot: Dict[str, Any]) -> Dict[str, Any]:
    """Strip internal _meta references before returning to the caller.

    Surfaces a couple of meta fields that are useful to the agent: the full
    original highlighted text, a scaffold-split summary when applicable, and
    an `is_signature` flag for slots that should be left blank.
    """
    public = {k: v for k, v in slot.items() if not k.startswith("_")}
    meta = slot.get("_meta", {})
    if meta.get("span_text"):
        public["span_text"] = meta["span_text"]
    scaffold = meta.get("scaffold")
    if scaffold:
        public["scaffold"] = {
            "kind": scaffold.get("kind"),
            "prefix": scaffold.get("pre", ""),
            "variable": scaffold.get("var", ""),
            "suffix": scaffold.get("suf", ""),
        }
    if meta.get("is_signature"):
        public["is_signature"] = True
        public["fill_policy"] = "leave_blank"
    if meta.get("original_label"):
        # Surfaces the raw highlighted text when the label was rewritten by
        # _enrich_meaningless_slot_labels (e.g. "必填" → "E-mail或传真").
        # The agent can use this to confirm what was originally in the cell.
        public["original_label"] = meta["original_label"]
    # For span-replacement slot kinds, surface the EXACT characters that get
    # replaced. Without this, the agent has only `context` (a snippet with
    # surrounding prose) to look at, and frequently includes adjacent
    # template prose in the slot_value — e.g. the slot at "...预付合同总额的
    # _____（其中合同总额的20%作为定金），..." should be filled with just
    # "90%", but agents have passed "90%（其中合同总额的20%作为定金）",
    # leaving the original parenthetical in place and producing
    # "90%（其中…20%作为定金）（其中…20%作为定金）". The `replaces` field
    # makes the fill target unambiguous.
    if slot.get("kind") == "underscores":
        para = meta.get("paragraph")
        start, end = meta.get("start"), meta.get("end")
        if para is not None and start is not None and end is not None:
            try:
                full = "".join((r.text or "") for r in para.runs)
                public["replaces"] = full[start:end]
            except Exception:
                pass
    if slot.get("kind") == "option_choice" and meta.get("options"):
        public["options"] = [
            {"index": opt["index"], "header": opt["header"], "preview": opt["preview"]}
            for opt in meta["options"]
        ]
        public["fill_policy"] = (
            "Pass the option index (1.." + str(len(meta["options"])) +
            ") or its Chinese label ('第一种', '第二种', ...) as slot_values[slot_id]. "
            "On fill, the chosen option's body is kept; the prompt, each option's "
            "header line, and all other options' bodies are removed."
        )
    return public


def _detect_option_choice_groups(body_paras) -> List[Dict[str, Any]]:
    """Scan body paragraphs for the 二选一/三选一 alternative-options pattern.

    Returns a list of group dicts. Each group has:
      - prompt_paragraph: the instruction paragraph ("（以下两种…请删除）")
      - options: list of {index, header (text), header_paragraph,
        body_paragraphs (list), preview (str)}

    Detection requires both a prompt paragraph AND at least 2 option headers
    (`第N种`) within the next K paragraphs. v1 only scans body paragraphs.
    """
    groups: List[Dict[str, Any]] = []
    K = 40
    i = 0
    while i < len(body_paras):
        p = body_paras[i]
        text = "".join((r.text or "") for r in p.runs).strip()
        if not text or not _OPTION_PROMPT_RE.search(text):
            i += 1
            continue
        # Found a prompt. Look ahead for option headers within K paragraphs.
        end_idx = min(i + 1 + K, len(body_paras))
        header_indices: List[Tuple[int, int]] = []  # (paragraph_index, option_index)
        for j in range(i + 1, end_idx):
            jt = "".join((r.text or "") for r in body_paras[j].runs).strip()
            if not jt:
                continue
            mh = _OPTION_HEADER_RE.match(jt)
            if mh:
                opt_idx_raw = mh.group(1)
                opt_idx = _CN_NUM_TO_INT.get(opt_idx_raw)
                if opt_idx is not None:
                    header_indices.append((j, opt_idx))
                continue
            # A section break before we found 2 headers aborts the group.
            if not header_indices and (
                _SECTION_BREAK_RE.match(jt)
                or _is_heading(body_paras[j])
                or _paragraph_has_list_numbering(body_paras[j])
            ):
                break
        if len(header_indices) < 2:
            i += 1
            continue
        # Walk options and collect each one's body paragraphs.
        options: List[Dict[str, Any]] = []
        for k, (hidx, opt_index) in enumerate(header_indices):
            header_text = "".join((r.text or "") for r in body_paras[hidx].runs).strip()
            # body: paragraphs after this header until next header / section
            # break / K limit / doc end
            body_paragraphs = []
            body_text_parts = []
            if k + 1 < len(header_indices):
                next_header_idx = header_indices[k + 1][0]
            else:
                next_header_idx = end_idx
            for m in range(hidx + 1, next_header_idx):
                mt = "".join((r.text or "") for r in body_paras[m].runs).strip()
                if (mt and _SECTION_BREAK_RE.match(mt)) or _is_heading(body_paras[m]) \
                        or _paragraph_has_list_numbering(body_paras[m]):
                    break
                body_paragraphs.append(body_paras[m])
                if mt:
                    body_text_parts.append(mt)
            preview = " / ".join(body_text_parts)[:80]
            options.append({
                "index": opt_index,
                "header": header_text,
                "header_paragraph": body_paras[hidx],
                "body_paragraphs": body_paragraphs,
                "preview": preview,
            })
        groups.append({
            "prompt_paragraph": p,
            "prompt_text": text,
            "options": options,
        })
        # Skip past the last option's body to avoid re-detecting nested patterns.
        last_body = options[-1]["body_paragraphs"]
        if last_body:
            last_idx = body_paras.index(last_body[-1])
            i = last_idx + 1
        else:
            i = header_indices[-1][0] + 1
    return groups


def _paragraph_element_id(paragraph) -> Optional[int]:
    """Identity that's stable across python-docx wrapper recreation.

    `doc.paragraphs` builds fresh Paragraph wrappers on each call, so
    `id(paragraph)` is not stable between scans. The underlying lxml
    `<w:p>` element IS stable for the lifetime of the document.
    """
    if paragraph is None:
        return None
    try:
        return id(paragraph._element)
    except Exception:
        return None


def _collect_option_choice_paragraph_ids(slots: List[Dict[str, Any]]) -> Set[int]:
    """Set of stable element-ids for every paragraph owned by an option_choice
    slot (prompt + every option's header + body). Used to suppress duplicate
    removal candidates whose drop is already handled by the slot fill."""
    consumed: Set[int] = set()
    for s in slots:
        if s.get("kind") != "option_choice":
            continue
        meta = s.get("_meta", {}) or {}
        eid = _paragraph_element_id(meta.get("prompt_paragraph"))
        if eid is not None:
            consumed.add(eid)
        for opt in meta.get("options") or []:
            eid = _paragraph_element_id(opt.get("header_paragraph"))
            if eid is not None:
                consumed.add(eid)
            for bp in opt.get("body_paragraphs") or []:
                eid = _paragraph_element_id(bp)
                if eid is not None:
                    consumed.add(eid)
    return consumed


# Leading list/section markers that should NOT appear in slot ids — they're
# repetitive ("1．", "（1）", "2．", "3．"...) and would make ids ambiguous.
_SLOT_ID_LEAD_MARKER_RE = re.compile(
    r"^[\s 　]*"
    r"(?:第[一二三四五六七八九十百零〇\d]+[条章节款项]|"
    r"[（(]\s*\d+\s*[）)]|"
    r"\d+[\.\．、:：])"
    r"[\s 　]*"
)

# Characters that should be stripped from a label before embedding it in a
# slot id — punctuation, brackets, whitespace, colons. CJK letters, ASCII
# letters, ASCII digits, and underscores are kept.
_SLOT_ID_DROP_RE = re.compile(
    r"[^A-Za-z0-9_一-鿿㐀-䶿]+"
)


def _build_canonical_id_map(
    original_path: str, partial_path: str
) -> Optional[Dict[str, str]]:
    """Map canonical slot ids (as returned by inspect_template on the ORIGINAL
    template) to their current ids in a partial-fill output document.

    Returns ``None`` if either file fails to scan. Returns a possibly-empty
    dict otherwise; entries for partial-doc slots that have no clear match
    are simply absent (caller falls back to the partial-doc id verbatim).

    Matching strategy: both docs are scanned in document order, then partial
    slots are paired with original slots by walking both lists and skipping
    over original slots that no longer appear in the partial (those are the
    already-filled ones). Pairing keys on (kind, label) — stable across
    fills because labels come from the surrounding prose, not the blank
    region itself.
    """
    try:
        from docx import Document
    except ImportError:
        return None
    try:
        orig_doc = Document(original_path)
        part_doc = Document(partial_path)
    except Exception:
        return None
    try:
        orig_slots = _scan_slots(orig_doc)
        part_slots = _scan_slots(part_doc)
    except Exception:
        return None
    mapping: Dict[str, str] = {}
    i = 0  # walker into orig_slots
    for ps in part_slots:
        pk = (ps["kind"], ps.get("label") or "")
        # Walk forward in originals until we find a (kind, label) match.
        # Original slots passed over are presumed already filled.
        j = i
        while j < len(orig_slots):
            os_ = orig_slots[j]
            ok = (os_["kind"], os_.get("label") or "")
            if ok == pk:
                break
            j += 1
        if j < len(orig_slots):
            mapping[orig_slots[j]["id"]] = ps["id"]
            i = j + 1
    # Also surface the inverse mapping for identity slots (original id == new
    # id) — these are slots whose counter didn't shift. Caller looks up the
    # canonical id directly; identity entries make the lookup uniform.
    part_ids = {ps["id"] for ps in part_slots}
    for os_ in orig_slots:
        if os_["id"] in part_ids and os_["id"] not in mapping:
            mapping[os_["id"]] = os_["id"]
    return mapping


def _make_slot_id(counter: int, label: Optional[str]) -> str:
    """Produce a stable, human-readable slot id like ``slot_023_签订时间``.

    The numeric prefix (zero-padded to 3 digits for sort stability) is the
    canonical disambiguator — labels in long Chinese form templates repeat
    heavily (1．/2．/3．). Including a sanitized label tail in the id makes
    it far easier for an LLM caller to map values to the correct slot, and
    the numeric prefix guarantees uniqueness even when labels collide.
    """
    short = ""
    if label:
        stripped = _SLOT_ID_LEAD_MARKER_RE.sub("", label).strip()
        cleaned = _SLOT_ID_DROP_RE.sub("", stripped)
        # Cap to ~10 CJK chars / 20 ASCII chars — enough to disambiguate
        # without bloating the id.
        if len(cleaned) > 20:
            cleaned = cleaned[:20]
        short = cleaned
    if short:
        return f"slot_{counter:03d}_{short}"
    return f"slot_{counter:03d}"


def _scan_slots(doc) -> List[Dict[str, Any]]:
    """Walk a python-docx Document and produce slot candidates with live refs.

    Slot kinds:
      - "underscores"  : a run of 3+ underscores; replaces in place
      - "label_blank"  : paragraph ending in "Label:"; value appended after
      - "empty_cell"   : table cell with no text, where row/column has a label
    """
    slots: List[Dict[str, Any]] = []
    counter = [0]

    def add(kind: str, label: Optional[str], context: str, meta: Dict[str, Any]) -> None:
        slot_id = _make_slot_id(counter[0], label)
        counter[0] += 1
        slots.append({
            "id": slot_id,
            "kind": kind,
            "label": label,
            "context": context,
            "_meta": meta,
        })

    # Body paragraphs (with neighbor lookahead for section_body_empty)
    body_paras = list(doc.paragraphs)
    seen_para_ids: Set[int] = set()
    # First pass: detect 二选一 / 三选一 option-choice groups and mark every
    # involved paragraph as consumed so the per-paragraph heuristics below
    # don't re-emit them as separate slots/removals. The composite slot
    # owns the drop of all unchosen-option paragraphs at fill time.
    option_groups = _detect_option_choice_groups(body_paras)
    for grp in option_groups:
        # Best-effort label: the nearest preceding non-empty heading-like paragraph.
        prompt_p = grp["prompt_paragraph"]
        prompt_idx = body_paras.index(prompt_p)
        label = None
        for back in range(prompt_idx - 1, max(-1, prompt_idx - 8), -1):
            bt = "".join((r.text or "") for r in body_paras[back].runs).strip()
            if bt and (_is_heading(body_paras[back]) or _SECTION_BREAK_RE.match(bt)):
                label = bt[:60]
                break
        if not label:
            label = f"{len(grp['options'])}选1"
        add(
            "option_choice",
            label,
            _snippet(grp["prompt_text"], 0, len(grp["prompt_text"]), radius=80),
            {
                "prompt_paragraph": prompt_p,
                "options": grp["options"],
            },
        )
        seen_para_ids.add(id(prompt_p))
        for opt in grp["options"]:
            seen_para_ids.add(id(opt["header_paragraph"]))
            for bp in opt["body_paragraphs"]:
                seen_para_ids.add(id(bp))
    for i, p in enumerate(body_paras):
        if id(p) in seen_para_ids:
            continue
        before = len(slots)
        _scan_paragraph_for_slots(p, add)
        # Suppress label_blank emitted for a top-level section heading whose
        # body lives in the following paragraphs (e.g. "四、交货方式及时间、
        # 地点：" followed by "交货方式：...", "交货时间：...", "交货地点：
        # ..."). Without this, a value gets appended after the heading's colon
        # AND the structured sub-fields below remain — duplicating content.
        # The heading marker may be invisible in run text — CN ministry
        # templates often render "四、" via Word's auto-numbering (w:numPr),
        # so we also accept paragraphs that carry numbering / Heading style.
        if len(slots) > before:
            new_slots = slots[before:]
            if any(s["kind"] == "label_blank" for s in new_slots):
                p_text = "".join((r.text or "") for r in p.runs).strip()
                if p_text.endswith(("：", ":")):
                    follow_p = None
                    follow_text = ""
                    for j in range(i + 1, min(i + 6, len(body_paras))):
                        nxt_text = "".join(
                            (r.text or "") for r in body_paras[j].runs
                        ).strip()
                        if nxt_text:
                            follow_p = body_paras[j]
                            follow_text = nxt_text
                            break
                    if follow_text and not _SECTION_BREAK_RE.match(follow_text):
                        p_ilvl = _paragraph_list_level(p)
                        n_ilvl = (
                            _paragraph_list_level(follow_p)
                            if follow_p is not None else None
                        )
                        # Heading-above-body via numbering: heading carries
                        # numbering; the next paragraph either has no
                        # numbering at all (typical for prose bodies under a
                        # numbered section) or sits at a strictly deeper ilvl
                        # (typical for ordered sub-clauses under a top-level
                        # 违约责任 / 售后服务 heading). Same-level sibling
                        # numbering ("1. 项目名称：" / "2. 项目编号：xxx") is
                        # excluded — those are peers, not parent/child.
                        heading_above_body = (
                            p_ilvl is not None
                            and (n_ilvl is None or n_ilvl > p_ilvl)
                        )
                        # Indicators that p is a section heading (not a peer
                        # of the lines below): visible section marker, Heading
                        # style, OR numbering with body at deeper / no level.
                        is_container = (
                            bool(_SECTION_BREAK_RE.match(p_text))
                            or _is_heading(p)
                            or heading_above_body
                        )
                        if is_container:
                            slots[:] = slots[:before] + [
                                s for s in new_slots if s["kind"] != "label_blank"
                            ]
        if len(slots) > before:
            seen_para_ids.add(id(p))
    # section_body_empty: heading followed by empty body paragraph
    for i, p in enumerate(body_paras[:-1]):
        if not _is_section_heading_like(p):
            continue
        nxt = body_paras[i + 1]
        if id(nxt) in seen_para_ids:
            continue
        if not _paragraph_is_empty(nxt):
            continue
        heading_text = "".join((r.text or "") for r in p.runs).strip()
        if not heading_text:
            continue
        # Suppress: heading is a list/table introduction (以下/如下/following...).
        # The real fill target is the table or list below, not this empty
        # paragraph — emitting a slot here would dump body text where the
        # user wants a new row in the table that follows.
        if _LIST_INTRO_RE.search(heading_text):
            continue
        # Suppress: a table immediately follows (with at most one empty
        # paragraph between). Same reasoning — the table is the right target.
        if _table_follows_paragraph(p, nxt):
            continue
        add("section_body_empty", heading_text, f"(empty body under heading '{heading_text}')", {
            "paragraph": nxt,
            "start": 0,
            "end": 0,
        })
        seen_para_ids.add(id(nxt))
    # Border-line slots: a paragraph with a horizontal border (bottom border)
    # renders as a fillable "underline" line even though it has no underscore
    # characters. Two layouts to support:
    #   (a) Label paragraph immediately followed by an empty bordered paragraph
    #       ("Patient Name:" / [empty + bottom-border]). The label paragraph
    #       was already emitted as label_blank; we redirect its fill into the
    #       bordered paragraph below so the value sits over the line instead
    #       of being appended after the colon.
    #   (b) A standalone empty bordered paragraph with no leading label — emit
    #       a stray slot so the agent can confirm intent.
    # Border on the label paragraph itself ("Patient Name:    " all in one
    # bordered paragraph) is handled in (c): mark the label_blank slot so its
    # fill is centered into the same paragraph rather than appended after.
    slots_by_paragraph_id: Dict[int, Dict[str, Any]] = {}
    for s in slots:
        meta = s.get("_meta", {})
        para = meta.get("paragraph")
        if para is not None and s["kind"] == "label_blank":
            slots_by_paragraph_id[id(para)] = s
    for i, p in enumerate(body_paras):
        if not _paragraph_has_horizontal_border(p):
            continue
        if _paragraph_is_empty(p):
            # (a) link to a preceding label_blank slot if its paragraph is the
            # previous one in body order
            linked = False
            if i > 0:
                prev = body_paras[i - 1]
                prev_slot = slots_by_paragraph_id.get(id(prev))
                if prev_slot is not None:
                    prev_slot["_meta"]["border_target"] = p
                    seen_para_ids.add(id(p))
                    linked = True
            if linked or id(p) in seen_para_ids:
                continue
            # (b) standalone empty bordered paragraph — emit a stray slot.
            # Use the previous non-empty paragraph (if any) as a label hint.
            label_hint: Optional[str] = None
            for j in range(i - 1, -1, -1):
                t = "".join((r.text or "") for r in body_paras[j].runs).strip()
                if t:
                    label_hint = t[:60]
                    break
            add("underscores", label_hint, "(bordered empty line)", {
                "paragraph": p,
                "start": 0,
                "end": 0,
                "is_signature": _looks_like_signature(label_hint, label_hint or "", 0, 0),
                "border_only": True,
            })
            seen_para_ids.add(id(p))
        else:
            # (c) bordered paragraph that already has its own content. If we
            # emitted a label_blank slot for it, mark the slot so the fill is
            # centered in the same paragraph instead of bumping right after
            # the label.
            existing = slots_by_paragraph_id.get(id(p))
            if existing is not None:
                existing["_meta"]["bordered_paragraph"] = True
    # Body tables
    for tbl in doc.tables:
        _scan_table_for_slots(tbl, add)
    # After all body-table scanning, enrich meaningless labels ("必填",
    # "待填", "请填", etc.) using the slot's table-row context. Without this,
    # the agent sees N highlighted "必填" slots in a 甲方/乙方 contact table
    # and has no idea which one is the email vs. address vs. phone — the
    # template author's intent ("required field, fill in") is preserved in
    # the row's label cell (e.g. "E-mail或传真"), not in the highlighted run.
    _enrich_meaningless_slot_labels(slots, doc)
    # Headers / footers
    for section in doc.sections:
        for hf in (
            section.header,
            getattr(section, "first_page_header", None),
            getattr(section, "even_page_header", None),
            section.footer,
            getattr(section, "first_page_footer", None),
            getattr(section, "even_page_footer", None),
        ):
            if hf is None:
                continue
            for p in hf.paragraphs:
                _scan_paragraph_for_slots(p, add)
            for tbl in hf.tables:
                _scan_table_for_slots(tbl, add)
    return slots


def _scan_paragraph_for_slots(paragraph, add: Callable[..., None]) -> None:
    text = "".join((r.text or "") for r in paragraph.runs)
    if not text:
        return
    # 0) Highlighted run spans — strongest authoring signal: "modify me".
    #    Consecutive highlighted runs merge into one slot. Highlight will be
    #    cleared on fill so the final document looks clean. If the span looks
    #    like "<variable><scaffold>" (15个工作日, ¥850, 50%, etc.), we record
    #    the scaffold split so a bare-number user reply still produces a
    #    correctly-scaffolded output.
    #
    #    Highlighted spans can coexist with underlined-whitespace and inline-
    #    whitespace blanks in the same paragraph (e.g. money lines with the
    #    drafting note highlighted at the end). We emit all of them and only
    #    skip downstream paragraph-level detection (underscores / angle /
    #    label-only / phrase bank) so we don't double-count.
    covered: List[Tuple[int, int]] = []

    def _overlaps_covered(start: int, end: int) -> bool:
        return any(not (end <= s or start >= e) for s, e in covered)

    highlighted_spans = _find_highlighted_spans(paragraph)
    for start, end, runs, span_text in highlighted_spans:
        # Suppress slot emission for highlighted drafting notes — spans whose
        # text is itself a deletion instruction ("空着，后盖章的一方手写即可",
        # "（请删除）", "（此句话非正文）"). _scan_removals picks these up via
        # the run-level instruction_run path (highlighted runs now count as
        # hintlike) and fill_template auto-applies the resulting removals, so
        # the drafting note gets cleared without the agent having to choose a
        # fill value. We still mark the span as 'covered' so downstream
        # underscore/inline-blank/label-only passes don't double-emit slots
        # for the same character range.
        if _RE_REMOVAL_BANK.search(span_text or ""):
            covered.append((start, end))
            continue
        label = _guess_label_before(text, start) or (span_text.strip() or None)
        ctx = _snippet(text, start, end)
        scaffold = _split_scaffold(span_text)
        add("highlighted", label, ctx, {
            "paragraph": paragraph,
            "start": start,
            "end": end,
            "runs": runs,
            "scaffold": scaffold,        # None if span isn't a known pattern
            "span_text": span_text,      # full original highlighted text
        })
        covered.append((start, end))
    # 0.5) Underlined whitespace spans — Word's "fill-in line" trick where the
    #      author selects a run of spaces and applies underline. There are no
    #      underscore characters in the document, but it looks like one to a
    #      human. Treat as an underscores-kind slot; centering pads with spaces
    #      (the inherited underline carries through so the line stays visible).
    raw_uw_spans = _find_underlined_whitespace_spans(paragraph)
    uw_spans = _expand_and_merge_uw_spans(text, raw_uw_spans)
    uw_emitted = False
    for start, end, runs, span_text in uw_spans:
        if _overlaps_covered(start, end):
            continue
        label = _guess_label_before(text, start)
        ctx = _snippet(text, start, end, radius=40)
        is_sig = _looks_like_signature(label, text, start, end)
        add("underscores", label, ctx, {
            "paragraph": paragraph,
            "start": start,
            "end": end,
            "is_signature": is_sig,
            "pad_char": " ",
            "source": "underlined_whitespace",
        })
        covered.append((start, end))
        uw_emitted = True
    # 0.6) Date scaffold — `<ws>年<ws>月[<ws>日]` left as fill space in CN
    #      forms. Detected before inline_blank so the WHOLE scaffold span
    #      (gaps + 年/月/日 markers) becomes one slot. Otherwise we'd only
    #      catch the leading gap and the agent's `2026年5月14日` value
    #      would land next to the leftover `年   月   日` template text.
    date_emitted = False
    for m in _DATE_SCAFFOLD_RE.finditer(text):
        ds, de = m.start(), m.end()
        if _overlaps_covered(ds, de):
            continue
        label = _guess_label_before(text, ds) or "日期"
        ctx = _snippet(text, ds, de, radius=40)
        is_sig = _looks_like_signature(label, text, ds, de)
        add("underscores", label, ctx, {
            "paragraph": paragraph,
            "start": ds,
            "end": de,
            "is_signature": is_sig,
            "pad_char": " ",
            "source": "date_scaffold",
        })
        covered.append((ds, de))
        date_emitted = True
    # 0.7) Inline whitespace blanks — no underline, no highlight, just spaces
    #      between a label colon / currency symbol and sentence punctuation
    #      or a unit marker. Common in CN government / contract templates.
    inline_emitted = False
    for m in _INLINE_BLANK_RE.finditer(text):
        gs, ge = m.start("gap"), m.end("gap")
        if _overlaps_covered(gs, ge):
            continue
        lm_text = m.group("lm")
        # When the left marker is a 大写/小写 amount label, override the
        # heuristic label with the clean marker text so the 大写/小写
        # reconciler can pair them by label alone. Without this, both
        # slots in a "...：（大写）___（小写）___" paragraph would carry
        # context strings containing "大写" and get misclassified as a pair
        # of 大写 slots.
        if _DAXIE_MARKER_RE.match(lm_text):
            label = "大写"
        elif _XIAOXIE_MARKER_RE.match(lm_text):
            label = "小写"
        else:
            label = _guess_label_before(text, gs)
        ctx = _snippet(text, gs, ge, radius=40)
        is_sig = _looks_like_signature(label, text, gs, ge)
        add("underscores", label, ctx, {
            "paragraph": paragraph,
            "start": gs,
            "end": ge,
            "is_signature": is_sig,
            "pad_char": " ",
            "source": "inline_whitespace_blank",
        })
        covered.append((gs, ge))
        inline_emitted = True
    if highlighted_spans or uw_emitted or date_emitted or inline_emitted:
        return
    # 1) underscore runs (may have multiple per paragraph)
    underscore_matches = list(_UNDERSCORE_RUN_RE.finditer(text))
    if underscore_matches:
        # Skip a paragraph that's *only* a long decorative underscore line
        # (e.g. a page divider) with no surrounding label or signature context.
        if _is_decorative_underscore_paragraph(text, underscore_matches):
            return
        # Group adjacent underscore runs that are separated only by decorator
        # characters (whitespace, '/', '\\', '.', '-', ...). A group of >1
        # entries is a "fill region" — '建筑面积：____/____平方米' or
        # '¥ ____ ____ ____'. Gaps containing Chinese / letters / digits
        # always break the group so labelled siblings stay separate.
        underscore_groups = _group_decorative_underscore_matches(
            text, underscore_matches
        )
        for group in underscore_groups:
            if len(group) == 1:
                m = group[0]
                # Absorb a trailing draft-mark + whitespace tail
                # ('____/         平方米' or '号/' at end of paragraph) into the
                # slot so the fill replaces them too.
                end_extended = _extend_span_right_past_draft_marks(text, m.end())
                label = _guess_label_before(text, m.start())
                context = _snippet(text, m.start(), end_extended)
                is_sig = _looks_like_signature(label, text, m.start(), end_extended)
                add("underscores", label, context, {
                    "paragraph": paragraph,
                    "start": m.start(),
                    "end": end_extended,
                    "is_signature": is_sig,
                })
                continue
            positions = [(m.start(), m.end()) for m in group]
            span_start = positions[0][0]
            span_end = positions[-1][1]
            label = _guess_label_before(text, span_start)
            ctx = _snippet(text, span_start, span_end, radius=60)
            is_sig = _looks_like_signature(label, text, span_start, span_end)
            # Inter-gap contains non-whitespace decorator chars ('/', '\\',
            # '.', '-')? Then those are draft marks inside one fill region:
            # emit ONE span covering everything so the fill absorbs them.
            # Pure-whitespace gaps stay composite (digit-cell pattern).
            has_decorator_gap = any(
                text[positions[i][1]:positions[i + 1][0]].strip()
                for i in range(len(positions) - 1)
            )
            if has_decorator_gap:
                span_end_extended = _extend_span_right_past_draft_marks(
                    text, span_end
                )
                add("underscores", label, ctx, {
                    "paragraph": paragraph,
                    "start": span_start,
                    "end": span_end_extended,
                    "is_signature": is_sig,
                    "pad_char": " ",
                })
            else:
                add("underscores", label, ctx, {
                    "paragraph": paragraph,
                    "start": span_start,
                    "end": span_end,
                    "positions": positions,   # multi-position composite
                    "composite": True,
                    "is_signature": is_sig,
                })
        return  # don't double-count this paragraph as label_blank
    # 2) angle-bracketed tokens like <your name> or 《姓名》
    angle_matches = list(_ANGLE_BRACKET_RE.finditer(text))
    if angle_matches:
        for m in angle_matches:
            inner = (m.group(1) or m.group(2) or "").strip()
            if not inner:
                continue
            ctx = _snippet(text, m.start(), m.end())
            add("angle_bracketed", inner, ctx, {
                "paragraph": paragraph,
                "start": m.start(),
                "end": m.end(),
            })
        return
    # 3) "Label:" at end of paragraph with no value after
    m = _LABEL_ONLY_RE.match(text)
    if m:
        # Skip when this paragraph is introducing a list/table — the content
        # belongs in the table or list that follows, not appended after the
        # heading's trailing colon.
        if _LIST_INTRO_RE.search(text):
            return
        label = m.group(1).strip()
        is_sig = _looks_like_signature(label, text, 0, len(text))
        add("label_blank", label, text, {"paragraph": paragraph, "is_signature": is_sig})
        return
    # 4) hint-text run (italic / grey) matching a placeholder-like phrase
    hint = _find_hint_run(paragraph)
    if hint is not None:
        run, run_start, run_end, snippet_text = hint
        label = _guess_label_before(text, run_start) or snippet_text.strip()
        ctx = _snippet(text, run_start, run_end)
        add("hint_text", label, ctx, {
            "paragraph": paragraph,
            "start": run_start,
            "end": run_end,
        })
        return
    # 5) whole-paragraph placeholder phrase (no italic required) — e.g.
    #    "Replace this with your bio." or "请填写姓名".
    pm = _RE_PHRASE_BANK.search(text)
    if pm:
        # If the matched phrase covers most of the paragraph, treat as a slot
        # that replaces the whole paragraph; otherwise just replace the span.
        cover = (pm.end() - pm.start()) / max(len(text.strip()), 1)
        if cover >= 0.6:
            label = _guess_label_before(text, pm.start()) or pm.group(0).strip()
            add("placeholder_phrase", label, _snippet(text, pm.start(), pm.end()), {
                "paragraph": paragraph,
                "start": 0,
                "end": len(text),
            })
        else:
            label = _guess_label_before(text, pm.start()) or pm.group(0).strip()
            add("placeholder_phrase", label, _snippet(text, pm.start(), pm.end()), {
                "paragraph": paragraph,
                "start": pm.start(),
                "end": pm.end(),
            })


def _cell_tc_element(cell):
    """Return the underlying <w:tc> XML element for a Cell wrapper.

    Note: python-docx returns the **anchor** tc on every continuation row of
    a vertically-merged region (so `cell._tc` is the same object across the
    merge). Callers that need to deduplicate per-merge should use the tc
    identity returned here as the dedupe key.
    """
    try:
        return cell._tc
    except Exception:
        return None


def _is_placeholder_cell(cell) -> bool:
    """True if the cell's content is a seal/date template placeholder
    (signature block waiting to be filled)."""
    text = _cell_text(cell)
    if not text.strip():
        return False
    if _SEAL_PLACEHOLDER_RE.search(text):
        return True
    if _DATE_PLACEHOLDER_RE.search(text):
        return True
    return False


def _scan_table_for_slots(tbl, add: Callable[..., None]) -> None:
    # Best-guess column headers from row 0 (text-only)
    header_labels: List[str] = []
    if tbl.rows:
        for cell in tbl.rows[0].cells:
            header_labels.append(_cell_text(cell).strip())

    # Track columns that already have a placeholder_cell slot. Empty cells
    # below such a column are usually layout phantoms (a non-merged empty
    # row directly after a vertically-merged anchor) and should not surface
    # as separate fillable slots. Also dedupe by the underlying <w:tc> XML
    # element so vertically-merged cells (which python-docx returns on every
    # spanned row as the same anchor tc) only get scanned once.
    #
    # The set stores the lxml elements themselves rather than `id(tc)` —
    # `id()` returns CPython memory addresses that can be recycled across
    # gc'd proxy objects, so different tcs in the same scan were colliding
    # and dropping legitimate slots. Storing the element references keeps
    # them alive (so they hash stably by node identity) and equality follows
    # underlying XML node identity.
    placeholder_columns: Set[int] = set()
    seen_cell_tcs: set = set()

    for ri, row in enumerate(tbl.rows):
        first_cell_text = _cell_text(row.cells[0]).strip() if row.cells else ""
        for ci, cell in enumerate(row.cells):
            # Nested tables first so order is stable
            for nested in cell.tables:
                _scan_table_for_slots(nested, add)
            tc = _cell_tc_element(cell)
            already_emitted_for_tc = tc is not None and tc in seen_cell_tcs
            # Placeholder cell (seal / stamp / date area). Emit once per merge
            # anchor by deduping on the underlying <w:tc> identity.
            if _is_placeholder_cell(cell) and not already_emitted_for_tc:
                row_label = first_cell_text if ci > 0 else ""
                col_label = header_labels[ci] if ci < len(header_labels) else ""
                cell_text = _cell_text(cell).strip()
                label = cell_text[:60] if cell_text else (col_label or row_label or "印章/日期")
                ctx_bits = []
                if row_label:
                    ctx_bits.append(f"row='{row_label}'")
                if col_label and col_label != cell_text:
                    ctx_bits.append(f"col='{col_label}'")
                context = (
                    "[seal/date cell" + (", " + ", ".join(ctx_bits) if ctx_bits else "")
                    + f": {cell_text[:60]}]"
                )
                add("placeholder_cell", label, context, {"cell": cell})
                placeholder_columns.add(ci)
                if tc is not None:
                    seen_cell_tcs.add(tc)
                continue
            if already_emitted_for_tc:
                # Continuation row of a cell we've already emitted (placeholder
                # or otherwise) — don't re-scan its paragraphs.
                continue
            if tc is not None:
                seen_cell_tcs.add(tc)
            for p in cell.paragraphs:
                _scan_paragraph_for_slots(p, add)
            # Empty cell heuristic: no text and no nested tables
            if not _cell_text(cell).strip() and not cell.tables:
                # Phantom cell under a column whose anchor was already emitted
                # as a placeholder slot — skip.
                if ci in placeholder_columns:
                    continue
                column_label = header_labels[ci] if ci < len(header_labels) else ""
                row_label = first_cell_text if ci > 0 else ""
                label = column_label or row_label
                if not label:
                    continue
                if ri == 0 and column_label == "":
                    # top-left empty cell with no label — skip
                    continue
                ctx_bits = []
                if row_label:
                    ctx_bits.append(f"row='{row_label}'")
                if column_label and ci != 0:
                    ctx_bits.append(f"col='{column_label}'")
                context = "[empty cell, " + ", ".join(ctx_bits) + "]" if ctx_bits else "[empty cell]"
                add("empty_cell", label, context, {"cell": cell})


def _cell_text(cell) -> str:
    return "\n".join((p.text or "") for p in cell.paragraphs)


# Labels that mean "fill in something here" but carry no field semantics — when
# a highlighted slot's label is just one of these, the agent has no idea what
# to put in. Used by _enrich_meaningless_slot_labels.
_MEANINGLESS_FILL_LABELS = frozenset({
    "必填", "必填项", "必填字段", "必须填写", "请填", "请填写", "请输入",
    "待填", "待填写", "待补充", "待完善", "待确认",
    "填写", "填入", "填空",
    "tbd", "to be filled", "to be completed", "fill in", "required",
    "n/a",
})


def _is_meaningless_fill_label(label: Optional[str]) -> bool:
    if not label:
        return False
    return label.strip().lower() in _MEANINGLESS_FILL_LABELS


def _paragraph_parent_cell_tc(paragraph):
    """Walk up paragraph._element ancestors to find the enclosing <w:tc>.
    Returns the lxml element or None if the paragraph isn't inside a cell."""
    try:
        el = paragraph._element
    except AttributeError:
        return None
    cur = el.getparent()
    while cur is not None:
        tag = getattr(cur, "tag", "")
        if isinstance(tag, str) and tag.endswith("}tc"):
            return cur
        cur = cur.getparent()
    return None


def _row_label_for_tc(target_tc, doc) -> Optional[str]:
    """Find the most informative label cell in the same row as `target_tc`.

    Walks every table (including nested) until it finds the row containing
    `target_tc`, then returns the text of the first cell to the left that is
    short enough to be a label, non-empty, distinct from the slot cell, and
    not a vMerge continuation. Skips merge-anchor party indicators ("甲方",
    "乙方") which carry no field semantics. Returns None when no usable
    label cell is in row order.
    """
    def _walk_tables(tables):
        for tbl in tables:
            for row in tbl.rows:
                cells = list(row.cells)
                target_idx = None
                for i, c in enumerate(cells):
                    if c._tc is target_tc:
                        target_idx = i
                        break
                if target_idx is None:
                    # Recurse into nested tables.
                    for c in cells:
                        for nested in c.tables:
                            hit = _walk_tables([nested])
                            if hit is not None:
                                return hit
                    continue
                # Found the row — scan cells to the left of the target for
                # the best label candidate (closest non-empty, non-anchor cell).
                seen_tcs: set = set()
                for j in range(target_idx - 1, -1, -1):
                    c = cells[j]
                    if c._tc is target_tc or c._tc in seen_tcs:
                        continue
                    seen_tcs.add(c._tc)
                    t = (_cell_text(c) or "").strip()
                    if not t:
                        continue
                    if t in ("甲方", "乙方", "甲方（买方）", "乙方（卖方）"):
                        # vMerge anchor for the party column — keep scanning.
                        continue
                    if len(t) > 40:
                        # Long prose cell — unlikely to be a label, stop.
                        return None
                    return t.replace("\n", " ").strip()
                return None
        return None

    return _walk_tables(doc.tables)


def _enrich_meaningless_slot_labels(slots: List[Dict[str, Any]], doc) -> None:
    """In-place: replace empty / meaningless slot labels with the row-label
    cell text when the slot lives inside a table cell. The original label is
    preserved in `_meta['original_label']` for debuggability."""
    for s in slots:
        # Only touch highlighted slots — other kinds (label_blank, empty_cell,
        # placeholder_cell, option_choice) already have meaningful labels from
        # their detection paths.
        if s.get("kind") != "highlighted":
            continue
        label = s.get("label")
        if not _is_meaningless_fill_label(label):
            continue
        meta = s.get("_meta") or {}
        paragraph = meta.get("paragraph")
        if paragraph is None:
            continue
        tc = _paragraph_parent_cell_tc(paragraph)
        if tc is None:
            continue
        row_label = _row_label_for_tc(tc, doc)
        if not row_label:
            continue
        meta["original_label"] = label
        s["label"] = row_label


def _guess_label_before(text: str, pos: int) -> Optional[str]:
    """Look at text[:pos] for a 'Label:' segment and return the label."""
    prefix = text[:pos]
    # Crop at the nearest preceding sentence terminator so a label like
    # '层数：' isn't dragged through the previous field's '____/____' or
    # the prior clause. Without this, slot ids for paragraphs with
    # multiple fields look like 'slot_001_________平方米层数'.
    for sep in ("；", ";", "。", "?", "？", "!", "！", "\n"):
        idx = prefix.rfind(sep)
        if idx != -1:
            prefix = prefix[idx + 1:]
    m = _LABEL_BEFORE_RE.search(prefix)
    if m:
        return m.group(1).strip()
    # Fallback: take last 30 chars before pos as context-only label hint
    tail = prefix.strip()
    if tail:
        snippet = tail[-30:]
        return snippet.lstrip("·•-—-– \t") or None
    return None


def _snippet(text: str, start: int, end: int, radius: int = 40) -> str:
    a = max(0, start - radius)
    b = min(len(text), end + radius)
    prefix = "…" if a > 0 else ""
    suffix = "…" if b < len(text) else ""
    return prefix + text[a:b] + suffix


# ---------------------------------------------------------------------- #
# Helpers — slot fill application                                         #
# ---------------------------------------------------------------------- #

def _slot_container_key(slot: Dict[str, Any]) -> Optional[Tuple[str, int]]:
    """Identifier that groups slots living in the same table cell or
    paragraph. Cell takes precedence — two slots in different paragraphs of
    the same cell still count as paired."""
    meta = slot.get("_meta") or {}
    cell = meta.get("cell")
    if cell is not None:
        return ("cell", id(cell))
    para = meta.get("paragraph")
    if para is not None:
        return ("paragraph", id(para))
    return None


def _reconcile_amount_pairs(
    slots: List[Dict[str, Any]],
    slot_values: Dict[str, Any],
) -> Tuple[Dict[str, Any], List[str]]:
    """Detect 大写 / 小写 amount slot pairs (same paragraph or table cell)
    and reconcile their values using deterministic Chinese ↔ Arabic
    conversion. Returns (possibly-updated slot_values, list of notes).

    Rules:
      - Both filled & agree → no-op.
      - Both filled & disagree → trust whichever side parses cleanly. If
        both parse, trust 大写 (capital-numeral grammar is self-checking).
      - Only one filled → compute and inject the other side.
      - Neither parses → leave untouched, append warning note.
    """
    notes: List[str] = []
    if not slot_values:
        return slot_values, notes
    new_values = dict(slot_values)

    groups: Dict[Tuple[str, int], List[Dict[str, Any]]] = defaultdict(list)
    for s in slots:
        key = _slot_container_key(s)
        if key is not None:
            groups[key].append(s)

    for group in groups.values():
        daxie = None
        xiaoxie = None
        for s in group:
            # Classify by label first so two slots in the same paragraph
            # (e.g. "...（大写）___（小写）___") get paired correctly. Using
            # label+context together causes the trailing 小写 slot's context
            # to match _DAXIE_LABEL_RE (the 大写 label is upstream in the
            # paragraph text), misclassifying both slots as 大写.
            label_text = s.get("label") or ""
            if _DAXIE_LABEL_RE.search(label_text):
                if daxie is None:
                    daxie = s
                continue
            if _XIAOXIE_LABEL_RE.search(label_text):
                if xiaoxie is None:
                    xiaoxie = s
                continue
            # Fallback: no usable label — fall back to context, but check
            # 小写 before 大写 since context strings often include both.
            ctx_text = s.get("context") or ""
            if _XIAOXIE_LABEL_RE.search(ctx_text) and not _DAXIE_LABEL_RE.search(ctx_text):
                if xiaoxie is None:
                    xiaoxie = s
            elif _DAXIE_LABEL_RE.search(ctx_text) and not _XIAOXIE_LABEL_RE.search(ctx_text):
                if daxie is None:
                    daxie = s
        if not daxie or not xiaoxie or daxie["id"] == xiaoxie["id"]:
            continue

        d_id, x_id = daxie["id"], xiaoxie["id"]
        d_val = new_values.get(d_id)
        x_val = new_values.get(x_id)
        d_dec = chinese_to_decimal(d_val) if d_val else None
        x_dec = parse_arabic_amount(x_val) if x_val else None

        if d_val and x_val:
            if d_dec is not None and x_dec is not None and d_dec == x_dec:
                continue
            if d_dec is not None and x_dec is not None:
                corrected = format_arabic_amount(d_dec)
                notes.append(
                    f"Amount mismatch reconciled ({d_id}/{x_id}): 大写={d_val!r} "
                    f"(parsed as {d_dec}), 小写 was {x_val!r}, corrected to {corrected!r}."
                )
                new_values[x_id] = corrected
            elif d_dec is not None:
                corrected = format_arabic_amount(d_dec)
                notes.append(
                    f"Amount: 小写 {x_val!r} not parseable; overwritten with "
                    f"{corrected!r} computed from 大写 {d_val!r}."
                )
                new_values[x_id] = corrected
            elif x_dec is not None:
                corrected = decimal_to_chinese(x_dec)
                notes.append(
                    f"Amount: 大写 {d_val!r} not parseable; overwritten with "
                    f"{corrected!r} computed from 小写 {x_val!r}."
                )
                new_values[d_id] = corrected
            else:
                notes.append(
                    f"Amount: neither 大写 {d_val!r} nor 小写 {x_val!r} parsed "
                    f"as a Chinese contract amount; left as-is."
                )
        elif d_val:
            if d_dec is not None:
                corrected = format_arabic_amount(d_dec)
                notes.append(
                    f"Amount: 小写 ({x_id}) inferred as {corrected!r} from 大写 {d_val!r}."
                )
                new_values[x_id] = corrected
        elif x_val:
            if x_dec is not None:
                corrected = decimal_to_chinese(x_dec)
                notes.append(
                    f"Amount: 大写 ({d_id}) inferred as {corrected!r} from 小写 {x_val!r}."
                )
                new_values[d_id] = corrected

    return new_values, notes


def _apply_slot_values(slots: List[Dict[str, Any]], slot_values: Dict[str, Any]) -> Dict[str, Any]:
    """Apply a {slot_id: value} mapping to a freshly scanned slot list.

    Underscore-kind slots are grouped by paragraph and applied in reverse
    text-position order so character offsets remain valid across multiple
    fills in the same paragraph.
    """
    by_id = {s["id"]: s for s in slots}
    # Forgiving lookup by leading numeric token. Lets legacy callers (and
    # LLMs that drop the descriptive tail) hit the right slot when they
    # pass "slot_4" / "slot_004" instead of the canonical "slot_004_签订时间".
    # Built only when there's a unique slot per number — collisions fall
    # through to unknown_slot_ids so we never silently mis-route a value.
    by_num: Dict[str, str] = {}
    num_re = re.compile(r"^slot_0*(\d+)")
    seen_nums: Set[str] = set()
    dup_nums: Set[str] = set()
    for s in slots:
        m = num_re.match(s["id"])
        if m:
            n = m.group(1)
            if n in seen_nums:
                dup_nums.add(n)
            seen_nums.add(n)
            by_num[n] = s["id"]
    for n in dup_nums:
        by_num.pop(n, None)
    # Refuse bare legacy keys (`slot_4`, `slot_004`) when the canonical id
    # carries a descriptive tail (`slot_004_签订时间`). Past incidents had the
    # LLM emit ~120 bare `slot_N` keys for a 79-slot template and rely on
    # numeric collision — values landed on semantically wrong slots. The bare
    # form is also forbidden by the prompt; rejecting at the boundary keeps a
    # buggy caller from silently producing a corrupted document.
    legacy_pattern = re.compile(r"^slot_\d+$")
    rejected_legacy: List[Dict[str, str]] = []
    for k in slot_values:
        if k in by_id:
            continue
        if not legacy_pattern.match(k):
            continue
        m = num_re.match(k)
        if not m:
            continue
        canon = by_num.get(m.group(1))
        if canon and canon != k and legacy_pattern.match(canon) is None:
            rejected_legacy.append({"given": k, "canonical": canon})
    if rejected_legacy:
        return {
            "filled_slot_ids": [],
            "unknown_slot_ids": [],
            "skipped_slot_ids": sorted(slot_values.keys()),
            "skipped_signature_slot_ids": [],
            "skipped_blank_slot_ids": [],
            "rejected_legacy_slot_ids": sorted(
                {item["given"] for item in rejected_legacy}
            ),
            "legacy_canonical_map": {
                item["given"]: item["canonical"] for item in rejected_legacy
            },
            "rejected": True,
            "reconciliation_notes": [
                "Refused to fill: caller used bare legacy slot ids "
                f"({sorted({item['given'] for item in rejected_legacy})}) "
                "instead of the descriptive ids returned by inspect_template. "
                "Re-call fill_template with the canonical ids (e.g. "
                + ", ".join(
                    f"{item['given']!r} → {item['canonical']!r}"
                    for item in rejected_legacy[:5]
                )
                + ")."
            ],
        }
    # Canonicalize legacy keys up front so amount reconciliation and the
    # main fill loop both see canonical ids.
    canonical_values: Dict[str, Any] = {}
    canonicalization_notes: List[str] = []
    for k, v in slot_values.items():
        if k in by_id:
            canonical_values[k] = v
            continue
        m = num_re.match(k)
        if m and m.group(1) in by_num:
            canon = by_num[m.group(1)]
            canonical_values[canon] = v
            if canon != k:
                canonicalization_notes.append(
                    f"Slot id {k!r} resolved to canonical id {canon!r}."
                )
        else:
            canonical_values[k] = v  # keep so it shows up in unknown_slot_ids
    slot_values = canonical_values
    slot_values, reconciliation_notes = _reconcile_amount_pairs(slots, slot_values)
    reconciliation_notes = canonicalization_notes + reconciliation_notes
    filled_ids: List[str] = []
    unknown_ids: List[str] = []
    skipped_signature_ids: List[str] = []
    skipped_blank_ids: List[str] = []

    # Group span-based fills per paragraph (underscores, angle_bracketed,
    # placeholder_phrase, hint_text, highlighted). Per-batch tuple carries the
    # kind + meta so the highlighted post-step can clear w:highlight without
    # touching other formatting.
    span_by_paragraph: Dict[int, List[Tuple[Any, int, int, str, str, str, Dict[str, Any]]]] = defaultdict(list)
    deferred: List[Tuple[str, Dict[str, Any], str, str]] = []  # (kind, meta, value, slot_id)

    for sid, value in slot_values.items():
        slot = by_id.get(sid)
        if not slot:
            unknown_ids.append(sid)
            continue
        meta = slot.get("_meta", {})
        kind = slot["kind"]
        # Signature / seal fields are left blank regardless of the value the
        # caller supplied — guard against accidental fills of 签字 / signature
        # lines (the human signs on paper after printing).
        if meta.get("is_signature"):
            skipped_signature_ids.append(sid)
            continue
        # Blank/None value means "don't fill this slot" — preserve the
        # original placeholder text (underscores, hint runs, whitespace
        # gaps). Earlier behavior collapsed the slot's underline / blank
        # region to empty when the agent passed value="" for a slot it
        # didn't actually want to fill (common in mutually-exclusive
        # option-style sections — e.g. "一次总付" vs "分期支付" where the
        # agent fills only one branch). 0 / False are kept as legitimate
        # fills (str(0) == "0").
        if value is None or (isinstance(value, str) and not value.strip()):
            skipped_blank_ids.append(sid)
            continue
        val_str = "" if value is None else str(value)
        if kind in ("underscores", "angle_bracketed", "placeholder_phrase",
                    "hint_text", "highlighted"):
            paragraph = meta["paragraph"]
            # For highlighted slots with a recognized scaffold (number+unit,
            # currency prefix, percentage, etc.), auto-attach the scaffold when
            # the user-provided value is a bare number — fixes the "15个工作日 →
            # 20" data-loss bug.
            effective_val = val_str
            if kind == "highlighted":
                scaffold = meta.get("scaffold")
                if scaffold:
                    effective_val = _reconstruct_with_scaffold(scaffold, val_str)
            # Composite underscore slot (digit-cell pattern): expand to one
            # span replacement per underscore position — first gets the value,
            # rest are blanked. Prevents the "8500" concatenated-into-every-
            # underscore bug.
            if kind == "underscores" and meta.get("composite"):
                positions = meta.get("positions", [])
                for i, (pstart, pend) in enumerate(positions):
                    fill_val = effective_val if i == 0 else ""
                    span_by_paragraph[id(paragraph)].append(
                        (paragraph, pstart, pend, fill_val, sid if i == 0 else None,
                         kind, meta)
                    )
            else:
                # Border-only slot: paragraph has a horizontal border but no
                # underscore characters. Write the value into the paragraph
                # itself, centered, so it sits above the border line.
                if kind == "underscores" and meta.get("border_only"):
                    deferred.append(("border_line", meta, effective_val, sid))
                    continue
                # For single-span underscore fills, center the value within
                # the original underscore width so the line stays balanced
                # (e.g. "____ Alice ____") instead of leaving the value
                # bumped against one side of the line. Numbered-list bodies
                # like '(1) ____________' get left-aligned fill instead, so
                # the value follows the marker rather than floating in the
                # middle of the line.
                fill_val = effective_val
                if kind == "underscores" and meta.get("source") != "date_scaffold":
                    paragraph_text = "".join((r.text or "") for r in paragraph.runs)
                    span_text = paragraph_text[meta["start"]:meta["end"]]
                    prefix = paragraph_text[:meta["start"]]
                    align = "left" if _prefix_is_only_list_marker(prefix) else "center"
                    fill_val = _center_in_underscore_span(
                        effective_val, span_text, meta.get("pad_char"), align=align
                    )
                span_by_paragraph[id(paragraph)].append(
                    (paragraph, meta["start"], meta["end"], fill_val, sid, kind, meta)
                )
        elif kind == "label_blank":
            deferred.append(("label_blank", meta, val_str, sid))
        elif kind == "empty_cell":
            deferred.append(("empty_cell", meta, val_str, sid))
        elif kind == "placeholder_cell":
            deferred.append(("placeholder_cell", meta, val_str, sid))
        elif kind == "section_body_empty":
            deferred.append(("section_body_empty", meta, val_str, sid))
        elif kind == "option_choice":
            options = meta.get("options") or []
            chosen = _normalize_option_choice(value, len(options))
            valid_indices = [opt["index"] for opt in options]
            if chosen is None or chosen not in valid_indices:
                reconciliation_notes.append(
                    f"Option-choice slot {sid}: value {value!r} did not match a "
                    f"valid option (expected one of {valid_indices} or labels "
                    f"like '第一种'/'第二种'). Slot skipped; document left unchanged."
                )
                continue
            choice_meta = dict(meta)
            choice_meta["chosen_index"] = chosen
            deferred.append(("option_choice", choice_meta, val_str, sid))
        else:
            unknown_ids.append(sid)

    for batch in span_by_paragraph.values():
        batch.sort(key=lambda x: x[1], reverse=True)  # latest start first
        for paragraph, start, end, val_str, sid, kind, meta in batch:
            if _replace_paragraph_span(paragraph, start, end, val_str):
                if kind == "highlighted":
                    # Clear highlight on every run that originally carried it
                    # in this span. Other run-level formatting (font, size,
                    # bold/italic/color) stays untouched.
                    for run in meta.get("runs", []):
                        _clear_highlight(run)
                if sid is not None:
                    filled_ids.append(sid)

    for kind, meta, val_str, sid in deferred:
        if kind == "label_blank":
            # Linked-border case: the next paragraph is an empty bordered line.
            # Write the value into THAT paragraph (centered) so it sits over
            # the visible line instead of appearing after the label's colon.
            target = meta.get("border_target")
            if target is not None:
                if _write_to_empty_paragraph(target, val_str):
                    _set_paragraph_centered(target)
                    filled_ids.append(sid)
                continue
            # Inline-border case: the label paragraph itself has a bottom
            # border. Center the whole paragraph so the value sits in the
            # middle of the line below the label.
            if meta.get("bordered_paragraph"):
                if _append_after_label(meta["paragraph"], val_str):
                    _set_paragraph_centered(meta["paragraph"])
                    filled_ids.append(sid)
                continue
            if _append_after_label(meta["paragraph"], val_str):
                filled_ids.append(sid)
        elif kind == "empty_cell":
            if _write_to_empty_cell(meta["cell"], val_str):
                filled_ids.append(sid)
        elif kind == "placeholder_cell":
            if _write_to_placeholder_cell(meta["cell"], val_str):
                filled_ids.append(sid)
        elif kind == "section_body_empty":
            if _fill_section_body_empty(meta["paragraph"], val_str):
                filled_ids.append(sid)
        elif kind == "border_line":
            if _write_to_empty_paragraph(meta["paragraph"], val_str):
                _set_paragraph_centered(meta["paragraph"])
                filled_ids.append(sid)
        elif kind == "option_choice":
            # Drop: prompt paragraph + every option's header paragraph
            # (the chosen option's header is structural — '第N种：…' — and
            # is removed too) + every non-chosen option's body paragraphs.
            chosen = meta["chosen_index"]
            to_drop = [meta["prompt_paragraph"]]
            for opt in meta["options"]:
                to_drop.append(opt["header_paragraph"])
                if opt["index"] != chosen:
                    to_drop.extend(opt["body_paragraphs"])
            dropped_any = False
            for p in to_drop:
                if p is not None and _drop_paragraph(p):
                    dropped_any = True
            if dropped_any:
                filled_ids.append(sid)
                reconciliation_notes.append(
                    f"Option-choice slot {sid}: kept 第{chosen}种, removed "
                    f"prompt + {len(meta['options']) - 1} other option(s)."
                )

    requested = set(slot_values.keys())
    signatures = set(skipped_signature_ids)
    blanks = set(skipped_blank_ids)
    skipped = sorted(
        (requested - set(filled_ids) - set(unknown_ids) - signatures - blanks)
        | signatures
        | blanks
    )
    outcome = {
        "filled_slot_ids": sorted(filled_ids),
        "unknown_slot_ids": sorted(unknown_ids),
        "skipped_slot_ids": skipped,
        "skipped_signature_slot_ids": sorted(signatures),
        "skipped_blank_slot_ids": sorted(blanks),
    }
    if reconciliation_notes:
        outcome["reconciliation_notes"] = reconciliation_notes
    return outcome


def _append_after_label(paragraph, value: str) -> bool:
    runs = paragraph.runs
    if not runs:
        paragraph.add_run(" " + value)
        return True
    # Append to the last non-empty run so we inherit its formatting
    for r in reversed(runs):
        if r.text:
            sep = "" if r.text.endswith((" ", "\t")) else " "
            r.text = r.text + sep + value
            return True
    runs[0].text = (runs[0].text or "") + " " + value
    return True


def _write_to_placeholder_cell(cell, value: str) -> bool:
    """Replace ALL text in a seal/date placeholder cell with `value`.

    Splits on '\\n' so multi-line values become multiple paragraphs.
    Existing paragraphs are reused (preserving alignment and first-run
    formatting); any leftover paragraphs are text-cleared but kept so the
    cell's vertical structure / merge boundaries stay intact.
    """
    lines = value.split("\n") if value else [""]
    paragraphs = list(cell.paragraphs)
    for i, line in enumerate(lines):
        if i < len(paragraphs):
            p = paragraphs[i]
            if p.runs:
                p.runs[0].text = line
                for r in p.runs[1:]:
                    r.text = ""
            else:
                p.add_run(line)
        else:
            cell.add_paragraph(line)
    for p in paragraphs[len(lines):]:
        for r in p.runs:
            r.text = ""
    return True


def _write_to_empty_cell(cell, value: str) -> bool:
    if cell.paragraphs:
        p = cell.paragraphs[0]
        if p.runs:
            p.runs[0].text = value
            for extra in p.runs[1:]:
                extra.text = ""
        else:
            p.add_run(value)
        return True
    cell.add_paragraph(value)
    return True


def _write_to_empty_paragraph(paragraph, value: str) -> bool:
    """Write text into a paragraph that's currently empty (section_body_empty)."""
    if paragraph.runs:
        paragraph.runs[0].text = value
        for extra in paragraph.runs[1:]:
            extra.text = ""
    else:
        paragraph.add_run(value)
    return True


# Leading numbered-list marker on a single line. Captures the numeric prefix
# in group 1 and the body in group 2. Accepts `1. body`, `1、body`, `(1) body`,
# `1) body`, `1) body`, `1．body`, etc. We don't accept `第N条 body` here —
# that's section-numbering, not a list-item style.
_NUMBERED_LINE_RE = re.compile(
    r"^\s*"
    r"(?:[(（])?\s*(\d{1,3})\s*[)）]?"   # number with optional brackets
    # Separator: ASCII '.', full-width '．', dunhao '、', or colon. The
    # following whitespace is optional — `1、甲方` (no space) is common in
    # CN prose. For unbracketed numbers without any separator we still
    # require at least one whitespace char so prose like `2024年` doesn't
    # match.
    r"(?:\s*[.．、:：]\s*|\s+)"
    r"(.+?)\s*$"
)


def _split_numbered_lines(value: str) -> Optional[List[str]]:
    """If `value` is a pure numbered list (every line a `N. body` item,
    numbers incrementing by 1 from any start), return the bodies. Otherwise
    None.

    Lenient about the starting number — agents occasionally continue the
    counter across sections (writing `8. a\\n9. b` for the second section).
    The rendered Word list will restart at 1 regardless, since the fill
    path creates a fresh <w:num> per list block.
    """
    if not value:
        return None
    lines = [ln for ln in value.splitlines() if ln.strip()]
    if len(lines) < 2:
        return None
    bodies: List[str] = []
    first_num: Optional[int] = None
    for i, line in enumerate(lines):
        m = _NUMBERED_LINE_RE.match(line)
        if not m:
            return None
        n = int(m.group(1))
        if first_num is None:
            first_num = n
        elif n != first_num + i:
            return None
        body = m.group(2).strip()
        if not body:
            return None
        bodies.append(body)
    return bodies


def _segment_value(value: str) -> List[Tuple[str, Any]]:
    """Split a fill value into a sequence of `(kind, content)` segments.

    - `('prose', text)`   — a plain paragraph (text may contain '\\n' for
      multiple consecutive non-list lines, kept together as separate
      paragraphs at render time).
    - `('list', bodies)`  — a contiguous block of numbered lines whose
      numbers increment by 1. `bodies` is the list of item bodies with
      their numeric prefix stripped.

    Handles:
      - intro prose followed by a list (`产品清单如下：\\n1. …`)
      - header prose between two independent lists
        (`甲方：\\n1.…\\n2.…\\n\\n乙方：\\n1.…\\n2.…`)
      - agents that don't restart at 1 between sections (numbers are
        accepted as long as they increment by 1)
    """
    if not value:
        return []
    segments: List[Tuple[str, Any]] = []
    prose_buf: List[str] = []
    list_buf: List[Tuple[int, str]] = []

    def flush_prose():
        if prose_buf:
            segments.append(("prose", "\n".join(prose_buf)))
            prose_buf.clear()

    def flush_list():
        if len(list_buf) >= 2:
            segments.append(("list", [body for (_, body) in list_buf]))
        else:
            # Single numbered line is not a list — fold back into prose.
            for n, body in list_buf:
                prose_buf.append(f"{n}. {body}")
        list_buf.clear()

    for raw_line in value.splitlines():
        line = raw_line.strip()
        if not line:
            # Blank line separates segments.
            flush_list()
            flush_prose()
            continue
        m = _NUMBERED_LINE_RE.match(line)
        if m:
            n = int(m.group(1))
            body = m.group(2).strip()
            if list_buf and n != list_buf[-1][0] + 1:
                # Counter broke — end this list, start fresh.
                flush_list()
            if not list_buf:
                flush_prose()
            list_buf.append((n, body))
        else:
            flush_list()
            prose_buf.append(line)
    flush_list()
    flush_prose()
    return segments


def _ensure_decimal_abstract_num(numbering_part) -> Optional[int]:
    """Return an abstractNumId whose level-0 renders as `1. 2. 3.` decimal.

    Reuses an existing definition when one matches; otherwise appends a
    new one. Cached on the numbering part so the same document only
    accumulates one new abstractNum no matter how many lists we add.
    """
    try:
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
    except ImportError:
        return None
    if numbering_part is None:
        return None
    cached = getattr(numbering_part, "_dts_decimal_abstract_id", None)
    if cached is not None:
        return cached
    el = numbering_part.element

    for absn in el.findall(qn("w:abstractNum")):
        aid = absn.get(qn("w:abstractNumId"))
        for lvl in absn.findall(qn("w:lvl")):
            if lvl.get(qn("w:ilvl")) != "0":
                continue
            numFmt = lvl.find(qn("w:numFmt"))
            lvlText = lvl.find(qn("w:lvlText"))
            fmt = numFmt.get(qn("w:val")) if numFmt is not None else None
            txt = lvlText.get(qn("w:val")) if lvlText is not None else None
            if fmt == "decimal" and txt in ("%1.", "%1．"):
                try:
                    cached = int(aid)
                except (TypeError, ValueError):
                    cached = None
                if cached is not None:
                    numbering_part._dts_decimal_abstract_id = cached
                    return cached
            break

    existing_abs = [
        int(a.get(qn("w:abstractNumId")))
        for a in el.findall(qn("w:abstractNum"))
        if (a.get(qn("w:abstractNumId")) or "").lstrip("-").isdigit()
    ]
    new_abs_id = (max(existing_abs) + 1) if existing_abs else 0
    absn = OxmlElement("w:abstractNum")
    absn.set(qn("w:abstractNumId"), str(new_abs_id))
    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start"); start.set(qn("w:val"), "1"); lvl.append(start)
    numFmt = OxmlElement("w:numFmt"); numFmt.set(qn("w:val"), "decimal"); lvl.append(numFmt)
    lvlText = OxmlElement("w:lvlText"); lvlText.set(qn("w:val"), "%1."); lvl.append(lvlText)
    lvlJc = OxmlElement("w:lvlJc"); lvlJc.set(qn("w:val"), "left"); lvl.append(lvlJc)
    pPr_lvl = OxmlElement("w:pPr")
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "720")
    ind.set(qn("w:hanging"), "360")
    pPr_lvl.append(ind)
    lvl.append(pPr_lvl)
    absn.append(lvl)
    first_num = el.find(qn("w:num"))
    if first_num is not None:
        first_num.addprevious(absn)
    else:
        el.append(absn)

    numbering_part._dts_decimal_abstract_id = new_abs_id
    return new_abs_id


def _new_decimal_num_id(paragraph) -> Optional[int]:
    """Create and return a FRESH `<w:num>` for decimal `1. 2. 3.` numbering.

    The underlying abstractNum is shared across the document (created once),
    but each call appends a brand-new `<w:num>` referencing it — that way
    every list gets its own counter and Word renders each list starting at 1.
    Returns None if the document has no numbering part.
    """
    try:
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
    except ImportError:
        return None
    try:
        numbering_part = paragraph.part.numbering_part
    except Exception:
        return None
    if numbering_part is None:
        return None
    abstract_id = _ensure_decimal_abstract_num(numbering_part)
    if abstract_id is None:
        return None
    el = numbering_part.element
    existing_num = [
        int(n.get(qn("w:numId")))
        for n in el.findall(qn("w:num"))
        if (n.get(qn("w:numId")) or "").lstrip("-").isdigit()
    ]
    new_num_id = (max(existing_num) + 1) if existing_num else 1
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(new_num_id))
    aref = OxmlElement("w:abstractNumId")
    aref.set(qn("w:val"), str(abstract_id))
    num.append(aref)
    # Explicit start-from-1 override. Without this, some Word versions
    # continue the counter across sibling <w:num> elements that share an
    # abstractNum — so section 五's list ends up numbered 9, 10, 11, 12 as
    # a continuation of section 二's 1, 2, 3, 4 even though we handed out
    # distinct numIds. Naming a startOverride forces Word to reset.
    lvlOverride = OxmlElement("w:lvlOverride")
    lvlOverride.set(qn("w:ilvl"), "0")
    startOverride = OxmlElement("w:startOverride")
    startOverride.set(qn("w:val"), "1")
    lvlOverride.append(startOverride)
    num.append(lvlOverride)
    el.append(num)
    return new_num_id


def _apply_list_numbering(paragraph, num_id: int, ilvl: int = 0) -> None:
    """Set <w:numPr><w:ilvl/><w:numId/></w:numPr> on the paragraph's pPr."""
    try:
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
    except ImportError:
        return
    pPr = paragraph._p.get_or_add_pPr()
    for old in pPr.findall(qn("w:numPr")):
        pPr.remove(old)
    numPr = OxmlElement("w:numPr")
    ilvl_el = OxmlElement("w:ilvl"); ilvl_el.set(qn("w:val"), str(ilvl))
    numId_el = OxmlElement("w:numId"); numId_el.set(qn("w:val"), str(num_id))
    numPr.append(ilvl_el)
    numPr.append(numId_el)
    pPr.append(numPr)


def _insert_paragraph_after(paragraph, text: str):
    """Insert a new <w:p> sibling immediately after `paragraph`, copying its
    pPr (so style / alignment carry over), and return the new Paragraph
    wrapper. The new paragraph contains a single run with `text`.
    """
    from docx.text.paragraph import Paragraph
    from copy import deepcopy
    from docx.oxml.ns import qn
    new_p = deepcopy(paragraph._p)
    # Drop existing runs (we want a fresh single run).
    for r in new_p.findall(qn("w:r")):
        new_p.remove(r)
    # Drop numPr so the caller can apply fresh numbering.
    pPr = new_p.find(qn("w:pPr"))
    if pPr is not None:
        for old in pPr.findall(qn("w:numPr")):
            pPr.remove(old)
    paragraph._p.addnext(new_p)
    wrapper = Paragraph(new_p, paragraph._parent)
    if text:
        wrapper.add_run(text)
    return wrapper


def _fill_section_body_empty(paragraph, value: str) -> bool:
    """Fill a section_body_empty slot.

    Segments `value` into prose blocks and numbered-list blocks (see
    `_segment_value`). The existing empty paragraph is reused for the
    first written line; subsequent lines are inserted after it. Each
    list block gets its OWN fresh numId so Word renders independent
    lists each starting at 1, regardless of what numbers the agent used.

    Falls back to a single-paragraph plain-text write when there's no
    structured content (single short prose value, no lists).
    """
    segments = _segment_value(value)
    if not segments:
        return _write_to_empty_paragraph(paragraph, value)
    if len(segments) == 1 and segments[0][0] == "prose" and "\n" not in segments[0][1]:
        # Fast path: a single short prose value (the common case for
        # short, non-list section bodies).
        return _write_to_empty_paragraph(paragraph, segments[0][1])

    # Flatten segments into an ordered list of `(text, num_id_or_None)`.
    plan: List[Tuple[str, Optional[int]]] = []
    for kind, content in segments:
        if kind == "prose":
            for line in str(content).split("\n"):
                line = line.strip()
                if line:
                    plan.append((line, None))
        elif kind == "list":
            list_num_id = _new_decimal_num_id(paragraph)
            for body in content:
                plan.append((body, list_num_id))
    if not plan:
        return _write_to_empty_paragraph(paragraph, value)

    first_text, first_num_id = plan[0]
    _write_to_empty_paragraph(paragraph, first_text)
    if first_num_id is not None:
        _apply_list_numbering(paragraph, first_num_id, ilvl=0)
    else:
        # Strip any pre-existing numPr that the template paragraph might
        # have inherited — prose paragraphs should not be auto-numbered.
        try:
            from docx.oxml.ns import qn
            pPr = paragraph._p.find(qn("w:pPr"))
            if pPr is not None:
                for old in pPr.findall(qn("w:numPr")):
                    pPr.remove(old)
        except Exception:
            pass
    prev = paragraph
    for text, num_id in plan[1:]:
        prev = _insert_paragraph_after(prev, text)
        if num_id is not None:
            _apply_list_numbering(prev, num_id, ilvl=0)
    return True


# ---------------------------------------------------------------------- #
# Helpers — highlighted runs                                              #
# ---------------------------------------------------------------------- #

def _run_is_highlighted(run) -> bool:
    """True if the run carries a Word highlight (yellow, green, cyan, etc.).

    `font.highlight_color` is a WD_COLOR_INDEX enum; None means no highlight,
    AUTO means an explicit "no highlight" — both are treated as not-highlighted.
    """
    try:
        hl = run.font.highlight_color
    except Exception:
        return False
    if hl is None:
        return False
    try:
        from docx.enum.text import WD_COLOR_INDEX
        if hl == WD_COLOR_INDEX.AUTO:
            return False
    except Exception:
        pass
    return True


def _clear_highlight(run) -> None:
    """Remove the w:highlight attribute from a run, preserving all other
    formatting (font, size, color, bold/italic, etc.). Idempotent.
    """
    try:
        from docx.enum.text import WD_COLOR_INDEX
        run.font.highlight_color = WD_COLOR_INDEX.AUTO
    except Exception:
        # Fall back to direct XML edit if the high-level setter isn't happy
        try:
            rPr = run._element.find(
                "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}rPr"
            )
            if rPr is not None:
                hl = rPr.find(
                    "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}highlight"
                )
                if hl is not None:
                    rPr.remove(hl)
        except Exception:
            pass


def _find_highlighted_spans(
    paragraph,
) -> List[Tuple[int, int, List[Any], str]]:
    """Find contiguous spans of highlighted runs in a paragraph.

    Returns a list of (start_offset, end_offset, runs, text) tuples in
    paragraph-text coordinates. Consecutive highlighted runs are merged into
    one span — a logical "edit this" region is often split into 2-3 runs by
    Word for unrelated reasons (e.g. autocorrect boundaries).
    """
    spans: List[Tuple[int, int, List[Any], str]] = []
    offset = 0
    cur_start: Optional[int] = None
    cur_runs: List[Any] = []
    cur_text: List[str] = []
    for run in paragraph.runs:
        rt = run.text or ""
        if _run_is_highlighted(run) and rt:
            if cur_start is None:
                cur_start = offset
            cur_runs.append(run)
            cur_text.append(rt)
        else:
            if cur_start is not None:
                spans.append((cur_start, offset, cur_runs, "".join(cur_text)))
                cur_start = None
                cur_runs = []
                cur_text = []
        offset += len(rt)
    if cur_start is not None:
        spans.append((cur_start, offset, cur_runs, "".join(cur_text)))
    return spans


# ---------------------------------------------------------------------- #
# Helpers — hint-text / heading / empty detection                         #
# ---------------------------------------------------------------------- #

def _is_heading(paragraph) -> bool:
    try:
        name = paragraph.style.name if paragraph.style else ""
    except Exception:
        name = ""
    return bool(name and _HEADING_STYLE_RE.match(name))


def _paragraph_has_list_numbering(paragraph) -> bool:
    """True when the paragraph carries auto list numbering (w:pPr/w:numPr).

    Chinese contract templates often style section headers like 不可抗力 /
    违约责任 with a numbered list (numId=1) instead of putting "九、" / "十、"
    in the run text, so a text-only regex won't see them as section breaks.
    """
    try:
        pPr = paragraph._p.pPr
    except Exception:
        return False
    return pPr is not None and pPr.numPr is not None


def _paragraph_list_level(paragraph) -> Optional[int]:
    """Return the numbering ilvl (0-based) for a numbered paragraph, else None.

    Used to distinguish a section heading from its body items: a "违约责任：" at
    ilvl=0 followed by "甲方应..." items at ilvl=1 means the heading is
    structurally a parent, not a sibling — so a value should fill the body
    items, not be appended after the heading's colon.
    """
    try:
        from docx.oxml.ns import qn
        pPr = paragraph._p.pPr
        if pPr is None or pPr.numPr is None:
            return None
        ilvl_el = pPr.numPr.find(qn("w:ilvl"))
        if ilvl_el is None:
            return 0
        val = ilvl_el.get(qn("w:val"))
        return int(val) if val is not None else 0
    except Exception:
        return None


def _is_section_heading_like(paragraph) -> bool:
    """True when the paragraph looks like a section heading.

    Covers two cases:
      1. Built-in Heading / 标题 style — `_is_heading` already handles this.
      2. Auto-numbered list paragraph (w:numPr) whose run text is short and
         has no colon — e.g. `不可抗力`, `违约责任`, `合同争议的解决方式`.
         These render with `九、` / `十、` prefixes from the numbering
         definition but the text itself is just the title. Body sub-items
         (`本合同技术开发风险责任由乙方承担。`) are longer and excluded.
    """
    if _is_heading(paragraph):
        return True
    if not _paragraph_has_list_numbering(paragraph):
        return False
    text = "".join((r.text or "") for r in paragraph.runs).strip()
    if not text:
        return False
    # Heading-like: short title (<= 40 chars), no inline colon or sentence
    # terminator, no leading sub-key like "10.1".
    if len(text) > 40:
        return False
    if re.search(r"[:：。；！？]", text):
        return False
    if re.match(r"^\s*\d+\.\d+", text):
        return False
    return True


_GAP_DECORATIVE_RE = re.compile(r"^[\s 　.\-—–·•・|/\\,;:：（）()\[\]]*$")


def _gaps_are_decorative(text: str, matches) -> bool:
    """True when the inter-match text between consecutive underscore runs is
    purely whitespace / decorator characters (no Chinese, no letters, no digits).
    Examples that PASS: '¥ ____ ____ ____', '____.____.____', '____ - ____ - ____'.
    Examples that FAIL: '姓名: ____ 年龄: ____ 性别: ____' (real labels in between).
    """
    for i in range(len(matches) - 1):
        gap = text[matches[i].end():matches[i + 1].start()]
        if not _GAP_DECORATIVE_RE.match(gap):
            return False
        # Also reject very long gaps even if they happen to match — likely
        # separate fields with a punctuation-only label.
        if len(gap) > 12:
            return False
    return True


def _group_decorative_underscore_matches(text, matches):
    """Group consecutive underscore matches that are separated only by
    decorator-only gaps (whitespace, '/', '\\', '.', '-', ...).

    Returns a list of groups; each group is a non-empty list of matches.
    Two consecutive matches join a group iff the text between them passes
    `_GAP_DECORATIVE_RE` and is ≤ 12 chars. A group of one match is just
    that match standing alone (no neighbors qualified to merge with it).
    """
    if not matches:
        return []
    groups = [[matches[0]]]
    for i in range(1, len(matches)):
        gap = text[matches[i - 1].end():matches[i].start()]
        if _GAP_DECORATIVE_RE.match(gap) and len(gap) <= 12:
            groups[-1].append(matches[i])
        else:
            groups.append([matches[i]])
    return groups


# A paragraph made up of only underscores plus pure decoration is a divider,
# not a fillable field. Removing underscore + decoration chars from such a
# paragraph leaves nothing behind.
_NON_UNDERSCORE_NON_DECOR_RE = re.compile(
    r"[_＿\s　.\-—–·•・|/\\,;:：（）()\[\]]+"
)


def _looks_like_signature(label: Optional[str], text: str, start: int, end: int) -> bool:
    """True if the slot's label or its near-context names a signature/seal field.

    Looks at the label, plus a small window of text on either side of the
    matched span (so '________ (signature)' is caught even though label
    detection only looks at text before the underscores).
    """
    if label and _SIGNATURE_RE.search(label):
        return True
    pre_window = text[max(0, start - 40):start]
    if _SIGNATURE_RE.search(pre_window):
        return True
    post_window = text[end:end + 40]
    if _SIGNATURE_RE.search(post_window):
        return True
    return False


def _is_decorative_underscore_paragraph(text: str, matches) -> bool:
    """True if the paragraph is a standalone underscore divider with no real
    label or signature context — i.e. content outside the underscores is only
    whitespace / punctuation. Avoids flagging dividers as fillable slots.
    """
    if not matches:
        return False
    stripped = text.strip()
    if not stripped:
        return False
    residue = _NON_UNDERSCORE_NON_DECOR_RE.sub("", stripped)
    if residue:
        return False
    total_underscores = sum(m.end() - m.start() for m in matches)
    return total_underscores >= 20


def _visual_width(s: str) -> int:
    """Approximate visual width of `s` in half-widths.

    Word renders proportional fonts: an ASCII space is ~1 half-width while a
    CJK ideograph / full-width punctuation glyph is ~2. Counting code points
    (len()) produces "centered" output that drifts off-center in the rendered
    document, especially in mixed Chinese / English contracts where a long
    Chinese value gets padded with too few ASCII spaces on each side.
    """
    w = 0
    for c in s:
        cp = ord(c)
        # CJK Unified Ideographs (incl. Ext A/B common ranges)
        if 0x4E00 <= cp <= 0x9FFF or 0x3400 <= cp <= 0x4DBF:
            w += 2
        # CJK symbols & punctuation, hiragana, katakana
        elif 0x3000 <= cp <= 0x33FF:
            w += 2
        # Full-width Latin / half-width katakana etc.
        elif 0xFF00 <= cp <= 0xFFEF:
            w += 2
        else:
            w += 1
    return w


# Enumeration / list markers at the very start of a paragraph. When a slot's
# paragraph begins with one of these, the underscore span is "the body of a
# numbered list item" and the value should sit *right after* the marker rather
# than be centered in the line.
#   (1), (a), （1）, （一）, 【1】, 1., 1)、, ①, 一、, 第一条, 1：
_LIST_MARKER_RE = re.compile(
    r"^\s*(?:"
    r"[（(【]\s*[0-9０-９一二三四五六七八九十百千零〇iIvVxXa-zA-Z]{1,6}\s*[)）】]"
    r"|[0-9０-９]{1,3}\s*[.、．:：)）]"
    r"|[一二三四五六七八九十百千零〇]{1,4}\s*[、.．:：)）]"
    r"|[①-⑳㈠-㈩❶-❿]"
    r"|第\s*[0-9０-９一二三四五六七八九十百千零〇]{1,4}\s*[条款项节章]"
    r")\s*"
)


def _prefix_is_only_list_marker(prefix: str) -> bool:
    """True if `prefix` consists solely of a list / enumeration marker plus
    optional whitespace — e.g. '（1）', '（1） ', '1. '. Used to switch
    underscore fills from centered to left-aligned only when the slot is
    the *body* of a numbered line (the marker is followed directly by the
    underscore span). Paragraphs like '2．研究开发经费由甲方 ___ ...' have
    real content between the marker and the slot, so they stay centered.
    """
    if not prefix:
        return False
    m = _LIST_MARKER_RE.match(prefix)
    if not m:
        return False
    return not prefix[m.end():].strip()


def _has_cjk(s: str) -> bool:
    for c in s:
        cp = ord(c)
        if 0x4E00 <= cp <= 0x9FFF or 0x3400 <= cp <= 0x4DBF or 0xF900 <= cp <= 0xFAFF:
            return True
    return False


def _build_space_padding(visual_half_widths: int, prefer_full_width: bool) -> str:
    """Produce a whitespace padding string that fills exactly `visual_half_widths`
    in Word. Uses ideographic spaces (each = 2 half-widths in Chinese fonts)
    when the value is CJK-heavy, with at most one trailing ASCII space to
    correct an odd parity. Falls back to plain ASCII spaces otherwise.
    """
    if visual_half_widths <= 0:
        return ""
    if prefer_full_width:
        fw = visual_half_widths // 2
        extra = visual_half_widths - fw * 2  # 0 or 1
        return ("　" * fw) + (" " * extra)
    return " " * visual_half_widths


def _center_in_underscore_span(
    value: str,
    span_text: str,
    pad_char: Optional[str] = None,
    align: str = "center",
) -> str:
    """Place `value` inside the span and pad to preserve the original line width.

    `align="center"` (default) splits padding on both sides; `align="left"`
    puts all padding on the right so the value sits at the start of the line
    — used for numbered-list bodies like '(1) ____________' where the value
    should appear right after the marker rather than floating in the middle.

    Padding is computed in **visual half-widths** (CJK / full-width glyphs = 2,
    ASCII = 1). For underline-character spans (`_`, `＿`) the same glyph is
    re-used as the pad. For underlined-whitespace spans the pad is regular
    spaces — but ideographic spaces when the value is CJK-heavy, because an
    ASCII space in a Chinese font renders narrower than half a CJK width and
    char-count centering drifts off in the rendered document. Falls back to
    the bare value when there is no room.
    """
    target_w = _visual_width(span_text)
    if target_w <= 0 or not value:
        return value
    value_w = _visual_width(value)
    # Reserve a single ASCII space on each side of the value for legibility.
    if value_w + 2 >= target_w:
        return value
    if pad_char is None:
        pad_char = "＿" if span_text.count("＿") > span_text.count("_") else "_"
    pad_visual = target_w - value_w - 2
    if align == "left":
        left_visual = 0
        right_visual = pad_visual
    else:
        left_visual = pad_visual // 2
        right_visual = pad_visual - left_visual
    is_space_pad = pad_char in (" ", "　", " ")
    if is_space_pad:
        use_fw = _has_cjk(value) or _has_cjk(span_text)
        left_pad = _build_space_padding(left_visual, use_fw)
        right_pad = _build_space_padding(right_visual, use_fw)
    else:
        pad_unit = _visual_width(pad_char) or 1
        left_pad = pad_char * (left_visual // pad_unit)
        right_pad = pad_char * (right_visual // pad_unit)
    return left_pad + " " + value + " " + right_pad


def _paragraph_has_horizontal_border(paragraph) -> bool:
    """True if the paragraph carries a visible bottom or between border
    (Word renders these as a horizontal rule, the most common author trick
    for 'fillable line' templates that contain no underscore characters).
    """
    try:
        from docx.oxml.ns import qn
    except Exception:
        return False
    try:
        pPr = paragraph._element.find(qn("w:pPr"))
        if pPr is None:
            return False
        pBdr = pPr.find(qn("w:pBdr"))
        if pBdr is None:
            return False
        for edge in ("w:bottom", "w:between", "w:top"):
            b = pBdr.find(qn(edge))
            if b is None:
                continue
            val = b.get(qn("w:val"))
            if val and val not in ("none", "nil"):
                return True
    except Exception:
        return False
    return False


def _find_underlined_whitespace_spans(
    paragraph,
) -> List[Tuple[int, int, List[Any], str]]:
    """Find consecutive runs whose visible text is only whitespace AND that
    carry an underline format. Word authors create fillable lines this way:
    select a long stretch of spaces / non-breaking spaces and apply underline.

    Returns (start_offset, end_offset, runs, text) tuples in paragraph-text
    coordinates. Requires total length ≥ 3 to filter incidental single-space
    underlines.
    """
    spans: List[Tuple[int, int, List[Any], str]] = []
    offset = 0
    cur_start: Optional[int] = None
    cur_runs: List[Any] = []
    cur_text: List[str] = []
    for run in paragraph.runs:
        rt = run.text or ""
        try:
            ul = bool(run.underline)
        except Exception:
            ul = False
        if ul and rt and not rt.strip():
            if cur_start is None:
                cur_start = offset
            cur_runs.append(run)
            cur_text.append(rt)
        else:
            if cur_start is not None:
                spans.append((cur_start, offset, cur_runs, "".join(cur_text)))
                cur_start = None
                cur_runs = []
                cur_text = []
        offset += len(rt)
    if cur_start is not None:
        spans.append((cur_start, offset, cur_runs, "".join(cur_text)))
    return [s for s in spans if len(s[3]) >= 3]


# Boundary tokens for extending an underlined-whitespace fill region.
# A slot's effective replace span extends across draft marks ('/','／'),
# whitespace, and inline sample text until one of these is reached. Longer
# alternatives (平方米, （大写）, etc.) appear before bare single-char units so
# Python's `re.match` consumes the longest valid token first. The same token
# set is used walking RIGHT (forward extension stops AT the match) and LEFT
# (backward extension stops AFTER the match).
_UW_BOUNDARY_TOKENS = (
    "平方米", "工作日", "小时", "分钟",
    "（大写）", "(大写)", "（小写）", "(小写)",
)
# Right-side boundary chars for an underlined-whitespace fill region.
# Unit chars `份` and `种` are common contract / form units that should
# always terminate a fill ("本合同一式X份", "三种方式"). Parens `（(）)`
# are boundaries so a gap immediately followed by an inline option marker
# (`（甲、乙、双）方所有`) doesn't get the option text + trailing prose
# swallowed into the slot. Comma and 顿号 are deliberately NOT in this set
# — they appear inside sample list text between merged UW gaps (see
# test_uw_spans_merge_across_sample_list).
_UW_BOUNDARY_SINGLE_CHARS = ":：;；。?？!！米元日天周月年个号份种（(）)％%¥$￥€£"
_UW_RIGHT_STOP_RE = re.compile(
    r"[" + re.escape(_UW_BOUNDARY_SINGLE_CHARS) + r"]"
    r"|" + r"|".join(re.escape(t) for t in _UW_BOUNDARY_TOKENS)
)

# Merge-blocker between two UW gaps. If ANY right-stop char (sentence
# terminator, unit char, paren, etc.) appears in the prose between gap1's
# trailing-whitespace edge and gap2's start, the gaps belong to different
# slots — don't merge across them. Reuses `_UW_RIGHT_STOP_RE` so any char
# that would naturally end a fill region also prevents an over-merge.
# 顿号 (`、`) is intentionally NOT in the stop set, since it's the canonical
# intra-list separator inside sample text between merged gaps.
_UW_MERGE_BLOCKER_RE = _UW_RIGHT_STOP_RE


def _extend_span_right_past_draft_marks(
    text: str, end: int, max_extension: int = 80
) -> int:
    """Walk RIGHT from `end` past any combination of whitespace and draft
    marks (`/`, `\\`, `／`, `＼`) until either a hard boundary is hit
    (sentence punctuation, unit token, amount marker — see
    `_UW_RIGHT_STOP_RE`) or no draft mark / whitespace remains.

    Returns the new end position. The walk only consumes whitespace and
    draft marks — any other character (including Chinese / letters /
    digits) ends the extension immediately, so we never absorb the next
    field's label or value.

    Used by the underscore-character scan path so trailing `/         平
    方米` regions following a `____` slot get folded into the slot — the
    underlined-whitespace path has its own (more general) extension via
    `_expand_and_merge_uw_spans`.
    """
    n = len(text)
    bound = min(n, end + max_extension)
    i = end
    saw_anything = False
    while i < bound:
        # Hit a hard right-stop boundary token? Stop here without consuming it.
        if _UW_RIGHT_STOP_RE.match(text, i):
            return i if saw_anything else end
        c = text[i]
        if c in " \t　/\\／＼":
            i += 1
            saw_anything = True
            continue
        # Anything else (Chinese, letters, digits, punctuation we didn't list)
        # ends the extension.
        return i if saw_anything else end
    return i if saw_anything else end


# Stranded draft marks: a `/` or `\\` (full or half width) right after a value
# char, followed by ≥3 whitespace chars, ending at a known boundary token or
# end-of-paragraph. Catches templates whose underscore blanks were already
# replaced by values but left the draft slash + gap behind, e.g.
#   '建筑面积：12000/        平方米；层数：6/              '
#   '批准文号：京发改〔2025〕第0128号/             '
# The two constraints — `(?<=\S)` before the mark, and a strict ≥3 whitespace
# gap followed by a boundary lookahead — keep real separators like
# '单价/数量', '2026/05/20', or '1/3' safe (no whitespace after the mark).
_STRANDED_DRAFT_MARK_RE = re.compile(
    r"(?<=\S)"
    r"(?P<mark>[/\\／＼])"
    r"(?P<gap>[ \t　]{3,})"
    r"(?=[" + re.escape(_UW_BOUNDARY_SINGLE_CHARS) + r"]"
    r"|" + r"|".join(re.escape(t) for t in _UW_BOUNDARY_TOKENS) +
    r"|$)"
)


def _strip_stranded_draft_marks(doc) -> int:
    """Walk every paragraph and erase stranded draft-mark+gap regions.
    Returns the count of erasures. Runs as a pre-save cleanup so the fix
    applies even when the scanner found no slot to attach to (template
    authors sometimes leave `数字/      平方米` style noise after pre-filled
    values, which would otherwise survive every fill)."""
    n = 0
    for paragraph in _iter_all_paragraphs(doc):
        text = paragraph.text
        if not text:
            continue
        matches = list(_STRANDED_DRAFT_MARK_RE.finditer(text))
        if not matches:
            continue
        # Apply in reverse so earlier offsets stay valid.
        for m in reversed(matches):
            _replace_paragraph_span(paragraph, m.start("mark"), m.end("gap"), "")
            n += 1
    return n


def _expand_and_merge_uw_spans(
    text: str,
    uw_spans: List[Tuple[int, int, List[Any], str]],
    max_left_distance: int = 40,
) -> List[Tuple[int, int, List[Any], str]]:
    """Expand each underlined-whitespace span to cover its full fill region.

    CN form templates commonly look like:
      "承包方式：  包工包料                          "  → sample + trailing gap
      "建筑面积：       /        平方米；层数：    /  "  → leading gap + slash + trailing gap
    The narrow `_find_underlined_whitespace_spans` only catches the long
    whitespace runs, leaving '/', sample text, and short leading gaps as
    literal content. This helper:

      - Walks LEFT from each span back to the nearest label colon ('：' or
        non-numeric ':') within `max_left_distance`, absorbing any sample
        text and short leading whitespace into the span.
      - Walks RIGHT past trailing draft marks and whitespace until hitting
        a hard punctuation boundary ('；', '。', '?', etc.), a unit token
        (平方米, 米, 元, 日, ...), an amount marker ((大写)/(小写)), or
        the start of the next uw span.
      - Merges consecutive spans whose right extension reaches the next
        span's start without crossing a hard boundary (so a "label：<gap>
        <sample>  <gap>" line becomes ONE slot covering the whole region).

    Returns expanded spans as `(left, right, merged_runs, text[left:right])`
    tuples in document order.
    """
    if not uw_spans:
        return []
    uw_spans = sorted(uw_spans, key=lambda s: s[0])
    n = len(text)

    def find_left_anchor(start_idx: int) -> int:
        # Walk LEFT from `start_idx` to find a label colon (`:` or `：`)
        # within `max_left_distance` chars. Returns the position just AFTER
        # the colon so the slot can absorb any sample text between the colon
        # and the gap.
        #
        # If we hit ANY other boundary first (sentence punctuation `。?？!！`,
        # clause separator `;；`, unit char `元/日/份/种/...`, paren `（(）)`,
        # or a multi-char unit token like `平方米`), abort expansion and
        # return `start_idx` unchanged. Those chars are hard stops — the
        # slot must not cross them, AND they don't signal "sample text
        # follows", so the slot shouldn't absorb anything past them either.
        # Previously we treated every boundary as an anchor-to-absorb-to,
        # which made `...解决。<prose>第<gap>种...` slurp the prose between
        # `。` and the gap into the slot.
        lo = max(0, start_idx - max_left_distance)
        i = start_idx
        while i > lo:
            for token in _UW_BOUNDARY_TOKENS:
                tlen = len(token)
                if i - tlen >= 0 and text[i - tlen:i] == token:
                    return start_idx
            c = text[i - 1]
            if c in (":", "：") and (i - 1 == 0 or not text[i - 2].isdigit()):
                return i
            if c in _UW_BOUNDARY_SINGLE_CHARS:
                return start_idx
            i -= 1
        return start_idx

    def find_right_extent(start_idx: int, bound: int, has_next: bool
                          ) -> Tuple[int, bool]:
        # Two-stage right walk. Returns (right, merge_with_next).
        #
        # Stage 1: consume trailing whitespace and draft marks (/ \ ／ ＼).
        #          These are ALWAYS safe to absorb — they're filler the
        #          template author left around the underlined line.
        # Stage 2: if a next UW gap exists in this paragraph AND no hard
        #          sentence boundary (. ? ! 。 ？ ！ ; ；) appears between
        #          us and it, the prose in between is sample text — merge
        #          to the next gap's start. Else stop at Stage 1's end.
        #
        # Previously the walk was "step char-by-char until any boundary in
        # the wide _UW_BOUNDARY set is hit." That treated everything between
        # the gap and the boundary as "absorbable sample text", which in
        # flowing CN prose (`<gap>的方式使用。`) silently swallowed entire
        # clauses. The structural problem: CN prose between a gap and the
        # next 。 rarely contains any boundary char, so the walk had no
        # natural stopping point. The two-stage approach makes the slot's
        # right edge the END of the underlined region itself — extending
        # only when there is positive evidence (a next gap, no sentence
        # break between) that more text belongs to the slot.
        n_text = len(text)
        i = start_idx
        while i < bound and i < n_text and text[i] in " \t　/\\／＼":
            i += 1
        extend_end = i
        if not has_next:
            return extend_end, False
        between = text[extend_end:bound]
        if _UW_MERGE_BLOCKER_RE.search(between):
            return extend_end, False
        return bound, True

    expanded: List[Tuple[int, int, List[Any], bool]] = []
    for idx, (start, end, runs, _stext) in enumerate(uw_spans):
        left = find_left_anchor(start)
        has_next = idx + 1 < len(uw_spans)
        next_start = uw_spans[idx + 1][0] if has_next else n
        right, merge_with_next = find_right_extent(end, next_start, has_next)
        expanded.append((left, right, list(runs), merge_with_next))

    merged: List[Tuple[int, int, List[Any], bool]] = []
    for left, right, runs, merge_with_next in expanded:
        if merged:
            p_left, p_right, p_runs, p_merge = merged[-1]
            if p_merge or left <= p_right:
                merged[-1] = (
                    p_left,
                    max(p_right, right),
                    p_runs + runs,
                    merge_with_next,
                )
                continue
        merged.append((left, right, runs, merge_with_next))

    return [(l, r, runs, text[l:r]) for l, r, runs, _ in merged]


def _set_paragraph_centered(paragraph) -> None:
    """Best-effort: set paragraph horizontal alignment to CENTER. Used when
    filling a border-line slot so the value sits over the middle of the line.
    """
    try:
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    except Exception:
        pass


def _table_follows_paragraph(heading_p, empty_p) -> bool:
    """True if a <w:tbl> element is the next sibling after empty_p — or after
    skipping additional empty paragraphs. The heading's parent is the body /
    cell; we walk forward in document order and bail when we hit the first
    non-empty element.
    """
    try:
        from docx.oxml.ns import qn
        elem = empty_p._element
        parent = elem.getparent()
        if parent is None:
            return False
        children = list(parent)
        try:
            idx = children.index(elem)
        except ValueError:
            return False
        TBL = qn("w:tbl")
        P = qn("w:p")
        for sib in children[idx + 1:]:
            tag = sib.tag
            if tag == TBL:
                return True
            if tag == P:
                # Look at this paragraph's text content; skip if empty
                text = "".join(
                    (t.text or "")
                    for t in sib.iter(qn("w:t"))
                )
                if text.strip():
                    return False
                # else continue scanning past more empty paragraphs
                continue
            # Section properties or other XML — keep scanning
    except Exception:
        return False
    return False


def _paragraph_is_empty(paragraph) -> bool:
    return not "".join((r.text or "") for r in paragraph.runs).strip()


def _run_is_hintlike(run) -> bool:
    """A run looks like an instruction/hint: italic, a light/grey font color,
    or a Word highlight color. Used as a gate on `_RE_REMOVAL_BANK` phrase
    matches — highlighted-but-otherwise-neutral runs containing phrases like
    "空着" / "由甲方手写" are drafting notes the template author wants the
    reader to act on, not template prose."""
    if getattr(run, "italic", False):
        return True
    try:
        if getattr(run.font, "highlight_color", None) is not None:
            return True
    except Exception:
        pass
    try:
        color = run.font.color
        if color is not None and color.rgb is not None:
            rgb = color.rgb
            # python-docx RGBColor is iterable as (r, g, b)
            r, g, b = int(rgb[0]), int(rgb[1]), int(rgb[2])
            # near-grey: channels close together and bright-ish (>= 0x80)
            if r >= 0x80 and abs(r - g) < 0x20 and abs(g - b) < 0x20 and abs(r - b) < 0x20:
                return True
            # explicit red-ish (often used for "delete me" notes) — caller
            # mostly cares about the phrase match, so we just treat red as
            # hint-eligible.
            if r >= 0xB0 and g < 0x60 and b < 0x60:
                return True
    except Exception:
        pass
    return False


def _find_hint_run(paragraph) -> Optional[Tuple[Any, int, int, str]]:
    """Find the first run in a paragraph that looks like instructional hint text.

    Returns (run, start_offset, end_offset, run_text) or None.
    """
    offset = 0
    for run in paragraph.runs:
        text = run.text or ""
        if not text:
            continue
        if _run_is_hintlike(run) and _looks_like_hint_text(text):
            return run, offset, offset + len(text), text
        offset += len(text)
    return None


def _looks_like_hint_text(text: str) -> bool:
    """True if a run's text matches the placeholder/hint phrase bank, or is
    a parenthesised / angle-bracketed snippet that reads like an instruction."""
    if not text or not text.strip():
        return False
    t = text.strip()
    if _RE_PHRASE_BANK.search(t):
        return True
    # parenthesised hints e.g. "(your bio here)" / "（请填写）" — short and starts/ends with bracket
    if (t.startswith(("(", "（", "[", "<", "《"))
        and t.endswith((")", "）", "]", ">", "》"))
        and len(t) <= 120):
        # Reject obvious citations like "(2024)" — must contain some letters
        inner = t[1:-1]
        if any(c.isalpha() for c in inner) or any("一" <= c <= "鿿" for c in inner):
            return True
    return False


# ---------------------------------------------------------------------- #
# Helpers — removal candidate scan                                        #
# ---------------------------------------------------------------------- #

def _scan_removals(doc) -> List[Dict[str, Any]]:
    """Find paragraphs that look like template instructions / meta-notes.

    Two kinds:
      - instruction_paragraph: whole paragraph matches the removal phrase bank.
      - instruction_run:       a paragraph whose hint-styled run (italic / grey
                               / red) contains removal-bank phrasing while the
                               surrounding runs are normal prose; the run is
                               flagged but the paragraph stays.
    """
    removals: List[Dict[str, Any]] = []
    counter = [0]

    def add(kind: str, text: str, reason: str, meta: Dict[str, Any]) -> None:
        rid = f"rm_{counter[0]}"
        counter[0] += 1
        removals.append({
            "id": rid,
            "kind": kind,
            "text": text,
            "reason": reason,
            "_meta": meta,
        })

    for p in _iter_all_paragraphs(doc):
        full = "".join((r.text or "") for r in p.runs)
        if not full.strip():
            continue
        m = _RE_REMOVAL_BANK.search(full)
        if not m:
            # Try styled-run-only removal: red/italic run with removal phrase
            for run in p.runs:
                rt = run.text or ""
                if not rt.strip():
                    continue
                if _run_is_hintlike(run) and _RE_REMOVAL_BANK.search(rt):
                    add(
                        "instruction_run",
                        _snippet(rt, 0, len(rt), radius=60),
                        "styled run with removal phrase",
                        {"paragraph": p, "run": run},
                    )
                    break
            continue
        # If the matched phrase covers most of the paragraph, mark it for full
        # deletion. Otherwise, treat the run containing the match as the
        # removal target — safer than nuking unrelated surrounding prose.
        cover = (m.end() - m.start()) / max(len(full.strip()), 1)
        # Parenthesized meta-notes: a paragraph whose entire content sits
        # inside one pair of parens / brackets and contains a removal marker
        # is itself the instruction — delete the whole thing regardless of
        # how short the matched phrase is. e.g.
        #   "（全文排版确保下表不跨页。此句话非正文，应删除）"
        stripped = full.strip()
        is_full_parenthetical = (
            len(stripped) >= 4
            and stripped[0] in "（(【[「『"
            and stripped[-1] in "）)】]」』"
        )
        if cover >= 0.5 or _all_runs_hintlike(p) or is_full_parenthetical:
            add(
                "instruction_paragraph",
                _snippet(full, m.start(), m.end(), radius=80),
                m.group(0)[:60],
                {"paragraph": p},
            )
        else:
            # Find the run containing the match
            offset = 0
            target_run = None
            for run in p.runs:
                rt = run.text or ""
                if offset <= m.start() < offset + len(rt):
                    target_run = run
                    break
                offset += len(rt)
            if target_run is not None:
                add(
                    "instruction_run",
                    _snippet(full, m.start(), m.end(), radius=60),
                    m.group(0)[:60],
                    {"paragraph": p, "run": target_run},
                )
    return removals


def _all_runs_hintlike(paragraph) -> bool:
    runs = [r for r in paragraph.runs if (r.text or "").strip()]
    if not runs:
        return False
    return all(_run_is_hintlike(r) for r in runs)


def _apply_removals_in_place(
    docx_path: Path, selected_keys: List[Tuple[str, str, str]], requested_ids: List[str]
) -> Dict[str, Any]:
    """Re-scan removals on the saved doc, match by (kind, reason, text), delete.

    The original inspect ids are scan-relative; if intermediate Jinja/bracket
    passes shifted paragraph structure we can't trust the cached _meta refs.
    Re-scan, match each requested key to at most one fresh removal, then apply.
    """
    try:
        from docx import Document
    except ImportError:
        return {"applied": [], "skipped": list(requested_ids)}
    try:
        doc = Document(str(docx_path))
    except Exception:
        return {"applied": [], "skipped": list(requested_ids)}

    fresh = _scan_removals(doc)
    # Index fresh by (kind, reason, text) → list of removals (in scan order)
    bucket: Dict[Tuple[str, str, str], List[Dict[str, Any]]] = defaultdict(list)
    for r in fresh:
        bucket[(r["kind"], r["reason"], r["text"])].append(r)

    applied: List[str] = []
    skipped: List[str] = []
    for rid, key in zip(requested_ids, selected_keys):
        candidates = bucket.get(key, [])
        if not candidates:
            skipped.append(rid)
            continue
        fresh_r = candidates.pop(0)
        meta = fresh_r.get("_meta", {})
        if fresh_r["kind"] == "instruction_paragraph":
            paragraph = meta["paragraph"]
            # "下表不跨页" / "keep table together" — before deleting the
            # instruction, harden the next table so Word renders it on one
            # page. Honors the author's layout intent rather than throwing it
            # away with the meta-note.
            full_text = "".join((r.text or "") for r in paragraph.runs)
            if _KEEP_TABLE_TOGETHER_RE.search(full_text):
                tbl_elem = _next_table_after(paragraph)
                if tbl_elem is not None:
                    _apply_keep_table_together(tbl_elem)
            if _drop_paragraph(paragraph):
                applied.append(rid)
            else:
                skipped.append(rid)
        elif fresh_r["kind"] == "instruction_run":
            run = meta["run"]
            try:
                run.text = ""
                # If paragraph is now effectively empty, drop it
                paragraph = meta["paragraph"]
                if _paragraph_is_empty(paragraph):
                    _drop_paragraph(paragraph)
                applied.append(rid)
            except Exception:
                skipped.append(rid)
        else:
            skipped.append(rid)

    try:
        doc.save(str(docx_path))
    except Exception:
        return {"applied": [], "skipped": list(requested_ids)}
    return {"applied": applied, "skipped": skipped}


def _drop_paragraph(paragraph) -> bool:
    """Remove a paragraph's XML element from its parent."""
    try:
        elem = paragraph._element
        parent = elem.getparent()
        if parent is None:
            return False
        parent.remove(elem)
        return True
    except Exception:
        return False


def _next_table_after(paragraph):
    """Return the next sibling <w:tbl> XML element after `paragraph`, or None.

    Walks forward through the paragraph's parent (body or cell). Used to
    locate the table that a `下表不跨页` instruction refers to.
    """
    try:
        elem = paragraph._element
        parent = elem.getparent()
    except Exception:
        return None
    if parent is None:
        return None
    ns_tbl = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tbl"
    seen = False
    for sibling in parent:
        if sibling is elem:
            seen = True
            continue
        if not seen:
            continue
        if sibling.tag == ns_tbl:
            return sibling
    return None


def _apply_keep_table_together(tbl_elem) -> bool:
    """Harden a <w:tbl> so Word renders the whole table on one page:

      - <w:cantSplit/> on every row's <w:trPr>  → no mid-row page break
      - <w:keepNext/>  on every paragraph in the table except the very
        last one → consecutive paragraphs (including paragraphs across
        rows) stay on the same page.

    Idempotent: existing cantSplit / keepNext elements are replaced.
    Returns True if at least one row was touched.
    """
    try:
        from docx.oxml import OxmlElement
    except ImportError:
        return False
    ns = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
    rows = tbl_elem.findall(f"{ns}tr")
    if not rows:
        return False
    for tr in rows:
        trPr = tr.find(f"{ns}trPr")
        if trPr is None:
            trPr = OxmlElement("w:trPr")
            # trPr must precede tc children
            tr.insert(0, trPr)
        for old in trPr.findall(f"{ns}cantSplit"):
            trPr.remove(old)
        trPr.append(OxmlElement("w:cantSplit"))
    paragraphs = list(tbl_elem.iter(f"{ns}p"))
    for p in paragraphs[:-1]:
        pPr = p.find(f"{ns}pPr")
        if pPr is None:
            pPr = OxmlElement("w:pPr")
            p.insert(0, pPr)
        for old in pPr.findall(f"{ns}keepNext"):
            pPr.remove(old)
        # keepNext should appear before most other pPr children for safety,
        # but Word tolerates either order; append for simplicity.
        pPr.append(OxmlElement("w:keepNext"))
    return True


if __name__ == "__main__":  # pragma: no cover
    import tempfile
    from docx import Document as _Doc

    skill = DocxTemplateSkill()
    tmp = Path(tempfile.mkdtemp(prefix="docxtpl_smoke_"))
    tpl_path = tmp / "tpl.docx"
    out_path = tmp / "out.docx"

    from docx.enum.text import WD_COLOR_INDEX as _WD
    d = _Doc()
    d.add_paragraph("Hello [NAME], today is [DATE].")
    d.add_paragraph("Patient Name: __________")
    d.add_paragraph("Diagnosis:")
    # Highlighted (single run) — text marked yellow that the user will replace
    p_hl1 = d.add_paragraph("Summary: ")
    r_hl1 = p_hl1.add_run("ABSTRACT_GOES_HERE")
    r_hl1.font.highlight_color = _WD.YELLOW
    p_hl1.add_run(" — keep this trailing prose bold").bold = True
    # Highlighted (multi-run span) — bold + italic split into 2 runs, both highlighted
    p_hl2 = d.add_paragraph("Title: ")
    r_hl2a = p_hl2.add_run("PART_ONE ")
    r_hl2a.font.highlight_color = _WD.GREEN
    r_hl2a.bold = True
    r_hl2b = p_hl2.add_run("PART_TWO")
    r_hl2b.font.highlight_color = _WD.GREEN
    r_hl2b.italic = True
    p_hl2.add_run(" (footer text)")
    # Angle-bracketed
    d.add_paragraph("Dear <your name>, welcome.")
    # Placeholder-phrase
    d.add_paragraph("Replace this with your bio here.")
    # Hint text (italic)
    p_hint = d.add_paragraph("Author: ")
    r_hint = p_hint.add_run("(insert your full name)")
    r_hint.italic = True
    # Section heading + empty body
    d.add_heading("Background", level=1)
    d.add_paragraph("")
    # Removal — whole-paragraph instruction
    d.add_paragraph("Delete this paragraph before submitting.")
    # Removal — Chinese instruction
    d.add_paragraph("请删除本段使用前。")
    # Table
    tbl = d.add_table(rows=2, cols=2)
    tbl.rows[0].cells[0].text = "Field"
    tbl.rows[0].cells[1].text = "Value"
    tbl.rows[1].cells[0].text = "Age"
    # tbl.rows[1].cells[1] intentionally empty
    d.save(str(tpl_path))

    inspect = skill.inspect_template(str(tpl_path))
    print("Inspect:")
    print(f"  message: {inspect['message']}")
    print("  slots:")
    for s in inspect["slots"]:
        print("   ", s)
    print("  removals:")
    for r in inspect["removals"]:
        print("   ", r)
    print("  bracket tokens:", inspect["bracket_tokens"])

    slot_vals = {s["id"]: f"<{s['kind']}>" for s in inspect["slots"]}
    removal_ids = [r["id"] for r in inspect["removals"]]
    result = skill.fill_template(
        str(tpl_path), str(out_path),
        context={"NAME": "Alice", "DATE": "2026-05-13"},
        slot_values=slot_vals,
        removal_ids=removal_ids,
    )
    print("Fill result:")
    print(f"  message: {result.get('message')}")
    print(f"  filled_keys: {result.get('filled_keys')}")
    print(f"  removals_applied: {result.get('removals_applied')}")
    print(f"  removals_skipped: {result.get('removals_skipped')}")
    if result.get('slot_fill'):
        print(f"  slot_fill: {result['slot_fill']}")
    print("Output content:")
    out_doc = _Doc(str(out_path))
    for p in out_doc.paragraphs:
        print(" ", repr(p.text))
    # Verify the highlighted slots had their highlight cleared
    print("Highlight verification (should all be False/None):")
    for p in out_doc.paragraphs:
        for r in p.runs:
            if r.text and "<highlighted>" in r.text:
                hl = r.font.highlight_color
                print(
                    f"   run={r.text!r:30s} highlight={hl!r:18s} "
                    f"bold={r.bold} italic={r.italic}"
                )
    # Verify surrounding formatting preserved
    print("Surrounding formatting (should keep bold/italic on neighbors):")
    for p in out_doc.paragraphs:
        if "keep this trailing prose bold" in (p.text or ""):
            for r in p.runs:
                print(f"   run={r.text!r:50s} bold={r.bold}")
        if "(footer text)" in (p.text or ""):
            for r in p.runs:
                print(f"   run={r.text!r:50s} bold={r.bold} italic={r.italic}")
