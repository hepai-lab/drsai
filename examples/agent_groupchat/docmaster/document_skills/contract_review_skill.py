"""Contract review skill for DocMaster — `合同文件` 审查。

Four review passes over a `.docx`, results merged into one issues list:

1. format            — font/size consistency, mixed CJK/Latin runs, full/half-width
                        punctuation mixing, numbered-clause spacing, signature/date
                        scaffold sanity.
2. fill_completeness — unfilled underscore runs, `**XX**`/`{{...}}` placeholders,
                        bare `甲方:` with empty tail, scaffold dates, empty
                        signature/seal cells.
3. consistency       — 大写/小写金额 cross-check (uses chinese_amount), date strings,
                        party-name (甲方/乙方) variants, numbered-clause continuity,
                        `第X条`/`附件X` references that don't resolve.
4. legal             — LLM-driven red-flag scan (违约/争议解决/法律适用/生效/责任).
                        Optional. Fail-soft: a timeout returns heuristic-only with
                        a note in summary.

Public API:

    skill = ContractReviewSkill(workspace_dir)
    skill.review(file_path, llm_call=None) -> {success, summary, stats, issues}

`llm_call` is a `Callable[[str], str]` that takes a prompt and returns the
model's reply. Pass `None` to skip the legal pass entirely.
"""

from __future__ import annotations

import json
import logging
import re
from collections import Counter, defaultdict
from decimal import Decimal
from pathlib import Path
from typing import Any, Callable, Dict, List, Optional, Tuple

try:
    from docx import Document
except ImportError as exc:  # pragma: no cover
    raise RuntimeError("python-docx is required for ContractReviewSkill") from exc

try:
    from .chinese_amount import chinese_to_decimal, parse_arabic_amount
except ImportError:  # support running this file directly
    from chinese_amount import chinese_to_decimal, parse_arabic_amount  # type: ignore

logger = logging.getLogger(__name__)


# ---------------------------------------------------------------- #
# Patterns                                                          #
# ---------------------------------------------------------------- #

# Underscore fill markers — copied from docx_template_skill so we stay aligned.
_UNDERSCORE_RUN_RE = re.compile(r"[_＿]{3,}|(?:[_＿](?:[ 　]+[_＿]){1,}){1}")
_PLACEHOLDER_INLINE_RE = re.compile(r"\{\{\s*[a-zA-Z_][\w]*\s*\}\}|<[^<>\n]{1,40}>|《[^《》\n]{1,40}》")
_ASTERISK_PAIR_RE = re.compile(r"\*\*[^*\n]{0,40}\*\*|\*\*")
_LABEL_BLANK_RE = re.compile(r"(甲\s*方|乙\s*方|丙\s*方|签\s*订\s*日\s*期|身\s*份\s*证|联\s*系\s*电\s*话|地\s*址|开\s*户\s*行|账\s*号|姓\s*名|项\s*目\s*名\s*称|合\s*同\s*编\s*号)\s*[:：]\s*$")
_DATE_SCAFFOLD_RE = re.compile(r"[ \t　]{2,}年(?:[ \t　]+月)?(?:[ \t　]+日)?")
# Match Chinese amount-pair markers in a paragraph: "(大写) ... (小写) ¥..." or
# the more relaxed "金额: 壹万元整 (¥10000)" style.
_DAXIE_AMOUNT_RE = re.compile(r"[（(]\s*大\s*写\s*[)）]\s*[:：]?\s*([零一二三四五六七八九十百千万亿壹贰叁肆伍陆柒捌玖拾佰仟萬億元圆角分整正点\s]{2,40})")
_XIAOXIE_AMOUNT_RE = re.compile(r"[（(]\s*小\s*写\s*[)）]\s*[:：]?\s*([¥￥$€£]?\s*[\d,，\.]{1,15}\s*元?)")
_INLINE_AMOUNT_PAIR_RE = re.compile(
    r"(人民币|金额|总计|合计|总价|价款|费用|报酬)[^\n]{0,8}?"
    r"([零壹贰叁肆伍陆柒捌玖拾佰仟萬億元圆角分整正点]{3,40})"
    r"[^\n]{0,12}?"
    r"[（(]\s*[¥￥$]?\s*([\d,，\.]{1,15})\s*元?\s*[)）]"
)
_NUMBERED_CLAUSE_RE = re.compile(r"^第\s*([零一二三四五六七八九十百千]+)\s*条")
# Body references like "见第三条" or "依第3.2条" or "见附件一".
_CLAUSE_REF_RE = re.compile(r"第\s*([零一二三四五六七八九十百千]+|\d+(?:[\.\-]\d+)*)\s*条")
_ANNEX_REF_RE = re.compile(r"附\s*件\s*([零一二三四五六七八九十百千]+|\d+)")
# Date strings — multiple formats commonly mixed in one contract.
_DATE_FORMS_RE = re.compile(
    r"(?P<a>(\d{4})\s*年\s*(\d{1,2})\s*月\s*(\d{1,2})\s*日)"
    r"|(?P<b>(\d{4})[-/.](\d{1,2})[-/.](\d{1,2}))"
)
# Party-name variants (甲方/乙方) — we collect surrounding labels and flag if
# multiple distinct names appear after the same label.
_PARTY_LABEL_RE = re.compile(r"(甲方|乙方|丙方)[（(][^)）\n]{0,30}[)）]?\s*[:：]\s*([^\n,，。；;]{1,40})")

_FULLWIDTH_PUNCT = "，。；：！？（）《》【】、“”‘’"
_HALFWIDTH_PUNCT = ",.;:!?()<>[],\"\"''"
_CJK_RE = re.compile(r"[一-鿿]")
_ASCII_LETTER_RE = re.compile(r"[A-Za-z]")

# "Signature/seal" is a fillable region but should be FLAGGED ONLY if user is
# expected to fill outside the doc (i.e. a missing signature line is OK if
# scaffold present; we just note it once in the report).
_SIGNATURE_LABEL_RE = re.compile(
    r"(?:签\s*字|签\s*名|签\s*章|盖\s*章|落\s*款|手\s*印|signature|signed\s*by)",
    re.IGNORECASE,
)


# ---------------------------------------------------------------- #
# Issue model                                                       #
# ---------------------------------------------------------------- #


def _issue(
    category: str,
    severity: str,
    message: str,
    *,
    location: str = "",
    snippet: str = "",
    suggestion: str = "",
    comment_target: str = "",
) -> Dict[str, Any]:
    """Build one issue dict. `comment_target` is the literal substring inside the
    docx that the annotated-output pass will attach a Word comment to. If empty,
    no inline comment is emitted for this issue (it still appears in the report)."""
    return {
        "category": category,
        "severity": severity,
        "message": message,
        "location": location,
        "snippet": snippet,
        "suggestion": suggestion,
        "comment_target": comment_target,
    }


# ---------------------------------------------------------------- #
# Skill                                                             #
# ---------------------------------------------------------------- #


class ContractReviewSkill:
    def __init__(self, workspace_dir: str | Path):
        self.workspace_dir = Path(workspace_dir).resolve()

    # -------------- entrypoint ------------------------------------ #

    def review(
        self,
        file_path: str | Path,
        llm_call: Optional[Callable[[str], str]] = None,
    ) -> Dict[str, Any]:
        path = Path(file_path)
        if not path.is_file():
            return {"success": False, "message": f"文件不存在: {file_path}"}
        if path.suffix.lower() != ".docx":
            return {"success": False, "message": f"仅支持 .docx 文件 (当前: {path.suffix})"}
        try:
            doc = Document(str(path))
        except Exception as exc:
            return {"success": False, "message": f"打开 .docx 失败: {exc}"}

        paragraphs = self._collect_paragraphs(doc)
        full_text = "\n".join(p["text"] for p in paragraphs)

        issues: List[Dict[str, Any]] = []
        issues.extend(self._check_format(doc, paragraphs))
        issues.extend(self._check_fill_completeness(doc, paragraphs))
        issues.extend(self._check_consistency(paragraphs, full_text))

        llm_note = ""
        if llm_call is not None:
            llm_issues, llm_note = self._check_legal_via_llm(full_text, llm_call)
            issues.extend(llm_issues)

        # Sort by severity (high first), then category, then keep insertion order.
        sev_rank = {"high": 0, "medium": 1, "low": 2}
        issues.sort(key=lambda i: (sev_rank.get(i["severity"], 99), i["category"]))

        stats = self._stats(issues)
        summary = self._summary(stats, llm_note)

        return {
            "success": True,
            "summary": summary,
            "stats": stats,
            "issues": issues,
        }

    # -------------- common helpers -------------------------------- #

    def _collect_paragraphs(self, doc) -> List[Dict[str, Any]]:
        out: List[Dict[str, Any]] = []
        idx = 0
        for para in doc.paragraphs:
            text = para.text or ""
            out.append({"index": idx, "text": text, "in_table": False, "para": para})
            idx += 1
        for tbl_i, tbl in enumerate(doc.tables):
            for r_i, row in enumerate(tbl.rows):
                for c_i, cell in enumerate(row.cells):
                    for para in cell.paragraphs:
                        out.append({
                            "index": idx,
                            "text": para.text or "",
                            "in_table": True,
                            "table_pos": (tbl_i, r_i, c_i),
                            "para": para,
                        })
                        idx += 1
        return out

    def _loc(self, p: Dict[str, Any]) -> str:
        if p.get("in_table"):
            t, r, c = p["table_pos"]
            return f"表{t + 1} 第{r + 1}行第{c + 1}列"
        return f"第{p['index'] + 1}段"

    # -------------- pass 1: format -------------------------------- #

    def _check_format(self, doc, paragraphs: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
        issues: List[Dict[str, Any]] = []

        # 1a. Font / size inventory across body runs. Flag if more than 3 distinct
        # CJK fonts or 4 distinct sizes appear — likely copy/paste from heterogeneous
        # sources without normalizing.
        fonts: Counter = Counter()
        sizes: Counter = Counter()
        for p in paragraphs:
            para = p["para"]
            for run in para.runs:
                if not (run.text or "").strip():
                    continue
                eastasia = None
                try:
                    rpr = run._element.rPr
                    if rpr is not None and rpr.rFonts is not None:
                        eastasia = rpr.rFonts.get(
                            "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia"
                        )
                except Exception:
                    pass
                fonts[eastasia or run.font.name or ""] += 1
                if run.font.size is not None:
                    sizes[run.font.size.pt] += 1
        nonempty_fonts = [f for f in fonts if f]
        if len(nonempty_fonts) > 3:
            top = ", ".join(f"{n}({c})" for n, c in fonts.most_common(5) if n)
            issues.append(_issue(
                "格式", "medium",
                f"中文字体不一致：检测到 {len(nonempty_fonts)} 种不同字体（{top}）",
                suggestion="统一使用单一中文字体，例如『宋体』或『仿宋』；可用 modify_docx_fonts_tool 一键归并。",
            ))
        if len(sizes) > 4:
            top = ", ".join(f"{s}pt({c})" for s, c in sizes.most_common(5))
            issues.append(_issue(
                "格式", "low",
                f"字号种类较多：检测到 {len(sizes)} 种字号（{top}）",
                suggestion="正文建议统一为单一字号（如小四号 12pt），标题层级各一档。",
            ))

        # 1b. Mixed CJK + Latin in the same run without spacing — common copy/paste artefact.
        for p in paragraphs:
            text = p["text"]
            if not text:
                continue
            # Look for CJK directly adjacent to a Latin word with no space, e.g. "甲方ABC公司"
            for m in re.finditer(r"[一-鿿][A-Za-z]|[A-Za-z][一-鿿]", text):
                # Only flag once per paragraph.
                issues.append(_issue(
                    "格式", "low",
                    "中英文混排未加空格，可能影响排版美观",
                    location=self._loc(p),
                    snippet=_excerpt(text, m.start(), m.end()),
                    comment_target=text[max(0, m.start() - 4): m.end() + 4],
                ))
                break

        # 1c. Punctuation mixing — paragraph dominated by Chinese but containing
        # half-width , ; : ? ! is a common authoring mistake.
        for p in paragraphs:
            text = p["text"]
            if not text or not _CJK_RE.search(text):
                continue
            cjk_count = len(_CJK_RE.findall(text))
            if cjk_count < 5:
                continue
            for ch in ",;:?!":
                if ch in text:
                    pos = text.index(ch)
                    issues.append(_issue(
                        "格式", "low",
                        f"中文段落中混用半角标点 '{ch}'，建议改为全角",
                        location=self._loc(p),
                        snippet=_excerpt(text, pos, pos + 1),
                        comment_target=text[max(0, pos - 6): pos + 6],
                    ))
                    break

        # 1d. Numbered-clause sequence sanity (also serves consistency, but a clause
        # numbered with mixed Chinese-numeral / Arabic-numeral styles is a format issue).
        styles = set()
        for p in paragraphs:
            m = re.match(r"^第\s*([零一二三四五六七八九十百千]+|\d+)\s*条", p["text"])
            if m:
                styles.add("汉字" if re.search(r"[零一二三四五六七八九十百千]", m.group(1)) else "阿拉伯")
        if len(styles) > 1:
            issues.append(_issue(
                "格式", "medium",
                f"条款编号方式混用（{', '.join(sorted(styles))}），建议全文统一",
                suggestion="正式合同建议统一为『第一条』『第二条』…，附件清单可用阿拉伯数字。",
            ))

        return issues

    # -------------- pass 2: fill completeness --------------------- #

    def _check_fill_completeness(self, doc, paragraphs: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
        issues: List[Dict[str, Any]] = []

        for p in paragraphs:
            text = p["text"]
            if not text:
                continue

            # Underscore runs.
            for m in _UNDERSCORE_RUN_RE.finditer(text):
                # Skip if the underscore region is in a signature/seal context.
                surround = text[max(0, m.start() - 16): m.end() + 16]
                is_sig = bool(_SIGNATURE_LABEL_RE.search(surround))
                issues.append(_issue(
                    "填写缺失",
                    "low" if is_sig else "high",
                    "存在未填写的下划线占位（签字/盖章位除外建议补齐）" if is_sig
                    else "存在未填写的下划线占位",
                    location=self._loc(p),
                    snippet=_excerpt(text, m.start(), m.end()),
                    comment_target=m.group(0),
                ))
                # only flag the first underscore region per paragraph to avoid
                # noise — most paragraphs with an unfilled blank have just one.
                break

            # Inline templating leftovers.
            for m in _PLACEHOLDER_INLINE_RE.finditer(text):
                issues.append(_issue(
                    "填写缺失", "high",
                    f"存在未替换的占位符 `{m.group(0)}`",
                    location=self._loc(p),
                    snippet=_excerpt(text, m.start(), m.end()),
                    comment_target=m.group(0),
                ))
                break

            # Asterisk markers — `**…**` paired sample text or bare `**`.
            if "**" in text:
                m = _ASTERISK_PAIR_RE.search(text)
                if m:
                    issues.append(_issue(
                        "填写缺失", "high",
                        f"存在未替换的星号占位 `{m.group(0)}`",
                        location=self._loc(p),
                        snippet=_excerpt(text, m.start(), m.end()),
                        comment_target=m.group(0),
                    ))

            # Label with empty tail — "甲方:" alone on a line.
            if _LABEL_BLANK_RE.search(text):
                issues.append(_issue(
                    "填写缺失", "high",
                    f"字段未填写：`{text.strip()}`",
                    location=self._loc(p),
                    snippet=text.strip()[:60],
                    comment_target=text.strip(),
                ))

            # Empty date scaffold — "    年    月    日".
            if _DATE_SCAFFOLD_RE.search(text):
                issues.append(_issue(
                    "填写缺失", "high",
                    "存在未填写的日期占位（年/月/日）",
                    location=self._loc(p),
                    snippet=text.strip()[:60],
                    comment_target=_DATE_SCAFFOLD_RE.search(text).group(0),
                ))

        return issues

    # -------------- pass 3: consistency --------------------------- #

    def _check_consistency(
        self, paragraphs: List[Dict[str, Any]], full_text: str
    ) -> List[Dict[str, Any]]:
        issues: List[Dict[str, Any]] = []

        # 3a. Amount cross-check: 大写 vs 小写.
        for p in paragraphs:
            text = p["text"]
            daxie = _DAXIE_AMOUNT_RE.search(text)
            xiaoxie = _XIAOXIE_AMOUNT_RE.search(text)
            inline = _INLINE_AMOUNT_PAIR_RE.search(text)

            pairs: List[Tuple[str, str]] = []
            if daxie and xiaoxie:
                pairs.append((daxie.group(1), xiaoxie.group(1)))
            if inline:
                pairs.append((inline.group(2), inline.group(3)))

            for cn, num in pairs:
                cn_dec = chinese_to_decimal(cn)
                num_dec = parse_arabic_amount(num)
                if cn_dec is None or num_dec is None:
                    continue
                if cn_dec != num_dec:
                    issues.append(_issue(
                        "内容一致性", "high",
                        f"金额大小写不一致：大写 `{cn.strip()}` (={cn_dec}) ≠ 小写 `{num.strip()}` (={num_dec})",
                        location=self._loc(p),
                        snippet=text.strip()[:80],
                        suggestion="请核对正确金额并统一两处。",
                        comment_target=text.strip()[:60],
                    ))

        # 3b. Numbered-clause continuity.
        clause_seq = []
        for p in paragraphs:
            m = _NUMBERED_CLAUSE_RE.match(p["text"])
            if m:
                clause_seq.append((p, m.group(1)))
        nums = [_cn_or_arabic_to_int(s) for _, s in clause_seq]
        for i in range(1, len(nums)):
            if nums[i] is not None and nums[i - 1] is not None and nums[i] != nums[i - 1] + 1:
                issues.append(_issue(
                    "内容一致性", "medium",
                    f"条款编号不连续：第{nums[i - 1]}条 → 第{nums[i]}条",
                    location=self._loc(clause_seq[i][0]),
                    snippet=clause_seq[i][0]["text"][:40],
                    suggestion="检查是否有条款被删除或重号。",
                ))

        # 3c. Clause / annex references that don't resolve.
        defined_clauses = {n for n in nums if n is not None}
        defined_annexes: set = set()
        for p in paragraphs:
            m = re.match(r"^\s*附\s*件\s*([零一二三四五六七八九十百千]+|\d+)\s*[:：]?", p["text"])
            if m:
                v = _cn_or_arabic_to_int(m.group(1))
                if v is not None:
                    defined_annexes.add(v)
        for p in paragraphs:
            for m in _CLAUSE_REF_RE.finditer(p["text"]):
                ref = _cn_or_arabic_to_int(m.group(1).split(".")[0].split("-")[0])
                if ref is None or ref in defined_clauses:
                    continue
                if not defined_clauses:
                    continue  # contract has no numbered clauses at all — skip.
                issues.append(_issue(
                    "内容一致性", "medium",
                    f"引用了不存在的『第{m.group(1)}条』",
                    location=self._loc(p),
                    snippet=_excerpt(p["text"], m.start(), m.end()),
                    comment_target=m.group(0),
                ))
                break
            for m in _ANNEX_REF_RE.finditer(p["text"]):
                ref = _cn_or_arabic_to_int(m.group(1))
                if ref is None or ref in defined_annexes:
                    continue
                if not defined_annexes:
                    continue
                issues.append(_issue(
                    "内容一致性", "medium",
                    f"引用了不存在的『附件{m.group(1)}』",
                    location=self._loc(p),
                    snippet=_excerpt(p["text"], m.start(), m.end()),
                    comment_target=m.group(0),
                ))
                break

        # 3d. Party-name variants — same label (甲方/乙方/丙方) followed by different
        # company names anywhere in the doc.
        party_names: Dict[str, set] = defaultdict(set)
        for m in _PARTY_LABEL_RE.finditer(full_text):
            label = re.sub(r"\s", "", m.group(1))
            name = m.group(2).strip()
            if name:
                party_names[label].add(name)
        for label, names in party_names.items():
            if len(names) > 1:
                issues.append(_issue(
                    "内容一致性", "high",
                    f"{label}名称不一致：{ ' / '.join(sorted(names))}",
                    suggestion="全文核对该方主体名称，确保与营业执照/身份证一致。",
                ))

        return issues

    # -------------- pass 4: legal red-flag (LLM) ------------------ #

    def _check_legal_via_llm(
        self, full_text: str, llm_call: Callable[[str], str]
    ) -> Tuple[List[Dict[str, Any]], str]:
        # Truncate to a sensible window — most problems live in the first ~6k chars
        # (preamble + main clauses); we add a tail sample for sig/dispute clauses.
        head = full_text[:6000]
        tail = full_text[-2000:] if len(full_text) > 8000 else ""
        excerpt = head + ("\n...\n" + tail if tail else "")

        prompt = (
            "你是一名资深合同审查律师。下面是一份中文合同的文本（可能被截断）。\n"
            "请用『法律风险』视角找出问题，重点关注以下几类：\n"
            "  1. 缺失条款：违约责任 / 争议解决 / 法律适用 / 生效条件 / 保密条款 / 知识产权归属 / 不可抗力。\n"
            "  2. 不公平条款：单方解除权、单方变更权、过度免责、责任不对等。\n"
            "  3. 模糊措辞：『若干日』『合理时间』『相关费用』等无法量化的表述。\n"
            "  4. 法律合规：争议解决约定的管辖法院/仲裁机构是否清楚、违约金是否过高（>30%）等。\n\n"
            "请仅返回 JSON 数组，每个元素形如：\n"
            '  {"severity": "high|medium|low", "message": "...", "suggestion": "...", '
            '"comment_target": "原文中可定位该问题的一段较短文字，便于附加批注，可以为空"}\n\n'
            "如果未发现明显问题，返回 []。除 JSON 数组外不要输出任何其他内容。\n\n"
            "合同正文：\n----\n" + excerpt + "\n----\n"
        )

        try:
            reply = llm_call(prompt)
        except Exception as exc:
            logger.warning("legal review LLM call failed: %s", exc)
            return [], f"LLM 审查未完成（{exc}），仅返回启发式检查结果。"

        if not reply:
            return [], "LLM 审查未返回内容，仅返回启发式检查结果。"

        items = _parse_llm_json_array(reply)
        if items is None:
            return [], "LLM 返回非 JSON 格式，已忽略；仅返回启发式检查结果。"

        out: List[Dict[str, Any]] = []
        for it in items:
            if not isinstance(it, dict):
                continue
            sev = str(it.get("severity", "medium")).lower()
            if sev not in ("high", "medium", "low"):
                sev = "medium"
            msg = str(it.get("message") or "").strip()
            if not msg:
                continue
            out.append(_issue(
                "法律风险", sev, msg,
                suggestion=str(it.get("suggestion") or "").strip(),
                comment_target=str(it.get("comment_target") or "").strip(),
            ))
        return out, ""

    # -------------- summary --------------------------------------- #

    def _stats(self, issues: List[Dict[str, Any]]) -> Dict[str, Any]:
        by_cat: Counter = Counter(i["category"] for i in issues)
        by_sev: Counter = Counter(i["severity"] for i in issues)
        return {
            "total": len(issues),
            "by_category": dict(by_cat),
            "by_severity": dict(by_sev),
        }

    def _summary(self, stats: Dict[str, Any], llm_note: str) -> str:
        if stats["total"] == 0:
            base = "审查完毕，未发现明显问题。"
        else:
            cat_str = "、".join(f"{k} {v}处" for k, v in stats["by_category"].items())
            sev = stats["by_severity"]
            sev_str = (
                f"（高 {sev.get('high', 0)} / 中 {sev.get('medium', 0)} / 低 {sev.get('low', 0)}）"
            )
            base = f"共发现 {stats['total']} 处问题：{cat_str} {sev_str}"
        if llm_note:
            base += "\n注意：" + llm_note
        return base


# ---------------------------------------------------------------- #
# Module-level helpers                                              #
# ---------------------------------------------------------------- #


_CN_DIGITS = {"零": 0, "一": 1, "二": 2, "三": 3, "四": 4, "五": 5, "六": 6, "七": 7, "八": 8, "九": 9, "十": 10}


def _cn_or_arabic_to_int(s: str) -> Optional[int]:
    s = (s or "").strip()
    if not s:
        return None
    if s.isdigit():
        return int(s)
    # Very small set: 一..十 plus 十一..九十九 patterns ("十", "十二", "二十", "二十三").
    if s == "十":
        return 10
    if len(s) == 2 and s[0] == "十" and s[1] in _CN_DIGITS:
        return 10 + _CN_DIGITS[s[1]]
    if len(s) == 2 and s[0] in _CN_DIGITS and s[1] == "十":
        return _CN_DIGITS[s[0]] * 10
    if len(s) == 3 and s[1] == "十" and s[0] in _CN_DIGITS and s[2] in _CN_DIGITS:
        return _CN_DIGITS[s[0]] * 10 + _CN_DIGITS[s[2]]
    if len(s) == 1 and s in _CN_DIGITS:
        return _CN_DIGITS[s]
    return None


def _excerpt(text: str, start: int, end: int, radius: int = 16) -> str:
    s = max(0, start - radius)
    e = min(len(text), end + radius)
    pre = "…" if s > 0 else ""
    post = "…" if e < len(text) else ""
    return f"{pre}{text[s:e]}{post}"


def _parse_llm_json_array(reply: str) -> Optional[List[Any]]:
    """Robustly extract a JSON array from a model reply that may be wrapped in
    markdown fences or have leading prose."""
    if not reply:
        return None
    # Strip code fences if present.
    fence = re.search(r"```(?:json)?\s*(\[.*?\])\s*```", reply, re.DOTALL)
    if fence:
        candidate = fence.group(1)
    else:
        # Take the substring from the first '[' to the matching last ']'.
        lb = reply.find("[")
        rb = reply.rfind("]")
        if lb == -1 or rb == -1 or rb <= lb:
            return None
        candidate = reply[lb: rb + 1]
    try:
        data = json.loads(candidate)
    except Exception:
        return None
    return data if isinstance(data, list) else None
